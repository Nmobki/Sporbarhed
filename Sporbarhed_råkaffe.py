#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import urllib
from datetime import datetime
import pandas as pd
from sqlalchemy import create_engine
import pyodbc
import docx
from docx.shared import Inches
import openpyxl
import networkx as nx


# =============================================================================
# Define functions
# =============================================================================

# Get section name for section from query
def get_section_name(section):
    df_temp_sections = df_sections.loc[df_sections['Sektion'] == section]
    x = df_temp_sections['Sektion navn'].iloc[0]
    if len(x) == 0 or len(x) > 31:
        return 'Sektion ' + str(section)
    else:
        return x

# Find statuscode for section log
def get_section_status_code(dataframe):
    if len(dataframe) == 0:
        return 1 # Empty dataframe
    else:
        return 99 # Continue

# Write into section log
def section_log_insert(section, statuscode, errorcode=None):
    df = pd.DataFrame(data={'Forespørgsels_id':req_id,
                            'Sektion':section,
                            'Statuskode':statuscode,
                            'Fejlkode_script':str(errorcode)}
                      , index=[0])
    df.to_sql('Sporbarhed_sektion_log', con=engine_04, schema='trc', if_exists='append', index=False)

# Write dataframe into Excel sheet
def insert_dataframe_into_excel (dataframe, sheetname, include_index):
    dataframe.to_excel(excel_writer, sheet_name=sheetname, index=include_index)

# Convert list into string for SQL IN operator
def string_to_sql(list_with_values):
    if len(list_with_values) == 0:
        return ''
    else:
        return "'{}'".format("','".join(list_with_values))

def number_format(value, number_type):
    try:
        if number_type == 'dec_2':
            return f'{round(value,2):,}'.replace(',', ';').replace('.', ',').replace(';', '.')
        elif number_type == 'dec_1':
            return f'{round(value,1):,}'.replace(',', ';').replace('.', ',').replace(';', '.')
        elif number_type == 'dec_0':
            return f'{int(round(value,0)):,}'.replace(',', ';').replace('.', ',').replace(';', '.')
        elif number_type == 'pct_2':
            return f'{round(value,4):.2%}'.replace(',', ';').replace('.', ',').replace(';', '.')
        elif number_type == 'pct_0':
            return f'{round(value,2):.0%}'.replace(',', ';').replace('.', ',').replace(';', '.')
        else:
            return value
    except:
        return value

# Prevent division by zero error
def zero_division(nominator, denominator, zero_return):
    dict_div = {'None':None,'Zero':0}
    if denominator in [0,None]:
        return dict_div[zero_return]
    else:
        return nominator / denominator

# Write into dbo.log
def log_insert(event, note):
    dict_log = {'Note': note
                ,'Event': event}
    pd.DataFrame(data=dict_log, index=[0]).to_sql('Log', con=engine_04, schema='dbo', if_exists='append', index=False)

# Get info from item table in Navision
def get_nav_item_info(item_no, field):
    if item_no in df_nav_items['Nummer'].tolist():
        df_temp = df_nav_items[df_nav_items['Nummer'] == item_no]
        return df_temp[field].iloc[0]
    else:
        return None

# Get info from assembly and production orders in Navision
def get_nav_order_info(order_no):
    if order_no in df_nav_order_info['Ordrenummer'].tolist():
        df_temp = df_nav_order_info[df_nav_order_info['Ordrenummer'] == order_no]
        return df_temp['Varenummer'].iloc[0]
    else:
        return None

# Convert placeholder values from dataframe to empty string for Word document
def convert_placeholders_word(string):
    if string in ['None','nan','NaT']:
        return ''
    else:
        return string

# Add dataframe to word document
def add_section_to_word(dataframe, section, pagebreak, rows_to_bold):
    # Add section header
    doc.add_heading(section, 1)
    # Add a table with an extra row for headers
    table = doc.add_table(dataframe.shape[0]+1, dataframe.shape[1])
    table.style = 'Table Grid'
    # Add headers to top row
    for i in range(dataframe.shape[-1]):
        table.cell(0,i).text = dataframe.columns[i]
    # Add data from dataframe to the table, replace supposed blank cells using function
    for x in range(dataframe.shape[0]):
        for y in range(dataframe.shape[-1]):
            table.cell(x+1,y).text =  convert_placeholders_word(str(dataframe.values[x,y]))
    # Bold total row if it exists
    for y in rows_to_bold:
        for x in range(dataframe.shape[1]):
            table.rows[y].cells[x].paragraphs[0].runs[0].font.bold = True
    # Add page break
    if pagebreak:
        doc.add_page_break()


# =============================================================================
# Variables for query connections
# =============================================================================
server_04 = 'sqlsrv04'
db_04 = 'BKI_Datastore'
con_04 = pyodbc.connect(f'DRIVER=SQL Server;SERVER={server_04};DATABASE={db_04};autocommit=True')
params_04 = urllib.parse.quote_plus(f'DRIVER=SQL Server Native Client 11.0;SERVER={server_04};DATABASE={db_04};Trusted_Connection=yes')
engine_04 = create_engine(f'mssql+pyodbc:///?odbc_connect={params_04}')
cursor_04 = con_04.cursor()

server_nav = r'SQLSRV03\NAVISION'
db_nav = 'NAV100-DRIFT'
con_nav = pyodbc.connect(f'DRIVER=SQL Server;SERVER={server_nav};DATABASE={db_nav};Trusted_Connection=yes')
params_nav = urllib.parse.quote_plus(f'DRIVER=SQL Server Native Client 11.0;SERVER={server_nav};DATABASE={db_nav};Trusted_Connection=yes')
engine_nav = create_engine(f'mssql+pyodbc:///?odbc_connect={params_nav}')

server_probat = '192.168.125.161'
db_probat = 'BKI_IMP_EXP'
con_probat = pyodbc.connect(f'DRIVER=SQL Server;SERVER={server_probat};DATABASE={db_probat};uid=bki_read;pwd=Probat2016')
params_probat = urllib.parse.quote_plus(f'DRIVER=SQL Server Native Client 11.0;SERVER={server_probat};DATABASE={db_probat};Trusted_Connection=yes')
engine_probat = create_engine(f'mssql+pyodbc:///?odbc_connect={params_probat}')

# =============================================================================
# Read data from request
# =============================================================================
query_ds_request =  """ SELECT TOP 1 [Id] ,[Forespørgselstype],[Rapport_modtager]
                    ,[Referencenummer] ,[Note_forespørgsel] ,[Modtagelse]  ,[Ordrerelationstype]
                    FROM [trc].[Sporbarhed_forespørgsel]
                    WHERE [Forespørgsel_igangsat] IS NULL """
df_request = pd.read_sql(query_ds_request, con_04)

# Exit script if no request data is found
if len(df_request) == 0:
    raise SystemExit(0)

# =============================================================================
# Set request variables
# =============================================================================
req_type = df_request.loc[0, 'Forespørgselstype']
req_reference_no = df_request.loc[0, 'Referencenummer']
req_recipients = df_request.loc[0, 'Rapport_modtager']
req_note = df_request.loc[0, 'Note_forespørgsel']
req_id = df_request.loc[0, 'Id']
req_modtagelse = df_request.loc[0, 'Modtagelse']
req_ordrelationstype = df_request.loc[0, 'Ordrerelationstype']

script_name = 'Sporbarhed_samlet.py'
timestamp = datetime.now().strftime('%d-%m-%Y %H:%M:%S')
orders_top_level = [req_reference_no]
orders_related = []

# =============================================================================
# Update request that it is initiated and write into log
# =============================================================================
# =============================================================================
# cursor_04.execute(f"""UPDATE [trc].[Sporbarhed_forespørgsel]
#                   SET [Forespørgsel_igangsat] = getdate()
#                   WHERE [Id] = {req_id} """)
# cursor_04.commit()
# =============================================================================
# log_insert(script_name, f'Request id: {req_id} initiated')

# =============================================================================
# Variables for files generated
# =============================================================================
filepath = r'\\filsrv01\BKI\11. Økonomi\04 - Controlling\NMO\4. Kvalitet\Sporbarhedstest\Tests via PowerApps'
file_name = f'Rapport_{req_reference_no}_{req_id}'

doc = docx.Document()
doc.add_heading(f'Rapport for produktionsordre {req_reference_no}',0)
doc.sections[0].header.paragraphs[0].text = f'{script_name}'
doc.sections[0].footer.paragraphs[0].text = f'{timestamp}'
doc.sections[0].page_width = docx.shared.Mm(297)
doc.sections[0].page_height = docx.shared.Mm(210)
doc.sections[0].top_margin = docx.shared.Mm(15)
doc.sections[0].bottom_margin = docx.shared.Mm(15)
doc.sections[0].left_margin = docx.shared.Mm(10)
doc.sections[0].right_margin = docx.shared.Mm(10)
doc.sections[0].orientation = docx.enum.section.WD_ORIENT.LANDSCAPE

doc_name = f'{file_name}.docx'
path_file_doc = filepath + r'\\' + doc_name

wb = openpyxl.Workbook()
wb_name = f'{file_name}.xlsx'
path_file_wb = filepath + r'\\' + wb_name
excel_writer = pd.ExcelWriter(path_file_wb, engine='xlsxwriter')

png_relations_name = f'{file_name}.png'
path_png_relations = filepath + r'\\' + png_relations_name

# =============================================================================
# Read setup for section for reporttypes. NAV querys with NOLOCK to prevent deadlocks
# =============================================================================
query_ds_reporttypes =  f"""SELECT SRS.[Sektion], SS.[Beskrivelse] AS [Sektion navn]
                       FROM [trc].[Sporbarhed_rapport_sektion] AS SRS
					   INNER JOIN [trc].[Sporbarhed_sektion] AS SS
					   ON SRS.[Sektion] = SS.[Id]
                       WHERE [Forespørgselstype] = {req_type} """
df_sections = pd.read_sql(query_ds_reporttypes, con_04)

# Query for Navision items, used for adding information to item numbers not queried directly from Navision
query_nav_items = """ SELECT [No_] AS [Nummer],[Description] AS [Beskrivelse]
                  ,[Item Category Code] AS [Varekategorikode]
				  ,CASE WHEN [Display Item] = 1 THEN 'Display'
				  WHEN [Item Category Code] = 'FÆR KAFFE' THEN 'Færdigkaffe'
				  WHEN [No_] LIKE '1040%' THEN 'Ristet kaffe'
				  WHEN [No_] LIKE '1050%' THEN 'Formalet kaffe'
				  WHEN [No_] LIKE '1020%' THEN 'Råkaffe'
				  ELSE [Item Category Code] END AS [Varetype]
                  FROM [dbo].[BKI foods a_s$Item] """
df_nav_items = pd.read_sql(query_nav_items, con_nav)

# Query for getting item numbers for production and assembly orders from Navision
query_nav_order_info = """ SELECT PAH.[No_] AS [Ordrenummer]
                       ,PAH.[Item No_] AS [Varenummer]
                       FROM [dbo].[BKI foods a_s$Posted Assembly Header] AS PAH
                       INNER JOIN [dbo].[BKI foods a_s$Item] AS I
                           ON PAH.[Item No_] = I.[No_]
                       WHERE I.[Item Category Code] = 'FÆR KAFFE'
                           AND I.[Display Item] = 1
                       UNION ALL
                       SELECT PO.[No_],PO.[Source No_]
                       FROM [dbo].[BKI foods a_s$Production Order] AS PO
                       INNER JOIN [dbo].[BKI foods a_s$Item] AS I
                           ON PO.[Source No_] = I.[No_]
                       WHERE PO.[Status] IN (2,3,4)
                           AND I.[Item Category Code] <> 'RÅKAFFE'
					   UNION ALL
					   SELECT PH.[No_], PL.[No_]
					   FROM [dbo].[BKI foods a_s$Purchase Header] AS PH
					   INNER JOIN [dbo].[BKI foods a_s$Purchase Line] AS PL
    					   ON PH.[No_] = PL.[Document No_]
						   AND PL.[Item Category Code] = 'RÅKAFFE' """
df_nav_order_info = pd.read_sql(query_nav_order_info, con_nav)


# =============================================================================
# For complete script below here
# =============================================================================

# Query section log for each section logged per script-run.
# Query is only executed at the end of each class
query_ds_section_log = f""" SELECT	SL.[Sektion] AS [Sektionskode]
                       ,S.[Beskrivelse] AS [Sektion],SS.[Beskrivelse] AS [Status]
                       ,SL.[Fejlkode_script] AS [Fejlkode script], SL.[Registreringstidspunkt]
                       FROM [trc].[Sporbarhed_sektion_log] AS SL
                       INNER JOIN [trc].[Sporbarhed_sektion] AS S
                         	ON SL.[Sektion] = S.[Id]
                       INNER JOIN [trc].[Sporbarhed_statuskode] AS SS
                            ON SL.[Statuskode] = SS.[Id]
                       WHERE SL.[Forespørgsels_id] = {req_id} """


# General info from Navision
query__nav_generelt = f""" SELECT TOP 1 PL.[Buy-from Vendor No_] AS [Leverandørnummer]
                	,V.[Name] AS [Leverandørnavn] ,PL.[No_] AS [Varenummer]
                    ,I.[Description] AS [Varenavn] ,I.[Mærkningsordning]
                    FROM [dbo].[BKI foods a_s$Purchase Line] AS PL
                    INNER JOIN [dbo].[BKI foods a_s$Item] AS I
                        ON PL.[No_] = I.[No_]
                    LEFT JOIN [dbo].[BKI foods a_s$Vendor] AS V
                        ON PL.[Buy-from Vendor No_] = V.[No_]
                    WHERE PL.[Type] = 2
                        AND PL.[Document No_] = '{req_reference_no}' """
df_nav_generelt = pd.read_sql(query__nav_generelt, con_nav)

# Get timestamp for last export of inventory from Probat
query_probat_inventory_timestamp = """ WITH [Tables] AS (
                                   SELECT MAX([RECORDING_DATE]) AS [Date]
                                    FROM [dbo].[PRO_EXP_PRODUCT_POS_INVENTORY]
                                    UNION ALL
                                    SELECT MAX([RECORDING_DATE])
                                    FROM [dbo].[PRO_EXP_WAREHOUSE_INVENTORY] )
                                    SELECT MIN([Date]) AS [Silobeholdning eksporteret]
                                    FROM [Tables] """
df_probat_inventory_timestamp = pd.read_sql(query_probat_inventory_timestamp, con_probat)

# Information from Probat for the receiving of coffee
query_probat_receiving = f""" IF '{req_modtagelse}' = 'None' -- Modtagelse ikke tastet
                         BEGIN
                         SELECT	CAST([DESTINATION] AS VARCHAR(20)) AS [Placering]
                         ,[RECORDING_DATE] AS [Dato] ,[PAPER_VALUE] / 10.0 AS [Kilo]
                		 ,NULL AS [Restlager]
                         FROM [dbo].[PRO_EXP_REC_ARRIVE]
                    	 WHERE CAST([CONTRACT_NO] AS VARCHAR(20)) = '{req_reference_no}'
                    	 UNION ALL
                    	 SELECT [Placering] ,NULL ,NULL ,SUM([Kilo]) AS [Kilo]
                    	 FROM [dbo].[Newest total inventory]
                    	 WHERE [Kontrakt] = '{req_reference_no}' 
                         AND [Placering] NOT LIKE '2__'
                    	 GROUP BY [Placering]
                         END
                         IF '{req_modtagelse}' <> 'None' -- Modtagelse er udfyldt
                         BEGIN
                         SELECT CAST([DESTINATION] AS VARCHAR(20)) AS [Placering]
                         ,[RECORDING_DATE] AS [Dato] ,[PAPER_VALUE] / 10.0 AS [Kilo]
                         ,NULL AS [Restlager]
                         FROM [dbo].[PRO_EXP_REC_ARRIVE]
                         WHERE CAST([CONTRACT_NO] AS VARCHAR(20)) = '{req_reference_no}'
                         AND CAST([DELIVERY_NAME] AS VARCHAR(20)) = '{req_modtagelse}'
                         UNION ALL
                         SELECT [Placering] ,NULL ,NULL ,SUM([Kilo]) AS [Kilo]
                         FROM [dbo].[Newest total inventory]
                         WHERE [Kontrakt] = '{req_reference_no}' 
                         AND CAST([Modtagelse] AS VARCHAR(20)) = '{req_modtagelse}'
                         AND [Placering] NOT LIKE '2__'
                         GROUP BY [Placering] END """
df_probat_receiving = pd.read_sql(query_probat_receiving, con_probat)

# Information from Probat for the processing of coffee
query_probat_processing = f""" IF 'None' = 'None' -- Ingen modtagelse tastet
                          BEGIN
                          SELECT [DESTINATION] AS [Silo]
                          ,DATEADD(D, DATEDIFF(D, 0, [START_TIME] ), 0) AS [Dato]
                          ,SUM([WEIGHT] / 10.0) AS [Kilo]
                          ,0 AS [Restlager]
                          FROM [dbo].[PRO_EXP_REC_SUM_DEST]
                          WHERE [CONTRACT_NO] = '{req_reference_no}' AND [DESTINATION] LIKE '2__'
                          GROUP BY [DESTINATION] ,DATEADD(D, DATEDIFF(D, 0, [START_TIME] ) ,0)
                          UNION ALL
                          SELECT [Placering] ,NULL ,0 ,SUM([Kilo])
                          FROM [dbo].[Newest total inventory]
                          WHERE [Kontrakt] = '{req_reference_no}' AND [Placering]  LIKE '2__'
                          GROUP BY [Placering]
                          END
                          IF 'None' <> 'None' -- Modtagelse tastet
                          BEGIN
                          SELECT [DESTINATION] AS [Silo]
                          ,DATEADD(D, DATEDIFF(D, 0, [START_TIME] ), 0) AS [Dato]
                          ,SUM([WEIGHT] / 10.0) AS [Kilo]
                          ,0 AS [Restlager]
                          FROM [dbo].[PRO_EXP_REC_SUM_DEST]
                          WHERE [CONTRACT_NO] = '{req_reference_no}'
                          AND [DESTINATION] LIKE '2__'
                          AND [DELIVERY_NAME] = '{req_modtagelse}'
                          GROUP BY [DESTINATION] ,DATEADD(D, DATEDIFF(D, 0, [START_TIME] ) ,0)
                          UNION ALL
                          SELECT [Placering] ,NULL ,0
                          ,SUM([Kilo]) AS [Kilo]
                          FROM [dbo].[Newest total inventory]
                          WHERE [Kontrakt] = '{req_reference_no}' AND [Placering]  LIKE '2__'
                          AND [Modtagelse] = '{req_modtagelse}'
                          GROUP BY [Placering]
                          END """
df_probat_processing = pd.read_sql(query_probat_processing, con_probat)

# Get order numbers the requested coffee has been used in
query_probat_used_in_roast = f""" IF 'None' = 'None' -- Ingen modtagelse tastet
                           BEGIN
                           SELECT [ORDER_NAME]
                           FROM [dbo].[PRO_EXP_ORDER_LOAD_R]
                           WHERE [S_CONTRACT_NO] = '{req_reference_no}'
                           GROUP BY [ORDER_NAME]
                           END
                           IF 'None' <> 'None' -- Modtagelse tastet
                           BEGIN
                           SELECT [ORDER_NAME]
                           FROM [dbo].[PRO_EXP_ORDER_LOAD_R]
                           WHERE [S_CONTRACT_NO] = '{req_reference_no}'
                           AND [S_DELIVERY_NAME] = '{req_modtagelse}'
                           GROUP BY [ORDER_NAME]
                           END """
df_probat_used_in_roast = pd.read_sql(query_probat_used_in_roast, con_probat)
# Convert orders to string for use in later queries
roast_orders = df_probat_used_in_roast['ORDER_NAME'].unique().tolist()
sql_roast_orders = string_to_sql(roast_orders)
# Green coffee used for roasting
query_probat_roast_input = f""" IF 'None' = 'None'
                            BEGIN
                            SELECT [CUSTOMER_CODE] AS [Varenummer]
                            ,[ORDER_NAME] AS [Ordrenummer]
                            ,[DESTINATION] AS [Rister]
                            ,SUM(CASE WHEN [S_CONTRACT_NO] = '{req_reference_no}'
                                THEN [WEIGHT] / 1000.0 ELSE 0 END) 
                                AS [Heraf kontrakt]
                            ,SUM([WEIGHT] / 1000.0) AS [Kilo råkaffe]
                            FROM [dbo].[PRO_EXP_ORDER_LOAD_R]
                            WHERE [ORDER_NAME] IN ({sql_roast_orders})
                            GROUP BY  [CUSTOMER_CODE] 
                            ,DATEADD(D, DATEDIFF(D, 0, [RECORDING_DATE] ), 0)
                            ,[ORDER_NAME] ,[DESTINATION]
                            END
                            IF 'None' <> 'None'
                            BEGIN
                            SELECT [CUSTOMER_CODE] AS [Varenummer]
                            ,[ORDER_NAME] AS [Ordrenummer]
                            ,[DESTINATION] AS [Rister]
                            ,SUM(CASE WHEN [S_CONTRACT_NO] = '{req_reference_no}' 
                                 AND [S_DELIVERY_NAME] = '{req_modtagelse}'
                          	   THEN [WEIGHT] / 1000.0
                          	   ELSE 0 END) AS [Heraf kontrakt]
                            ,SUM([WEIGHT] / 1000.0) AS [Kilo råkaffe]
                            FROM [dbo].[PRO_EXP_ORDER_LOAD_R]
                            WHERE [ORDER_NAME] IN ({sql_roast_orders})
                            GROUP BY  [CUSTOMER_CODE]
                            ,DATEADD(D, DATEDIFF(D, 0, [RECORDING_DATE] ), 0)
                            ,[ORDER_NAME] ,[DESTINATION] END """
# Only try to read query if any orders exist
if len(sql_roast_orders) > 0:
    df_probat_roast_input = pd.read_sql(query_probat_roast_input, con_probat)
else:
    df_probat_roast_input = pd.DataFrame()
# Output from roasters
query_probat_roast_output = f""" SELECT
                            DATEADD(D, DATEDIFF(D, 0, [RECORDING_DATE] ), 0) AS [Dato]
                            ,[DEST_NAME] AS [Silo] ,[ORDER_NAME] AS [Ordrenummer]
                            ,SUM([WEIGHT]) / 1000.0 AS [Kilo ristet]
                            FROM [dbo].[PRO_EXP_ORDER_UNLOAD_R]
                            WHERE [ORDER_NAME] IN ({sql_roast_orders})
                            GROUP BY DATEADD(D, DATEDIFF(D, 0, [RECORDING_DATE] ), 0)
                            ,[DEST_NAME] ,[ORDER_NAME] """
# Only try to read query if any orders exist
if len(sql_roast_orders) > 0:
    df_probat_roast_output = pd.read_sql(query_probat_roast_output, con_probat)
else:
    df_probat_roast_output = pd.DataFrame()

# Read grinding orders form Probat
query_probat_grinding_input = f""" SELECT [ORDER_NAME] AS [Ordrenummer]
                        	  ,[CUSTOMER_CODE] AS [Varenummer]
                        	  ,DATEADD(D, DATEDIFF(D, 0, [RECORDING_DATE] ), 0) AS [Dato]
                        	  ,[DESTINATION] AS [Mølle]
                        	  FROM [dbo].[PRO_EXP_ORDER_LOAD_G]
                        	  WHERE [S_ORDER_NAME] IN ({sql_roast_orders})
                        	  GROUP BY [ORDER_NAME],[CUSTOMER_CODE]
                        	  ,DATEADD(D, DATEDIFF(D, 0, [RECORDING_DATE] ), 0)
                        	  ,[DESTINATION] """
# Only try to read query if any orders exist
if len(sql_roast_orders) > 0:
    df_probat_grinding_input = pd.read_sql(query_probat_grinding_input, con_probat)
else:
    df_probat_grinding_input = pd.DataFrame()

# Convert orders to string for use in grinder output query
grinder_orders = df_probat_grinding_input['Ordrenummer'].unique().tolist()
sql_grinder_orders = string_to_sql(grinder_orders)

# Get output from grinders
query_probat_grinding_output = f""" SELECT [ORDER_NAME] AS [Ordrenummer]
                               ,SUM([WEIGHT] / 1000.0) AS [Kilo]
                               ,[DEST_NAME] AS [Silo]
                               FROM [dbo].[PRO_EXP_ORDER_UNLOAD_G]
                               WHERE [ORDER_NAME] IN ({sql_grinder_orders})
                               GROUP BY [ORDER_NAME],[DEST_NAME] """
# Only try to read query if any orders exist
if len(sql_grinder_orders) > 0:
    df_probat_grinding_output = pd.read_sql(query_probat_grinding_output, con_probat)
else:
    df_probat_grinding_output = pd.DataFrame()

# Get order relations from Probat for finished goods if possible
query_probat_orders = f""" IF 'None' = 'None' -- Modtagelse ikke defineret
                      BEGIN
                      -- Formalet kaffe
                      SELECT PG.[ORDER_NAME] AS [Ordrenummer],PG.[S_ORDER_NAME] AS [Relateret ordre],'Probat formalet pakkelinje' AS [Kilde]
                      FROM [dbo].[PRO_EXP_ORDER_LOAD_R] AS LR
                      INNER JOIN [dbo].[PRO_EXP_ORDER_LOAD_G] AS LG
                    	ON LR.[ORDER_NAME] = LG.[S_ORDER_NAME]
                      INNER JOIN [dbo].[PRO_EXP_ORDER_SEND_PG] AS PG
                    	ON LG.[ORDER_NAME] = PG.[S_ORDER_NAME]
                      WHERE LR.[S_CONTRACT_NO] = '{req_reference_no}'
                    	AND PG.[ORDER_NAME] <> ''
                      GROUP BY PG.[ORDER_NAME],PG.[S_ORDER_NAME]                      	
                      UNION ALL
                      -- Helbønne
                      SELECT PB.[ORDER_NAME] AS [Ordrenummer],PB.[S_ORDER_NAME] AS [Relateret ordre],'Probat helbønne pakkelinje'
                      FROM [dbo].[PRO_EXP_ORDER_LOAD_R] AS LR
                      INNER JOIN [dbo].[PRO_EXP_ORDER_SEND_PB] AS PB
                    	ON LR.[ORDER_NAME] = PB.[S_ORDER_NAME]
                      WHERE LR.[S_CONTRACT_NO] = '{req_reference_no}'
                    	AND PB.[ORDER_NAME] <> ''
                      GROUP BY PB.[ORDER_NAME],PB.[S_ORDER_NAME]
                      UNION ALL
					  -- Mølleordrer
                      SELECT LG.[ORDER_NAME] AS [Ordrenummer],LG.[S_ORDER_NAME] AS [Relateret ordre],'Probat mølle' AS [Kilde]
                      FROM [dbo].[PRO_EXP_ORDER_LOAD_R] AS LR
                      INNER JOIN [dbo].[PRO_EXP_ORDER_LOAD_G] AS LG
                    	ON LR.[ORDER_NAME] = LG.[S_ORDER_NAME]
                      WHERE LR.[S_CONTRACT_NO] = '{req_reference_no}'
                    	AND LG.[ORDER_NAME] <> ''
                      GROUP BY LG.[ORDER_NAME],LG.[S_ORDER_NAME]                      	
                      END
                      IF 'None' <> 'None' -- Modtagelse defineret
                      BEGIN
                      -- Formalet kaffe
                      SELECT PG.[ORDER_NAME] AS [Ordrenummer],PG.[S_ORDER_NAME] AS [Relateret ordre],'Probat formalet pakkelinje' AS [Kilde]
                      FROM [dbo].[PRO_EXP_ORDER_LOAD_R] AS LR
                      INNER JOIN [dbo].[PRO_EXP_ORDER_LOAD_G] AS LG
                    	ON LR.[ORDER_NAME] = LG.[S_ORDER_NAME]
                      INNER JOIN [dbo].[PRO_EXP_ORDER_SEND_PG] AS PG
                    	ON LG.[ORDER_NAME] = PG.[S_ORDER_NAME]
                      WHERE LR.[S_CONTRACT_NO] = '{req_reference_no}'
                    	AND LR.[S_DELIVERY_NAME] = '{req_modtagelse}'
                    	AND PG.[ORDER_NAME] <> ''
                      GROUP BY PG.[ORDER_NAME],PG.[S_ORDER_NAME]
                      UNION ALL
                      -- Helbønne
                      SELECT PB.[ORDER_NAME] AS [Ordrenummer],PB.[S_ORDER_NAME] AS [Relateret ordre],'Probat helbønne pakkelinje'
                      FROM [dbo].[PRO_EXP_ORDER_LOAD_R] AS LR
                      INNER JOIN [dbo].[PRO_EXP_ORDER_SEND_PB] AS PB
                    	ON LR.[ORDER_NAME] = PB.[S_ORDER_NAME]
                      WHERE LR.[S_CONTRACT_NO] = '{req_reference_no}'
                    	AND LR.[S_DELIVERY_NAME] = '{req_modtagelse}'
                    	AND PB.[ORDER_NAME] <> ''
                      GROUP BY PB.[ORDER_NAME],PB.[S_ORDER_NAME]
                      UNION ALL
					  -- Mølleordrer
                      SELECT LG.[ORDER_NAME] AS [Ordrenummer],LG.[S_ORDER_NAME] AS [Relateret ordre],'Probat mølle' AS [Kilde]
                      FROM [dbo].[PRO_EXP_ORDER_LOAD_R] AS LR
                      INNER JOIN [dbo].[PRO_EXP_ORDER_LOAD_G] AS LG
                    	ON LR.[ORDER_NAME] = LG.[S_ORDER_NAME]
                      WHERE LR.[S_CONTRACT_NO] = '{req_reference_no}'
                    	AND LR.[S_DELIVERY_NAME] = '{req_modtagelse}'
                    	AND LG.[ORDER_NAME] <> ''
                      GROUP BY LG.[ORDER_NAME],LG.[S_ORDER_NAME]
                      END """
df_probat_orders = pd.read_sql(query_probat_orders, con_probat)
df_probat_orders_top = df_probat_orders.loc[df_probat_orders['Kilde'] != 'Probat mølle']

# Join previous found orders to one list for query below
sql_related_orders = string_to_sql(roast_orders + grinder_orders)
# Get related orders from Navision
query_nav_order_related = f""" SELECT [Prod_ Order No_] AS [Ordrenummer]
                           ,[Reserved Prod_ Order No_] AS [Relateret ordre]
                           ,'Navision reservationer' AS [Kilde]
                           FROM [dbo].[BKI foods a_s$Reserved Prod_ Order No_]
                           WHERE [Reserved Prod_ Order No_] IN 
                           ({sql_related_orders})
                           AND [Invalid] = 0 """
df_nav_order_related = pd.read_sql(query_nav_order_related, con_nav)

# Get list of orders and append to lists if they do not already exist
# Merge Probat and NAV orders before merging
nav_orders_top = df_nav_order_related['Ordrenummer'].unique().tolist()
nav_orders_related = df_nav_order_related['Relateret ordre'].unique().tolist()
probat_orders_top = df_probat_orders_top['Ordrenummer'].unique().tolist()
probat_orders_related = df_probat_orders_top['Relateret ordre'].unique().tolist()

# Create strings dependent on request relationsship type, defined when report is requested by user
if req_ordrelationstype == 0: # All
    temp_orders_top = probat_orders_top + nav_orders_top
    temp_orders_related = probat_orders_related + nav_orders_related
elif req_ordrelationstype == 1: # Just Probat
    temp_orders_top = probat_orders_top
    temp_orders_related = probat_orders_related
elif req_ordrelationstype == 2: # Just Navision
    temp_orders_top = nav_orders_top
    temp_orders_related = nav_orders_related

# If order doesn't exist in list, append:
for order in temp_orders_top:
    if order not in  orders_top_level:
        orders_top_level.append(order)

for order in temp_orders_related:
    if order not in orders_related:
        orders_related.append(order)
# String used for querying Navision, only finished goods
req_orders_total = string_to_sql(orders_top_level)

# Recursive query to find all relevant produced orders related to the requested order
# First is identified all lotnumbers related to the orders identified through NAV reservations (only production orders)
# Next is a recursive part which identifies any document numbers which have consumed these lotnumbers (ILE_C)
# Which is then queried again to find all lotnumbers produced on the orders from which these lotnumbers originally came.
query_nav_færdigvaretilgang = f""" WITH [LOT_ORG] AS ( SELECT [Lot No_],[Document No_]
                              FROM [dbo].[BKI foods a_s$Item Ledger Entry] (NOLOCK)
                              WHERE [Order No_] IN ({req_orders_total})
                              AND [Entry Type] = 6
                              UNION ALL
                              SELECT ILE_O.[Lot No_],ILE_O.[Document No_]
                              FROM [LOT_ORG]
                              INNER JOIN [dbo].[BKI foods a_s$Item Ledger Entry] (NOLOCK) AS ILE_C
                                  ON [LOT_ORG].[Lot No_] = ILE_C.[Lot No_]
                                  AND [ILE_C].[Entry Type] IN (5,8)
                              INNER JOIN [dbo].[BKI foods a_s$Item Ledger Entry] (NOLOCK) AS ILE_O
                            	  ON ILE_C.[Document No_] = ILE_O.[Document No_]
                                  AND ILE_O.[Entry Type] IN (6,9)
                              INNER JOIN [dbo].[BKI foods a_s$Item] (NOLOCK) AS I
									  ON ILE_O.[Item No_] = I.[No_]
								  WHERE I.[Item Category Code] = 'FÆR KAFFE')
                              ,[LOT_SINGLE] AS ( SELECT [Lot No_],[Document No_] AS [Ordrenummer]
                              FROM [LOT_ORG] GROUP BY [Lot No_],[Document No_] )
                              SELECT ILE.[Item No_] AS [Varenummer],I.[Description] AS [Varenavn],[LOT_SINGLE].[Ordrenummer]
                        	  ,SUM(CASE WHEN ILE.[Entry Type] IN (0,6,9)
                        		THEN ILE.[Quantity] * I.[Net Weight]
                        		ELSE 0 END) AS [Produceret]
                        	,SUM(CASE WHEN ILE.[Entry Type] = 1
                        		THEN ILE.[Quantity] * I.[Net Weight] * -1
                        		ELSE 0 END) AS [Salg]
                        	,SUM(CASE WHEN ILE.[Entry Type] NOT IN (0,1,6,9)
                        		THEN ILE.[Quantity] * I.[Net Weight] * -1
                        		ELSE 0 END) AS [Regulering & ompak]
                        	,SUM(ILE.[Remaining Quantity] * I.[Net Weight]) AS [Restlager]
                            FROM [dbo].[BKI foods a_s$Item Ledger Entry] (NOLOCK) AS ILE
                            INNER JOIN [dbo].[BKI foods a_s$Item] (NOLOCK) AS I
                            	ON ILE.[Item No_] = I.[No_]
                            INNER JOIN [LOT_SINGLE]
                            	ON ILE.[Lot No_] = [LOT_SINGLE].[Lot No_]
                            GROUP BY ILE.[Item No_],I.[Description],[LOT_SINGLE].[Ordrenummer] """
df_nav_færdigvaretilgang = pd.read_sql(query_nav_færdigvaretilgang, con_nav)

# Recursive query to get all customer who purchased identified lotnumbers.
# See explanation of query above
query_nav_debitorer = f""" WITH [LOT_ORG] AS ( SELECT [Lot No_],[Document No_]
                      FROM [dbo].[BKI foods a_s$Item Ledger Entry] (NOLOCK)
                      WHERE [Order No_] IN({req_orders_total}) AND [Entry Type] = 6
                      UNION ALL
                      SELECT ILE_O.[Lot No_],ILE_O.[Document No_]
                      FROM [LOT_ORG]
                      INNER JOIN [dbo].[BKI foods a_s$Item Ledger Entry] (NOLOCK) AS ILE_C
                          ON [LOT_ORG].[Lot No_] = ILE_C.[Lot No_]
                    	  AND [ILE_C].[Entry Type] IN (5,8)
                      INNER JOIN [dbo].[BKI foods a_s$Item Ledger Entry] (NOLOCK) AS ILE_O
                    	  ON ILE_C.[Document No_] = ILE_O.[Document No_]
                    	  AND ILE_O.[Entry Type] IN (6,9) 
                      INNER JOIN [dbo].[BKI foods a_s$Item] (NOLOCK) AS I
							  ON ILE_O.[Item No_] = I.[No_]
						  WHERE I.[Item Category Code] = 'FÆR KAFFE')
                      ,[LOT_SINGLE] AS ( SELECT [Lot No_],[Document No_] AS [Produktionsordrenummer]
                      FROM [LOT_ORG] GROUP BY [Lot No_],[Document No_] )
                      SELECT C.[No_] AS [Debitornummer],C.[Name] AS [Debitornavn], [LOT_SINGLE].[Produktionsordrenummer]
                    	  ,ILE.[Posting Date] AS [Dato]
                    	  ,ILE.[Item No_] AS [Varenummer]
                    	  ,SUM(ILE.[Quantity] * -1) AS [Enheder]
                    	  ,SUM(ILE.[Quantity] * I.[Net Weight] * -1) AS [Kilo]
                      FROM [dbo].[BKI foods a_s$Item Ledger Entry] (NOLOCK) AS ILE
                      INNER JOIN [dbo].[BKI foods a_s$Item] (NOLOCK) AS I
                    	  ON ILE.[Item No_] = I.[No_]
                      INNER JOIN [LOT_SINGLE]
                      	  ON ILE.[Lot No_] = [LOT_SINGLE].[Lot No_]
                      INNER JOIN [dbo].[BKI foods a_s$Customer] (NOLOCK) AS C
                    	  ON ILE.[Source No_] = C.[No_]
                      WHERE ILE.[Entry Type] = 1
                      GROUP BY  C.[No_] ,C.[Name],ILE.[Posting Date],ILE.[Item No_],[LOT_SINGLE].[Produktionsordrenummer] """
df_nav_debitorer = pd.read_sql(query_nav_debitorer, con_nav)

# Query to show relation between requested order and any orders which have used it as components
query_nav_orders = f""" WITH [LOT_ORG] AS ( SELECT [Lot No_]
                              FROM [dbo].[BKI foods a_s$Item Ledger Entry] (NOLOCK)
                              WHERE [Order No_] IN ({req_orders_total})
                              AND [Entry Type] = 6
                              UNION ALL
                              SELECT ILE_O.[Lot No_]
                              FROM [LOT_ORG]
                              INNER JOIN [dbo].[BKI foods a_s$Item Ledger Entry] (NOLOCK) AS ILE_C
                                  ON [LOT_ORG].[Lot No_] = ILE_C.[Lot No_]
                                  AND [ILE_C].[Entry Type] IN (5,8)
                              INNER JOIN [dbo].[BKI foods a_s$Item Ledger Entry] (NOLOCK) AS ILE_O
                            	  ON ILE_C.[Document No_] = ILE_O.[Document No_]
                                  AND ILE_O.[Entry Type] IN (6,9)
                              INNER JOIN [dbo].[BKI foods a_s$Item] (NOLOCK) AS I
        						  ON ILE_O.[Item No_] = I.[No_]
    						  WHERE I.[Item Category Code] = 'FÆR KAFFE')
                              ,[DOC_CONS] AS ( SELECT [Lot No_], [Document No_]
                              FROM [dbo].[BKI foods a_s$Item Ledger Entry] (NOLOCK)
                              WHERE [Entry Type] IN (5,8)
                              GROUP BY [Lot No_], [Document No_] )
                              ,[DOC_OUT] AS ( SELECT [Lot No_], [Document No_]
                              FROM [dbo].[BKI foods a_s$Item Ledger Entry] (NOLOCK)
                              WHERE [Entry Type] IN (6,9)
                              GROUP BY [Lot No_], [Document No_] )
                              SELECT DO.[Document No_] AS [Relateret ordre]
                              ,DC.[Document No_] AS [Ordrenummer]
                              ,'Navision forbrug' AS [Kilde]
                              FROM [LOT_ORG] AS L
                              INNER JOIN [DOC_OUT] AS DO
                                  ON L.[Lot No_] = DO.[Lot No_]
                              LEFT JOIN [DOC_CONS] AS DC
                                  ON L.[Lot No_] = DC.[Lot No_]
                              WHERE DC.[Document No_] IS NOT NULL
                              GROUP BY DO.[Document No_] ,DC.[Document No_] """
df_nav_orders = pd.read_sql(query_nav_orders, con_nav)

# Query to get karakterer saved in BKI_Datastore
query_ds_karakterer = f""" IF 'None' = 'None'
                      BEGIN
                      SELECT SK.[Id],RRP.[Id] AS [Risteri id],SK.[Bruger] AS [Person],SK.[Dato] AS [Registreringstidspunkt]
                      ,SK.[Smag_Syre] AS [Syre],SK.[Smag_Krop] AS [Krop],SK.[Smag_Aroma] AS [Aroma],SK.[Smag_Eftersmag] AS [Eftersmag]
                      ,SK.[Smag_Robusta] AS [Robusta],ISNULL(S.[Beskrivelse],'Ej smagt') AS [Status],SK.[Bemærkning]
                      FROM [cof].[Smageskema] AS SK
                      LEFT JOIN [cof].[Risteri_modtagelse_registrering] AS RMR
                    	ON SK.[Id_org] = RMR.[Id]
                        AND RMR.[Id_org_kildenummer] = 3
                      LEFT JOIN [cof].[Risteri_råkaffe_planlægning] AS RRP
                    	ON RMR.[Id_org] = RRP.[Id]
                      LEFT JOIN [cof].[Status] AS S
                    	ON SK.[Status] = S.[Id]
                      WHERE SK.[Kontraktnummer] = '{req_reference_no}'
                      END
                      IF 'None' <> 'None'
                      BEGIN
                      SELECT SK.[Id],RRP.[Id] AS [Risteri id],SK.[Bruger] AS [Person],SK.[Dato] AS [Registreringstidspunkt]
                      ,SK.[Smag_Syre] AS [Syre],SK.[Smag_Krop] AS [Krop],SK.[Smag_Aroma] AS [Aroma],SK.[Smag_Eftersmag] AS [Eftersmag]
                      ,SK.[Smag_Robusta] AS [Robusta],ISNULL(S.[Beskrivelse],'Ej smagt') AS [Status],SK.[Bemærkning]
                      FROM [cof].[Smageskema] AS SK
                      LEFT JOIN [cof].[Risteri_modtagelse_registrering] AS RMR
                    	ON SK.[Id_org] = RMR.[Id]
                        AND RMR.[Id_org_kildenummer] = 3
                      LEFT JOIN [cof].[Risteri_råkaffe_planlægning] AS RRP
                    	ON RMR.[Id_org] = RRP.[Id]
                      LEFT JOIN [cof].[Status] AS S
                    	ON SK.[Status] = S.[Id]
                      WHERE SK.[Kontraktnummer] = '{req_reference_no}'
                    	AND RRP.[Delivery] = '{req_modtagelse}'
                      END """
df_ds_karakterer = pd.read_sql(query_ds_karakterer, con_04)

query_probat_gc_samples = f""" IF 'None' = 'None'
                        BEGIN
                        SELECT [RECORDING_DATE] AS [Dato],[SAMPLE_ID] AS [Probat id],[VOLUME] AS [Volumen]
                        ,[HUMIDITY_1] AS [Vandprocent 1],[HUMIDITY_2] AS [Vandprocent 2],[HUMIDITY_3] AS [Vandprocent 3]
                        ,[USERNAME] AS [Bruger],[INFO] AS [Bemærkning]
                        FROM [dbo].[PRO_EXP_SAMPLE_RECEIVING]
                        WHERE [PRO_EXPORT_GENERAL_ID] IN (SELECT MAX([PRO_EXPORT_GENERAL_ID]) FROM [dbo].[PRO_EXP_SAMPLE_RECEIVING] GROUP BY [SAMPLE_ID])
                        	AND [CONTRACT_NO] = '{req_reference_no}'
                        END
                        IF 'None' <> 'None'
                        BEGIN
                        SELECT [RECORDING_DATE] AS [Dato],[SAMPLE_ID] AS [Probat id],[VOLUME] AS [Volumen]
                        ,[HUMIDITY_1] AS [Vandprocent 1],[HUMIDITY_2] AS [Vandprocent 2],[HUMIDITY_3] AS [Vandprocent 3]
                        ,[USERNAME] AS [Bruger],[INFO] AS [Bemærkning]
                        FROM [dbo].[PRO_EXP_SAMPLE_RECEIVING]
                        WHERE [PRO_EXPORT_GENERAL_ID] IN (SELECT MAX([PRO_EXPORT_GENERAL_ID]) FROM [dbo].[PRO_EXP_SAMPLE_RECEIVING] GROUP BY [SAMPLE_ID])
                        	AND [CONTRACT_NO] = '{req_reference_no}'
                            AND [DELIVERY_NAME] = '{req_modtagelse}'
                        END """
df_probat_gc_samples = pd.read_sql(query_probat_gc_samples, con_probat)


# =============================================================================
# Section 1: Generelt
# =============================================================================
section_id = 1
section_name = get_section_name(section_id)
column_order = ['Kontraktnummer','Modtagelse','Varenummer','Varenavn','Mærkningsordning','Leverandørnummer'
                ,'Leverandørnavn','Silobeholdning eksporteret']

if get_section_status_code(df_nav_generelt) == 99:
    try:
        df_nav_generelt['Kontraktnummer'] = req_reference_no
        df_nav_generelt['Modtagelse'] = req_modtagelse
        df_nav_generelt['Silobeholdning eksporteret'] = df_probat_inventory_timestamp['Silobeholdning eksporteret'].iloc[0]
        # Apply column formating
        df_nav_generelt['Silobeholdning eksporteret'] = df_nav_generelt['Silobeholdning eksporteret'].dt.strftime('%d-%m-%Y %H:%M')
        # Transpose dataframe
        df_nav_generelt = df_nav_generelt[column_order].transpose()
        df_nav_generelt = df_nav_generelt.reset_index()
        df_nav_generelt.columns = ['Sektion','Værdi']
        # Write results to Word and Excel
        insert_dataframe_into_excel(df_nav_generelt, section_name, True)
        add_section_to_word(df_nav_generelt, section_name, True, [0])
        # Write status into log
        section_log_insert(section_id, 0)
    except Exception as e: # Insert error into log
        section_log_insert(section_id, 2, e)
else: # Write into log if no data is found or section is out of scope
    section_log_insert(section_id, get_section_status_code(df_nav_generelt))

# =============================================================================
# Section 21: Modtagelse
# =============================================================================
section_id = 21
section_name = get_section_name(section_id)
column_order = ['Placering','Dato','Kilo','Restlager']
columns_1_dec = ['Kilo','Restlager']

if get_section_status_code(df_probat_receiving) == 99:
    try:
        # Create total for dataframe
        dict_modtagelse_total = {'Kilo': [df_probat_receiving['Kilo'].sum()],
                                 'Restlager': [df_probat_receiving['Restlager'].sum()]}
        # Create temp dataframe including total
        df_temp_total = pd.concat([df_probat_receiving,
                                   pd.DataFrame.from_dict(data=dict_modtagelse_total, orient = 'columns')])
        # Apply column formating
        df_temp_total['Dato'] = df_temp_total['Dato'].dt.strftime('%d-%m-%Y')
        for col in columns_1_dec:
            df_temp_total[col] = df_temp_total[col].apply(lambda x: number_format(x, 'dec_1'))
        df_temp_total = df_temp_total[column_order]
        # Write results to Word and Excel
        insert_dataframe_into_excel(df_temp_total, section_name, False)
        add_section_to_word(df_temp_total, section_name, True, [-1,0])
        # Write status into log
        section_log_insert(section_id, 0)
    except Exception as e: # Insert error into log
        section_log_insert(section_id, 2, e)
else: # Write into log if no data is found or section is out of scope
    section_log_insert(section_id, get_section_status_code(df_probat_receiving))

# =============================================================================
# Section 20: Rensning
# =============================================================================
section_id = 20
section_name = get_section_name(section_id)
column_order = ['Silo','Dato','Kilo','Restlager']
columns_1_dec = ['Kilo','Restlager']
if get_section_status_code(df_probat_processing) == 99:
    try:
        # Apply column formating for date column before concat
        df_probat_processing['Dato'] = df_probat_processing['Dato'].dt.strftime('%d-%m-%Y')
        df_probat_processing.fillna('', inplace=True)
        #Concat dates into one strng and sum numeric columns if they can be grouped
        df_probat_processing = df_probat_processing.groupby('Silo').agg(
            {'Kilo': 'sum',
             'Restlager': 'sum',
             'Dato': lambda x: ','.join(sorted(pd.Series.unique(x)))
             }).reset_index()
        df_probat_processing['Dato'] = df_probat_processing['Dato'].apply(lambda x: x.rstrip(','))
        df_probat_processing['Dato'] = df_probat_processing['Dato'].apply(lambda x: x.lstrip(','))
        # Create total for dataframe
        dict_modtagelse_total = {'Kilo': [df_probat_processing['Kilo'].sum()],
                                 'Restlager': [df_probat_processing['Restlager'].sum()]}
        # Create temp dataframe including total
        df_temp_total = pd.concat([df_probat_processing,
                                   pd.DataFrame.from_dict(data=dict_modtagelse_total, orient = 'columns')])
        for col in columns_1_dec:
            df_temp_total[col] = df_temp_total[col].apply(lambda x: number_format(x, 'dec_1'))
        df_temp_total = df_temp_total[column_order]
        # Write results to Word and Excel
        insert_dataframe_into_excel(df_temp_total, section_name, False)
        add_section_to_word(df_temp_total, section_name, True, [-1,0])
        # Write status into log
        section_log_insert(section_id, 0)
    except Exception as e: # Insert error into log
        section_log_insert(section_id, 2, e)
else: # Write into log if no data is found or section is out of scope
    section_log_insert(section_id, get_section_status_code(df_probat_processing))

# =============================================================================
# Section 5: Risteordrer
# =============================================================================
section_id = 5
section_name = get_section_name(section_id)
column_order = ['Varenummer','Varenavn','Dato','Rister','Ordrenummer','Silo',
                'Kilo råkaffe','Heraf kontrakt','Kilo ristet']
columns_1_dec = ['Kilo råkaffe','Heraf kontrakt','Kilo ristet']

if get_section_status_code(df_probat_roast_output) == 99:
    try:
        # Apply column formating for date column before concat
        df_probat_roast_output['Dato'] = df_probat_roast_output['Dato'].dt.strftime('%d-%m-%Y')
        # Concat dates into one strng and sum numeric columns if they can be grouped
        df_probat_roast_output = df_probat_roast_output.groupby('Ordrenummer').agg(
            {'Kilo ristet': 'sum',
             'Dato': lambda x: ','.join(sorted(pd.Series.unique(x))),
             'Silo': lambda x: ','.join(sorted(pd.Series.unique(x)))
             }).reset_index()
        # Join roast output to input for one table
        df_probat_roast_total = pd.merge(df_probat_roast_input,
                                         df_probat_roast_output,
                                         left_on = 'Ordrenummer',
                                         right_on = 'Ordrenummer',
                                         how = 'left',
                                         suffixes = ('' ,'_R')
                                         )
        #Column formating and lookups
        df_probat_roast_total['Dato'] = df_probat_roast_total['Dato'].apply(lambda x: x.rstrip(','))
        df_probat_roast_total['Dato'] = df_probat_roast_total['Dato'].apply(lambda x: x.lstrip(','))
        df_probat_roast_total['Varenavn'] = df_probat_roast_total['Varenummer'].apply(get_nav_item_info, field='Beskrivelse')
        # Create total for dataframe
        dict_risteordrer_total = {'Kilo råkaffe': df_probat_roast_total['Kilo råkaffe'].sum(),
                                 'Heraf kontrakt': df_probat_roast_total['Heraf kontrakt'].sum(),
                                 'Kilo ristet': df_probat_roast_total['Kilo ristet'].sum()
                                 }

        # Create temp dataframe including total
        df_temp_total = pd.concat([df_probat_roast_total,
                               pd.DataFrame([dict_risteordrer_total])])
        for col in columns_1_dec:
            df_temp_total[col] = df_temp_total[col].apply(lambda x: number_format(x, 'dec_1'))
        df_temp_total = df_temp_total[column_order]
        df_temp_total.sort_values(by=['Varenummer'] ,inplace=True)
        # Write results to Word and Excel
        insert_dataframe_into_excel(df_temp_total, section_name, False)
        add_section_to_word(df_temp_total, section_name, True, [-1,0])
        # Write status into log
        section_log_insert(section_id, 0)
    except Exception as e: # Insert error into log
        section_log_insert(section_id, 2, e)
else: # Write into log if no data is found or section is out of scope
    section_log_insert(section_id, get_section_status_code(df_probat_roast_output))


# =============================================================================
# Section 4: Mølleordrer
# =============================================================================
section_id = 4
section_name = get_section_name(section_id)
column_order = ['Varenummer','Varenavn','Ordrenummer','Dato','Silo','Kilo']
columns_1_dec = ['Kilo']

if get_section_status_code(df_probat_grinding_input) == 99:
    try:
        # Apply column formating for date column before concat
        df_probat_grinding_input['Dato'] = df_probat_grinding_input['Dato'].dt.strftime('%d-%m-%Y')
        # Concat dates into one string and sum numeric columns if they can be grouped
        df_probat_grinding_input = df_probat_grinding_input.groupby(['Ordrenummer','Varenummer','Mølle']).agg(
            {'Dato': lambda x: ','.join(sorted(pd.Series.unique(x)))
            }).reset_index()
        df_probat_grinding_output = df_probat_grinding_output.groupby('Ordrenummer').agg(
            {'Kilo': 'sum',
             'Silo': lambda x: ','.join(sorted(pd.Series.unique(x)))
            }).reset_index()
        # Join roast output to input for one table
        df_probat_grinding_total = pd.merge(df_probat_grinding_input,
                                         df_probat_grinding_output,
                                         left_on = 'Ordrenummer',
                                         right_on = 'Ordrenummer',
                                         how = 'left',
                                         suffixes = ('' ,'_R')
                                         )
        #Column formating and lookups
        df_probat_grinding_total['Dato'] = df_probat_grinding_total['Dato'].apply(lambda x: x.rstrip(','))
        df_probat_grinding_total['Dato'] = df_probat_grinding_total['Dato'].apply(lambda x: x.lstrip(','))
        df_probat_grinding_total['Silo'] = df_probat_grinding_total['Silo'].apply(lambda x: x.rstrip(','))
        df_probat_grinding_total['Silo'] = df_probat_grinding_total['Silo'].apply(lambda x: x.lstrip(','))
        df_probat_grinding_total['Varenavn'] = df_probat_grinding_total['Varenummer'].apply(get_nav_item_info, field='Beskrivelse')
        # Create total for dataframe
        dict_mølleordrer_total = {'Kilo': df_probat_grinding_total['Kilo'].sum()}

        # Create temp dataframe including total
        df_temp_total = pd.concat([df_probat_grinding_total,
                               pd.DataFrame([dict_mølleordrer_total])])
        for col in columns_1_dec:
            df_temp_total[col] = df_temp_total[col].apply(lambda x: number_format(x, 'dec_1'))
        df_temp_total = df_temp_total[column_order]
        df_temp_total.sort_values(by=['Varenummer'] ,inplace=True)
        # Write results to Word and Excel
        insert_dataframe_into_excel(df_temp_total, section_name, False)
        add_section_to_word(df_temp_total, section_name, True, [-1,0])
        # Write status into log
        section_log_insert(section_id, 0)
    except Exception as e: # Insert error into log
        section_log_insert(section_id, 2, e)
else: # Write into log if no data is found or section is out of scope
    section_log_insert(section_id, get_section_status_code(df_probat_roast_output))

# =============================================================================
# Section 3: Færdigvaretilgang
# =============================================================================
section_id = 3
section_name = get_section_name(section_id)
column_order = ['Varenummer','Varenavn','Ordrenummer','Produceret','Salg','Restlager','Regulering & ompak']
columns_1_dec = ['Produceret','Salg','Restlager','Regulering & ompak']

if get_section_status_code(df_nav_færdigvaretilgang) == 99:
    try:
        # Concat order numbers to one string
        df_nav_færdigvaretilgang = df_nav_færdigvaretilgang.groupby(['Varenummer','Varenavn']).agg(
            {'Ordrenummer': lambda x: ','.join(sorted(pd.Series.unique(x))),
             'Produceret': 'sum',
             'Salg': 'sum',
             'Restlager': 'sum',
             'Regulering & ompak': 'sum'
            }).reset_index()
        df_nav_færdigvaretilgang['Ordrenummer'] = df_nav_færdigvaretilgang['Ordrenummer'].apply(lambda x: x.rstrip(','))
        df_nav_færdigvaretilgang['Ordrenummer'] = df_nav_færdigvaretilgang['Ordrenummer'].apply(lambda x: x.lstrip(','))
        # Create total for dataframe
        dict_færdigvare_total = {'Produceret': [df_nav_færdigvaretilgang['Produceret'].sum()],
                                 'Salg': [df_nav_færdigvaretilgang['Salg'].sum()],
                                 'Restlager': [df_nav_færdigvaretilgang['Restlager'].sum()],
                                 'Regulering & ompak': [df_nav_færdigvaretilgang['Regulering & ompak'].sum()]}
        df_temp_total = pd.concat([df_nav_færdigvaretilgang,
                                  pd.DataFrame.from_dict(data=dict_færdigvare_total, orient='columns')])
        df_temp_total = df_temp_total[column_order]
        df_temp_total.sort_values(by=['Varenummer'], inplace=True)
        # Data formating
        for col in columns_1_dec:
            df_temp_total[col] = df_temp_total[col].apply(lambda x: number_format(x, 'dec_1'))
        # Write results to Word and Excel
        insert_dataframe_into_excel (df_temp_total, section_name, False)
        add_section_to_word(df_temp_total, section_name, True, [-1,0])
        # Write status into log
        section_log_insert(section_id, 0)
    except Exception as e: # Insert error into log
        section_log_insert(section_id, 2, e)
else: # Write into log if no data is found or section is out of scope
    section_log_insert(section_id, get_section_status_code(df_temp_total))

# =============================================================================
# Section 7: Debitorer
# =============================================================================
section_id = 7
section_name = get_section_name(section_id)
column_order = ['Debitornummer','Debitornavn','Dato','Varenummer','Varenavn','Produktionsordrenummer',
                'Enheder','Kilo']
columns_1_dec = ['Enheder','Kilo']

if get_section_status_code(df_nav_debitorer) == 99:
    try:
        # Concat Order nos to one string
        df_nav_debitorer = df_nav_debitorer.groupby(['Debitornummer','Debitornavn','Dato','Varenummer']).agg(
            {'Produktionsordrenummer': lambda x: ','.join(sorted(pd.Series.unique(x))),
             'Enheder': 'sum',
             'Kilo': 'sum'
            }).reset_index()
        df_nav_debitorer['Produktionsordrenummer'] = df_nav_debitorer['Produktionsordrenummer'].apply(lambda x: x.rstrip(','))
        df_nav_debitorer['Produktionsordrenummer'] = df_nav_debitorer['Produktionsordrenummer'].apply(lambda x: x.lstrip(','))
        # Create total for dataframe
        dict_debitor_total = {'Enheder': [df_nav_debitorer['Enheder'].sum()],
                              'Kilo':[df_nav_debitorer['Kilo'].sum()]}
        # Add varenavn
        df_nav_debitorer['Varenavn'] = df_nav_debitorer['Varenummer'].apply(get_nav_item_info, field='Beskrivelse')
        # Look up column values and string format datecolumn for export
        df_nav_debitorer['Dato'] = df_nav_debitorer['Dato'].dt.strftime('%d-%m-%Y')
        # Create temp dataframe with total
        df_temp_total = pd.concat([df_nav_debitorer, pd.DataFrame.from_dict(data=dict_debitor_total, orient='columns')])
        df_temp_total = df_temp_total[column_order]
        df_temp_total.sort_values(by=['Varenummer','Debitornummer','Dato'], inplace=True)
        # Data formating
        for col in columns_1_dec:
            df_temp_total[col] = df_temp_total[col].apply(lambda x: number_format(x, 'dec_1'))
        # Write results to Word and Excel
        insert_dataframe_into_excel (df_temp_total, section_name, False)
        add_section_to_word(df_temp_total, section_name, True, [-1,0])
        # Write status into log
        section_log_insert(section_id, 0)
    except Exception as e: # Insert error into log
        section_log_insert(section_id, 2, e)
else: # Write into log if no data is found or section is out of scope
    section_log_insert(section_id, get_section_status_code(df_temp_total))

# =============================================================================
# Section 8: Massebalance
# =============================================================================
section_id = 8
section_name = get_section_name(section_id)
columns_1_dec = ['[1] Kontrakt','[2] Renset','[3] Restlager','[4] Difference','[6] Anvendt til produktion',
                 '[7] Difference','[9] Ristet kaffe','[10] Difference','[12] Færdigvareproduktion',
                 '[13] Difference','[15] Salg','[16] Regulering & ompak','[17] Restlager','[18] Difference']
columns_2_pct = ['[5] Difference pct','[8] Difference pct','[11] Difference pct','[14] Difference pct','[19] Difference pct']

dict_massebalance = {'[1] Kontrakt': df_probat_receiving['Kilo'].sum(),
                     '[2] Renset': df_probat_processing['Kilo'].sum(),
                     '[3] Restlager': df_probat_processing['Restlager'].sum(),
                     '[4] Difference': None,
                     '[5] Difference pct': None,
                     '[6] Anvendt til produktion': df_probat_roast_total['Heraf kontrakt'].sum(),
                     '[7] Difference': None,
                     '[8] Difference pct': None,
                     '[9] Ristet kaffe': df_probat_roast_total['Kilo ristet'].sum(),
                     '[10] Difference': None,
                     '[11] Difference pct': None,
                     '[12] Færdigvareproduktion': df_nav_færdigvaretilgang['Produceret'].sum(),
                     '[13] Difference': None,
                     '[14] Difference pct': None,
                     '[15] Salg': df_nav_færdigvaretilgang['Salg'].sum(),
                     '[16] Regulering & ompak': df_nav_færdigvaretilgang['Regulering & ompak'].sum(),
                     '[17] Restlager': df_nav_færdigvaretilgang['Restlager'].sum(),
                     '[18] Difference': None,
                     '[19] Difference pct': None
                    }
# Calculate differences and percentages before converting to dataframe:
dict_massebalance['[4] Difference'] = dict_massebalance['[1] Kontrakt'] - dict_massebalance['[2] Renset'] - dict_massebalance['[3] Restlager']
dict_massebalance['[5] Difference pct'] = zero_division(dict_massebalance['[4] Difference'], dict_massebalance['[1] Kontrakt'], 'None')
dict_massebalance['[7] Difference'] = ( dict_massebalance['[2] Renset'] - dict_massebalance['[3] Restlager']
                                        - dict_massebalance['[6] Anvendt til produktion'] )
dict_massebalance['[8] Difference pct'] = zero_division(dict_massebalance['[7] Difference'],
                                                        dict_massebalance['[2] Renset'] - dict_massebalance['[3] Restlager'], 'None')
dict_massebalance['[10] Difference'] = dict_massebalance['[2] Renset'] - dict_massebalance['[3] Restlager'] - dict_massebalance['[9] Ristet kaffe']
dict_massebalance['[11] Difference pct'] = zero_division(dict_massebalance['[10] Difference'],
                                                        dict_massebalance['[2] Renset'] - dict_massebalance['[3] Restlager'], 'None')
dict_massebalance['[13] Difference'] = dict_massebalance['[9] Ristet kaffe'] - dict_massebalance['[12] Færdigvareproduktion']
dict_massebalance['[14] Difference pct'] = zero_division(dict_massebalance['[13] Difference'], dict_massebalance['[12] Færdigvareproduktion'], 'None')
dict_massebalance['[18] Difference'] = ( dict_massebalance['[12] Færdigvareproduktion'] - dict_massebalance['[15] Salg']
                                         - dict_massebalance['[16] Regulering & ompak'] - dict_massebalance['[17] Restlager'] )
dict_massebalance['[19] Difference pct'] = zero_division(dict_massebalance['[18] Difference'],
                                                         dict_massebalance['[12] Færdigvareproduktion'], 'None')
#Number formating
for col in columns_1_dec:
    dict_massebalance[col] = number_format(dict_massebalance[col] ,'dec_1')
for col in columns_2_pct:
    dict_massebalance[col] = number_format(dict_massebalance[col] ,'pct_2')

df_massebalance = pd.DataFrame.from_dict(data=dict_massebalance, orient='index').reset_index()
df_massebalance.columns = ['Sektion','Værdi']
df_massebalance['Note'] = [None,None,None,'[1]-[2]-[3]','[4]/[1]',None,'[2]-[3]-[6]','[7]/([2]-[3])',None,
                           '[2]-[3]-[9]','[10]/([2]-[3])',None,'[9]-[12]','[13]/[12]',None,None,None,
                           '[12]-[15]-[16]-[17]','[18]/[12]']
df_massebalance['Bemærkning'] = None

if get_section_status_code(df_massebalance) == 99:
    try:
        # Write results to Word and Excel
        insert_dataframe_into_excel (df_massebalance, section_name, True)
        add_section_to_word(df_massebalance, section_name, True, [0,4,5,7,8,10,11,13,14,18,19])
        # Write status into log
        section_log_insert(section_id, 0)
    except Exception as e: # Insert error into log
        section_log_insert(section_id, 2, e)
else: # Write into log if no data is found or section is out of scope
    section_log_insert(section_id, get_section_status_code(df_massebalance))

# =============================================================================
# Section 2: Relaterede ordrer Kontrakt --> færdigvare
# =============================================================================
section_id = 2
section_name = get_section_name(section_id)
column_order = ['Ordrenummer','Varenummer','Navn','Relateret ordre',
                'Relateret vare','Relateret navn','Kilde']
df_temp_orders = pd.concat([df_nav_orders,df_probat_orders,df_nav_order_related])

if get_section_status_code(df_temp_orders) == 99:
    try:
        df_temp_orders['Varenummer'] = df_temp_orders['Ordrenummer'].apply(lambda x: get_nav_order_info(x))
        df_temp_orders['Navn'] = df_temp_orders['Varenummer'].apply(lambda x: get_nav_item_info(x, 'Beskrivelse'))
        df_temp_orders['Relateret vare'] = df_temp_orders['Relateret ordre'].apply(lambda x: get_nav_order_info(x))
        df_temp_orders['Relateret navn'] = df_temp_orders['Relateret vare'].apply(lambda x: get_nav_item_info(x, 'Beskrivelse'))
        # Remove orders not existing in NAV and sort columns and rows
        df_temp_orders.dropna(inplace=True)
        df_temp_orders = df_temp_orders[column_order]
        df_temp_orders.sort_values(by=['Ordrenummer','Relateret ordre'], inplace=True)
        # Write results to Word and Excel
        insert_dataframe_into_excel (df_temp_orders, section_name, False)
        add_section_to_word(df_temp_orders, section_name, True, [0])
        # Write status into log
        section_log_insert(section_id, 0)
# =================================================================
# Section 19: Relation visualization
# =================================================================
        #Try to create .png with relations illustrated and add to .docx as well
        try:
            df_temp_order_relation = df_temp_orders[['Ordrenummer','Varenummer','Relateret ordre','Relateret vare']]
            df_temp_order_relation['Ordretype'] = df_temp_order_relation['Varenummer'].apply(lambda x: get_nav_item_info(x, 'Varetype'))
            df_temp_order_relation['Relateret ordretype'] = df_temp_order_relation['Relateret vare'].apply(lambda x: get_nav_item_info(x, 'Varetype'))
            df_temp_order_relation['Primær'] = df_temp_order_relation['Ordretype'] + '\n' + df_temp_order_relation['Ordrenummer']
            df_temp_order_relation['Sekundær'] = df_temp_order_relation['Relateret ordretype'] + '\n' + df_temp_order_relation['Relateret ordre']
            df_temp_order_relation = df_temp_order_relation[['Primær','Sekundær']]
            # Add green coffees
            df_temp_gc_orders = pd.DataFrame(columns=['Primær','Sekundær'])
            df_temp_gc_orders['Primær'] = 'Ristet kaffe' + '\n' + df_probat_roast_input['Ordrenummer']
            df_temp_gc_orders['Sekundær'] = 'Råkaffe' + '\n' + req_reference_no
            df_order_relations = pd.concat([df_temp_order_relation,df_temp_gc_orders])
            # Create relation visualization
            array_for_drawing = list(df_order_relations.itertuples(index=False, name=None))
            graph = nx.DiGraph()
            graph.add_edges_from(array_for_drawing)
            relations_plot = nx.drawing.nx_pydot.to_pydot(graph)
            relations_plot.write_png(path_png_relations)
            # Add image to word document
            doc.add_picture(path_png_relations, width=Inches(11.0), height=Inches(6.50))
            # Write to log
            section_log_insert(19, 0)
        except Exception as e: # Insert error into log. Same section_id as others..
            section_log_insert(19, 2, e)
    except Exception as e: # Insert error into log
        section_log_insert(section_id, 2, e)
else: # Write into log if no data is found or section is out of scope
    section_log_insert(section_id, get_section_status_code(df_temp_orders))

# =============================================================================
# Section 12: Karakterer
# =============================================================================
section_id = 12
section_name = get_section_name(section_id)
column_order = ['Id','Risteri id','Person','Registreringstidspunkt','Syre','Krop','Aroma','Eftersmag','Robusta','Status','Bemærkning']

if get_section_status_code(df_ds_karakterer) == 99:
    try:
        # String format datecolumn for export
        df_ds_karakterer['Registreringstidspunkt'] = df_ds_karakterer['Registreringstidspunkt'].dt.strftime('%d-%m-%Y')
        df_ds_karakterer.sort_values(by=['Id'], inplace=True)
        # Write results to Word and Excel
        insert_dataframe_into_excel (df_ds_karakterer, section_name, False)
        add_section_to_word(df_ds_karakterer, section_name, False, [0])
        # Write status into log
        section_log_insert(section_id, 0)
    except Exception as e: # Insert error into log
        section_log_insert(section_id, 2, e)
else: # Write into log if no data is found or section is out of scope
    section_log_insert(section_id, get_section_status_code(df_ds_karakterer))


# =============================================================================
# Section 22: Probat samples
# =============================================================================
section_id = 22
section_name = get_section_name(section_id)
column_order = ['Dato','Probat id','Volumen,Vandprocent 1','Vandprocent 2','Vandprocent 3'
                ,'Bruger','Bemærkning']

if get_section_status_code(df_probat_gc_samples) == 99:
    try:
        # String format datecolumn for export
        df_probat_gc_samples['Dato'] = df_probat_gc_samples['Dato'].dt.strftime('%d-%m-%Y')
        df_probat_gc_samples.sort_values(by=['Probat id'], inplace=True)
        # Write results to Word and Excel
        insert_dataframe_into_excel (df_probat_gc_samples, section_name, False)
        add_section_to_word(df_probat_gc_samples, section_name, False, [0])
        # Write status into log
        section_log_insert(section_id, 0)
    except Exception as e: # Insert error into log
        section_log_insert(section_id, 2, e)
else: # Write into log if no data is found or section is out of scope
    section_log_insert(section_id, get_section_status_code(df_probat_gc_samples))


# =============================================================================
# Section 18: Sektionslog
# =============================================================================
section_id = 18
df_section_log = pd.read_sql(query_ds_section_log, con_04)
section_name = get_section_name(section_id)

if get_section_status_code(df_section_log) == 99:
    try:
        df_section_log['Registreringstidspunkt'] = df_section_log['Registreringstidspunkt'].dt.strftime('%H:%M%:%S')
        df_section_log.sort_values(by=['Sektionskode'], inplace=True)
        # Write results to Word and Excel
        insert_dataframe_into_excel (df_section_log, section_name, False)
        add_section_to_word(df_section_log, section_name, False, [0])
        # Write status into log
        section_log_insert(section_id, 0)
    except Exception as e: # Insert error into log
        section_log_insert(section_id, 2, e)
else: # Write into log if no data is found or section is out of scope
    section_log_insert(section_id, get_section_status_code(df_section_log))

#Save files
excel_writer.save()
# log_insert(script_name, f'Excel file {file_name} created')

doc.save(path_file_doc)
# log_insert(script_name, f'Word document {file_name} created')









# =============================================================================
# Write into email log
# =============================================================================
dict_email_log = {'Filsti': filepath
                  ,'Filnavn': file_name
                  ,'Modtager': req_recipients
                  ,'Emne': f'Anmodet rapport for ordre {req_reference_no}'
                  ,'Forespørgsels_id': req_id
                  ,'Note':req_note}
# pd.DataFrame(data=dict_email_log, index=[0]).to_sql('Sporbarhed_email_log', con=engine_04, schema='trc', if_exists='append', index=False)
# log_insert(script_name, f'Request id: {req_id} inserted into [trc].[Email_log]')

# =============================================================================
# Update request that dataprocessing has been completed
# =============================================================================
cursor_04.execute(f"""UPDATE [trc].[Sporbarhed_forespørgsel]
                  SET Data_færdigbehandlet = 1
                  WHERE [Id] = {req_id}""")
cursor_04.commit()
# log_insert(script_name, f'Request id: {req_id} completed')

# Exit script
raise SystemExit(0)
