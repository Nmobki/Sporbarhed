#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import urllib
from datetime import datetime
import pandas as pd
from sqlalchemy import create_engine
import pyodbc
import docx
from docx.shared import Inches
import openpyxl
import networkx as nx
import Sporbarhed_råkaffe


# =============================================================================
# Define functions
# =============================================================================

# Get section name for section from query
def get_section_name(section):
    df_temp_sections = df_sections.loc[df_sections['Sektion'] == section]
    x = df_temp_sections['Sektion navn'].iloc[0]
    if len(x) == 0 or len(x) > 31:
        return 'Sektion ' + str(section)
    else:
        return x

# Find statuscode for section log
def get_section_status_code(dataframe):
    if len(dataframe) == 0:
        return 1 # Empty dataframe
    else:
        return 99 # Continue

# Write into section log
def section_log_insert(section, statuscode, errorcode=None):
    df = pd.DataFrame(data={'Forespørgsels_id':req_id,
                            'Sektion':section,
                            'Statuskode':statuscode,
                            'Fejlkode_script':str(errorcode)}
                      , index=[0])
    df.to_sql('Sporbarhed_sektion_log', con=engine_04, schema='trc', if_exists='append', index=False)

# Write dataframe into Excel sheet
def insert_dataframe_into_excel (dataframe, sheetname, include_index):
    dataframe.to_excel(excel_writer, sheet_name=sheetname, index=include_index)

# Convert list into string for SQL IN operator
def string_to_sql(list_with_values):
    if len(list_with_values) == 0:
        return ''
    else:
        return "'{}'".format("','".join(list_with_values))

def number_format(value, number_type):
    try:
        if number_type == 'dec_2':
            return f'{round(value,2):,}'.replace(',', ';').replace('.', ',').replace(';', '.')
        elif number_type == 'dec_1':
            return f'{round(value,1):,}'.replace(',', ';').replace('.', ',').replace(';', '.')
        elif number_type == 'dec_0':
            return f'{int(round(value,0)):,}'.replace(',', ';').replace('.', ',').replace(';', '.')
        elif number_type == 'pct_2':
            return f'{round(value,4):.2%}'.replace(',', ';').replace('.', ',').replace(';', '.')
        elif number_type == 'pct_0':
            return f'{round(value,2):.0%}'.replace(',', ';').replace('.', ',').replace(';', '.')
        else:
            return value
    except:
        return value

# Prevent division by zero error
def zero_division(nominator, denominator, zero_return):
    dict = {'None':None,'Zero':0}
    if denominator in [0,None]:
        return dict[zero_return]
    else:
        return nominator / denominator

# Convert placeholder values from dataframe to empty string for Word document
def convert_placeholders_word(string):
    if string in ['None','nan','NaT']:
        return ''
    else:
        return string

# Strip comma from commaseparated strings
def strip_comma_from_string(text):
    text = text.rstrip(',')
    return text.lstrip(',')

# Write into dbo.log
def log_insert(event, note):
    dict_log = {'Note': note
                ,'Event': event}
    pd.DataFrame(data=dict_log, index=[0]).to_sql('Log', con=engine_04, schema='dbo', if_exists='append', index=False)

# Get info from item table in Navision
def get_nav_item_info(item_no, field):
    if item_no in df_nav_items['Nummer'].tolist():
        df_temp = df_nav_items[df_nav_items['Nummer'] == item_no]
        return df_temp[field].iloc[0]
    else:
        return None

# Get info from assembly and production orders in Navision
def get_nav_order_info(order_no):
    if order_no in df_nav_order_info['Ordrenummer'].tolist(): 
        df_temp = df_nav_order_info[df_nav_order_info['Ordrenummer'] == order_no]
        return df_temp['Varenummer'].iloc[0]
    else:
        return None

# Add dataframe to word document
def add_section_to_word(dataframe, section, pagebreak, rows_to_bold):
    # Add section header
    doc.add_heading(section, 1)
    # Add a table with an extra row for headers
    table = doc.add_table(dataframe.shape[0]+1, dataframe.shape[1])
    table.style = 'Table Grid'
    # Add headers to top row
    for i in range(dataframe.shape[-1]):
        table.cell(0,i).text = dataframe.columns[i]
    # Add data from dataframe to the table, replace supposed blank cells using function
    for x in range(dataframe.shape[0]):
        for y in range(dataframe.shape[-1]):
            table.cell(x+1,y).text =  convert_placeholders_word(str(dataframe.values[x,y]))
    # Bold total row if it exists
    for y in rows_to_bold:
        for x in range(dataframe.shape[1]):
            table.rows[y].cells[x].paragraphs[0].runs[0].font.bold = True
    # Add page break
    if pagebreak:
        doc.add_page_break()


# =============================================================================
# Variables for query connections
# =============================================================================
server_04 = 'sqlsrv04'
db_04 = 'BKI_Datastore'
con_04 = pyodbc.connect(f'DRIVER=SQL Server;SERVER={server_04};DATABASE={db_04};autocommit=True')
params_04 = urllib.parse.quote_plus(f'DRIVER=SQL Server Native Client 11.0;SERVER={server_04};DATABASE={db_04};Trusted_Connection=yes')
engine_04 = create_engine(f'mssql+pyodbc:///?odbc_connect={params_04}')
cursor_04 = con_04.cursor()

server_nav = r'SQLSRV03\NAVISION'
db_nav = 'NAV100-DRIFT'
con_nav = pyodbc.connect(f'DRIVER=SQL Server;SERVER={server_nav};DATABASE={db_nav};Trusted_Connection=yes')
params_nav = urllib.parse.quote_plus(f'DRIVER=SQL Server Native Client 11.0;SERVER={server_nav};DATABASE={db_nav};Trusted_Connection=yes')
engine_nav = create_engine(f'mssql+pyodbc:///?odbc_connect={params_nav}')

server_comscale = r'comscale-bki\sqlexpress'
db_comscale = 'ComScaleDB'
con_comscale = pyodbc.connect(f'DRIVER=SQL Server;SERVER={server_comscale};DATABASE={db_comscale}')
params_comscale = urllib.parse.quote_plus(f'DRIVER=SQL Server Native Client 11.0;SERVER={server_comscale};DATABASE={db_comscale};Trusted_Connection=yes')
engine_comscale = create_engine(f'mssql+pyodbc:///?odbc_connect={params_comscale}')

server_probat = '192.168.125.161'
db_probat = 'BKI_IMP_EXP'
con_probat = pyodbc.connect(f'DRIVER=SQL Server;SERVER={server_probat};DATABASE={db_probat};uid=bki_read;pwd=Probat2016')
params_probat = urllib.parse.quote_plus(f'DRIVER=SQL Server Native Client 11.0;SERVER={server_probat};DATABASE={db_probat};Trusted_Connection=yes')
engine_probat = create_engine(f'mssql+pyodbc:///?odbc_connect={params_probat}')

# =============================================================================
# Read data from request
# =============================================================================
query_ds_request =  """ SELECT TOP 1 [Id] ,[Forespørgselstype],[Rapport_modtager]
                    ,[Referencenummer] ,[Note_forespørgsel] ,[Modtagelse]  ,[Ordrerelationstype]
                    FROM [trc].[Sporbarhed_forespørgsel]
                    WHERE [Forespørgsel_igangsat] IS NULL """
df_request = pd.read_sql(query_ds_request, con_04)

# Exit script if no request data is found
if len(df_request) == 0:
    raise SystemExit(0)

# =============================================================================
# Set request variables
# =============================================================================
req_type = df_request.loc[0, 'Forespørgselstype']
req_reference_no = df_request.loc[0, 'Referencenummer'].rstrip(' ')
req_recipients = df_request.loc[0, 'Rapport_modtager']
req_note = df_request.loc[0, 'Note_forespørgsel']
req_id = df_request.loc[0, 'Id']
req_modtagelse = df_request.loc[0, 'Modtagelse']
req_ordrelationstype = df_request.loc[0, 'Ordrerelationstype']

script_name = 'Sporbarhed_samlet.py'
dict_email_emne = {
    0: f'Anmodet rapport for ordre {req_reference_no}'
    ,1: f'Anmodet rapport for parti {req_reference_no}'
    ,2: 'Anmodet rapport for opspræt'
    ,3: f'Anmodet rapport for handelsvare {req_reference_no}'
    }
timestamp = datetime.now().strftime('%d-%m-%Y %H:%M:%S')
orders_top_level = [req_reference_no]
orders_related = []

# =============================================================================
# Update request that it is initiated and write into log
# =============================================================================
cursor_04.execute(f"""UPDATE [trc].[Sporbarhed_forespørgsel]
                  SET [Forespørgsel_igangsat] = getdate()
                  WHERE [Id] = {req_id} """)
cursor_04.commit()
log_insert(script_name, f'Request id: {req_id} initiated')

# =============================================================================
# Variables for files generated
# =============================================================================
filepath = r'\\filsrv01\BKI\11. Økonomi\04 - Controlling\NMO\4. Kvalitet\Sporbarhedstest\Tests via PowerApps'
file_name = f'Rapport_{req_reference_no}_{req_id}'

doc = docx.Document()
doc.add_heading(f'Rapport for produktionsordre {req_reference_no}',0)
doc.sections[0].header.paragraphs[0].text = f'{script_name}'
doc.sections[0].footer.paragraphs[0].text = f'{timestamp}'
doc.sections[0].page_width = docx.shared.Mm(297)
doc.sections[0].page_height = docx.shared.Mm(210)
doc.sections[0].top_margin = docx.shared.Mm(15)
doc.sections[0].bottom_margin = docx.shared.Mm(15)
doc.sections[0].left_margin = docx.shared.Mm(10)
doc.sections[0].right_margin = docx.shared.Mm(10)
doc.sections[0].orientation = docx.enum.section.WD_ORIENT.LANDSCAPE

doc_name = f'{file_name}.docx'
path_file_doc = filepath + r'\\' + doc_name

wb = openpyxl.Workbook()
wb_name = f'{file_name}.xlsx'
path_file_wb = filepath + r'\\' + wb_name
excel_writer = pd.ExcelWriter(path_file_wb, engine='xlsxwriter')

png_relations_name = f'{file_name}.png'
path_png_relations = filepath + r'\\' + png_relations_name

# =============================================================================
# Read setup for section for reporttypes. NAV querys with NOLOCK to prevent deadlocks
# =============================================================================
query_ds_reporttypes =  f"""SELECT SRS.[Sektion], SS.[Beskrivelse] AS [Sektion navn]
                       FROM [trc].[Sporbarhed_rapport_sektion] AS SRS
					   INNER JOIN [trc].[Sporbarhed_sektion] AS SS
					   ON SRS.[Sektion] = SS.[Id]
                       WHERE [Forespørgselstype] = {req_type} """
df_sections = pd.read_sql(query_ds_reporttypes, con_04)

# Query for Navision items, used for adding information to item numbers not queried directly from Navision
query_nav_items = """ SELECT [No_] AS [Nummer],[Description] AS [Beskrivelse]
                  ,[Item Category Code] AS [Varekategorikode]
				  ,CASE WHEN [Display Item] = 1 THEN 'Display'
				  WHEN [Item Category Code] = 'FÆR KAFFE' THEN 'Færdigkaffe'
				  WHEN [No_] LIKE '1040%' THEN 'Ristet kaffe'
				  WHEN [No_] LIKE '1050%' THEN 'Formalet kaffe'
				  WHEN [No_] LIKE '1020%' THEN 'Råkaffe'
				  ELSE [Item Category Code] END AS [Varetype]
                  FROM [dbo].[BKI foods a_s$Item] """
df_nav_items = pd.read_sql(query_nav_items, con_nav)

# Query for getting item numbers for production and assembly orders from Navision
query_nav_order_info = """ SELECT PAH.[No_] AS [Ordrenummer]
                       ,PAH.[Item No_] AS [Varenummer]
                       FROM [dbo].[BKI foods a_s$Posted Assembly Header] AS PAH
                       INNER JOIN [dbo].[BKI foods a_s$Item] AS I
                           ON PAH.[Item No_] = I.[No_]
                       WHERE I.[Item Category Code] = 'FÆR KAFFE'
                           AND I.[Display Item] = 1
                       UNION ALL
                       SELECT PO.[No_],PO.[Source No_]
                       FROM [dbo].[BKI foods a_s$Production Order] AS PO
                       INNER JOIN [dbo].[BKI foods a_s$Item] AS I
                           ON PO.[Source No_] = I.[No_]
                       WHERE PO.[Status] IN (2,3,4)
                           AND I.[Item Category Code] <> 'RÅKAFFE' """
df_nav_order_info = pd.read_sql(query_nav_order_info, con_nav)

# Query section log for each section logged per script-run.
# Query is only executed at the end of each function
query_ds_section_log = f""" SELECT	SL.[Sektion] AS [Sektionskode]
                       ,S.[Beskrivelse] AS [Sektion],SS.[Beskrivelse] AS [Status]
                       ,SL.[Fejlkode_script] AS [Fejlkode script], SL.[Registreringstidspunkt]
                       FROM [trc].[Sporbarhed_sektion_log] AS SL
                       INNER JOIN [trc].[Sporbarhed_sektion] AS S
                         	ON SL.[Sektion] = S.[Id]
                       INNER JOIN [trc].[Sporbarhed_statuskode] AS SS
                            ON SL.[Statuskode] = SS.[Id]
                       WHERE SL.[Forespørgsels_id] = {req_id} """

def rapport_færdigkaffe():
    # Query to read various information from BKI_Datastore for the order requested in the report
    query_ds_generelt = f""" WITH [KP] AS ( SELECT [Ordrenummer]
                    	,SUM( CASE WHEN [Prøvetype] = 0 THEN [Antal_prøver] ELSE 0 END) AS [Kontrolprøve]
                    	,SUM( CASE WHEN [Prøvetype] = 1 THEN [Antal_prøver] ELSE 0 END) AS [Referenceprøve]
                    	,SUM( CASE WHEN [Prøvetype] = 2 THEN [Antal_prøver] ELSE 0 END) AS [Henstandsprøve]
                        FROM [cof].[Kontrolskema_prøver]
                        GROUP BY [Ordrenummer] )
                        ,[SK] AS ( SELECT [Referencenummer] ,MAX([Status]) AS [Status]
                        FROM [cof].[Smageskema] WHERE [Referencetype] = 2
                        GROUP BY [Referencenummer] )
                        SELECT SF.[Referencenummer] AS [Ordrenummer] ,SF.[Pakketidspunkt] ,KH.[Igangsat_af] AS [Igangsat af]
                        ,KH.[Silo_opstart] AS [Opstartssilo] ,KH.[Taravægt] ,KH.[Nitrogen] / 100.0 AS [Nitrogen]
                        ,KH.[Bemærkning] AS [Bemærkning opstart] ,ISNULL(KP.[Kontrolprøve] ,0) AS [Kontrolprøver]
                        ,ISNULL(KP.[Referenceprøve] ,0) AS [Referenceprøver]
                        ,ISNULL(KP.[Henstandsprøve] ,0) AS [Henstandsprøver]
                        ,CASE WHEN SK.[Status] = 1 THEN 'Godkendt' WHEN SK.[Status] = 0 THEN 'Afvist'
                        ELSE 'Ej smagt' END AS [Smagning status], KH.[Pakkelinje]
                        FROM [trc].[Sporbarhed_forespørgsel] AS SF
                        LEFT JOIN [cof].[Kontrolskema_hoved] AS KH 
                            ON SF.[Referencenummer] = KH.[Ordrenummer]
                        LEFT JOIN [KP] 
                            ON SF.[Referencenummer] = KP.[Ordrenummer]
                        LEFT JOIN [SK] 
                            ON SF.[Referencenummer] = SK.[Referencenummer]
                        WHERE SF.[Id] = {req_id} """
    df_results_generelt = pd.read_sql(query_ds_generelt, con_04)
    
    production_machine = df_results_generelt['Pakkelinje'].iloc[0]
    
    # Query to get all samples registrered for the requested order.
    # Dataframe to be filtered later on to split by sample type.
    query_ds_samples = f""" SELECT KP.[Id],KP.[Ordrenummer],KP.[Registreringstidspunkt]
                	   ,KP.[Registreret_af] AS [Operatør],KP.[Bemærkning]
                       ,KP.[Prøvetype] AS [Prøvetype int],P.[Beskrivelse] AS [Prøvetype]
                       ,CASE WHEN KP.[Kontrol_mærkning] = 1 THEN 'Ok' 
                       WHEN KP.[Kontrol_mærkning] = 0 THEN 'Ej ok' END AS [Mærkning]
                       ,CASE WHEN KP.[Kontrol_rygsvejning] = 1	THEN 'Ok'
                       WHEN KP.[Kontrol_rygsvejning] = 0 THEN 'Ej ok' END AS [Rygsvejsning]
                       ,CASE WHEN KP.[Kontrol_ventil] = 1 THEN 'Ok'
                       WHEN KP.[Kontrol_ventil] = 0 THEN 'Ej ok' END AS [Ventil]
                       ,CASE WHEN KP.[Kontrol_peelbar] = 1	THEN 'Ok'
                       WHEN KP.[Kontrol_peelbar] = 0 THEN 'Ej ok' END AS [Peelbar]
                       ,CASE WHEN KP.[Kontrol_tintie] = 1 THEN 'Ok'
                       WHEN KP.[Kontrol_tintie] = 0 THEN 'Ej ok' END AS [Tintie]
    				   ,CASE WHEN KP.[Kontrol_tæthed] = 1 THEN 'Ok'
                       WHEN KP.[Kontrol_tæthed] = 0 THEN 'Ej ok' END AS [Tæthed]
                       ,KP.[Vægt_aflæst] AS [Vægt],KP.[Kontrol_ilt] / 100.0 AS [Ilt],KP.[Silo]
                       ,CASE WHEN SK.[Status] = 1 THEN 'Godkendt' WHEN SK.[Status] = 0
                       THEN 'Afvist' ELSE 'Ej smagt' END AS [Smagning status]
    				   ,KP.[Antal_prøver] AS [Antal prøver]
                       FROM [cof].[Kontrolskema_prøver] AS KP
                       INNER JOIN [cof].[Prøvetype] AS P 
                           ON KP.[Prøvetype] = P.[Id]
                       LEFT JOIN [cof].[Smageskema] AS SK
                           ON KP.[Id] = SK.[Id_org]
                           AND SK.[Id_org_kildenummer] = 6
                       WHERE KP.[Ordrenummer] = '{req_reference_no}' """
    df_prøver = pd.read_sql(query_ds_samples, con_04)
    
    # All grades given for the requested order. Coalesce is to ensure that query
    # returns no results if record exists but no grades have been given
    query_ds_karakterer = f""" SELECT [Id] ,[Dato] ,[Bruger] ,[Smag_Syre] AS [Syre]
                          ,[Smag_Krop] AS [Krop] ,[Smag_Aroma] AS [Aroma] 
                          ,[Smag_Eftersmag] AS [Eftersmag],[Smag_Robusta] AS [Robusta] ,[Bemærkning]
                          FROM [cof].[Smageskema]
                          WHERE [Referencetype] = 2	
                              AND [Referencenummer] = '{req_reference_no}'
                              AND COALESCE([Smag_Syre],[Smag_Krop],[Smag_Aroma],
                                [Smag_Eftersmag],[Smag_Robusta]) IS NOT NULL"""
    df_karakterer = pd.read_sql(query_ds_karakterer, con_04)
    
    # If lotnumbers from requested order have been checked for leakage the information
    # from the check is returned with this query. Will often return no results
    query_ds_vacslip = """ SELECT [Registreringstidspunkt] AS [Kontroltidspunkt]
                       ,[Initialer] AS [Kontrolleret af],[Lotnummer]
                       ,[Pallenummer],[Antal_poser] AS [Antal leakers]
                       ,[Bemærkning] AS [Kontrol bemærkning]
    				   ,CASE WHEN [Overført_email_log] = 1 THEN
    				   'Over grænseværdi' ELSE 'Ok' END AS [Resultat af kontrol]
                       FROM [cof].[Vac_slip] """
    df_ds_vacslip = pd.read_sql(query_ds_vacslip, con_04)
    
    # Primary packaging material - valve for bag
    query_ds_ventil = f""" SELECT [Varenummer] ,[Batchnr_stregkode] AS [Lotnummer]
                      FROM [cof].[Ventil_registrering]
                      WHERE [Ordrenummer] = '{req_reference_no}' """
    df_ds_ventil = pd.read_sql(query_ds_ventil, con_04)
    
    # Order statistics from Comscale. Only for good bags (trade)
    query_com_statistics = f""" WITH CTE AS ( SELECT SD.[Nominal] ,SD.[Tare]
                           ,SUM( SD.[MeanValueTrade] * SD.[CounterGoodTrade] ) AS [Total vægt]
                           ,SUM( SD.[StandardDeviationTrade] * SD.[CounterGoodTrade] ) AS [Std afv]
                           ,SUM( SD.[CounterGoodTrade] ) AS [Antal poser]
                           FROM [ComScaleDB].[dbo].[StatisticData] AS SD
                           INNER JOIN [dbo].[Statistic] AS S ON SD.[Statistic_ID] = S.[ID]
                           WHERE S.[Order] = '{req_reference_no}' AND lower(S.[ArticleNumber]) NOT LIKE '%k'
                           GROUP BY S.[Order],SD.[Nominal],SD.[Tare] )
                           SELECT CTE.[Total vægt] / 1000.0 AS [Total vægt kg],CTE.[Antal poser]
                           ,CASE WHEN CTE.[Antal poser] = 0 
                           THEN NULL ELSE CTE.[Total vægt] / CTE.[Antal poser] END AS [Middelvægt g]
                           ,CASE WHEN CTE.[Antal poser] = 0 
                           THEN NULL ELSE CTE.[Std afv] / CTE.[Antal poser] END AS [Standardafvigelse g]
                           ,CASE WHEN CTE.[Antal poser] = 0 
                           THEN NULL ELSE CTE.[Total vægt] / CTE.[Antal poser] END - CTE.[Nominal] AS [Gns. godvægt per enhed g]
                           ,CTE.[Total vægt] - CTE.[Nominal] * CTE.[Antal poser] AS [Godvægt total g]
                           ,CTE.[Nominal] AS [Nominel vægt g],CTE.[Tare] AS [Taravægt g]
                           FROM CTE """
    df_com_statistics = pd.read_sql(query_com_statistics, con_comscale)
    
    # Query to pull various information from Navision for the requested order.
    query_nav_generelt = f""" WITH [RECEPT] AS (
                         SELECT	POC.[Prod_ Order No_],I.[No_]
                         FROM [dbo].[BKI foods a_s$Prod_ Order Component] (NOLOCK) AS POC
                         INNER JOIN [dbo].[BKI foods a_s$Item] (NOLOCK) AS I
                        	ON POC.[Item No_] = I.[No_]
                         WHERE [POC].[Prod_ Order Line No_] = 10000
                        	AND I.[Sequence Code] = 4)
                         ,[ILE] AS ( SELECT [Order No_],MIN([Posting Date]) AS [Posting Date]
                         ,SUM(CASE WHEN [Entry Type] = 5 AND [Location Code] = 'REWORK' 
                              THEN [Quantity] ELSE 0 END) AS [Rework forbrug]
                         ,SUM(CASE WHEN [Entry Type] = 6 AND [Location Code] = 'REWORK' 
                              THEN [Quantity] ELSE 0 END) AS [Rework afgang]
                         ,SUM(CASE WHEN [Entry Type] = 5 AND [Location Code] = 'SLAT' 
                              THEN [Quantity] ELSE 0 END) AS [Slat forbrug]
                         ,SUM(CASE WHEN [Entry Type] = 6 AND [Location Code] = 'SLAT' 
                              THEN [Quantity] ELSE 0 END) AS [Slat afgang]
                         FROM [dbo].[BKI foods a_s$Item Ledger Entry] (NOLOCK)
                         WHERE [Order Type] = 1 GROUP BY [Order No_] )
                         SELECT PO.[Source No_] AS [Varenummer]
                         ,I.[Description] AS [Varenavn]
                         ,I.[Base Unit of Measure] AS [Basisenhed]
                         ,CASE WHEN PO.[Status] = 0 THEN 'Simuleret'
                         WHEN PO.[Status] = 1 THEN 'Planlagt'
                         WHEN PO.[Status] = 2 THEN 'Fastlagt'
                         WHEN PO.[Status] = 3 THEN 'Frigivet'
                         WHEN PO.[Status] = 4 THEN 'Færdig'
                         END AS [Prod.ordre status]
                         ,ICR.[Cross-Reference No_] AS [Stregkode]
                         ,RECEPT.[No_] AS [Receptnummer],ILE.[Rework afgang]
                         ,ILE.[Posting Date] AS [Produktionsdato]
                         ,ILE.[Rework forbrug],ILE.[Slat afgang],ILE.[Slat forbrug]
                         FROM [dbo].[BKI foods a_s$Production Order] (NOLOCK) AS PO
                         INNER JOIN [dbo].[BKI foods a_s$Item] (NOLOCK) AS I
                        	ON PO.[Source No_] = I.[No_]
                         LEFT JOIN [dbo].[BKI foods a_s$Item Cross Reference] (NOLOCK) AS ICR
                        	ON I.[No_] = ICR.[Item No_] AND ICR.[Unit of Measure] = 'PS'
                        	AND ICR.[Cross-Reference Type] = 3
                         LEFT JOIN [RECEPT] ON PO.[No_] = RECEPT.[Prod_ Order No_]
                         LEFT JOIN [ILE] ON PO.[No_] = ILE.[Order No_]
                         WHERE I.[Item Category Code] = 'FÆR KAFFE' AND PO.[No_] = '{req_reference_no}' """
    df_nav_generelt = pd.read_sql(query_nav_generelt, con_nav)
    
    production_date = df_nav_generelt['Produktionsdato'].iloc[0]
    
    # Control of scales in packing area, 3 days back and 1 day ahead of production date
    query_ds_vægtkontrol = f""" SELECT V.[Registreringstidspunkt]
                           ,V.[Registreret_af] AS [Registreret af],V.[Vægt],V.[Serienummer]
                           ,CASE WHEN V.[Status] = 1 THEN 'Ok' ELSE 'Ej ok' END AS [Status]
                           FROM [cof].[Vægtkontrol] AS V
                           INNER JOIN [cof].[Serienummer_pakkelinje] AS SP
                           ON V.[Serienummer] = SP.[Serienummer]
                           WHERE SP.[Pakkelinje] = '{production_machine}'
                           AND DATEADD(d, DATEDIFF(d, 0, V.[Registreringstidspunkt] ), 0) 
                           BETWEEN DATEADD(d,-3, '{production_date}') AND DATEADD(d, 1, '{production_date}') """
    df_ds_vægtkontrol = pd.read_sql(query_ds_vægtkontrol, con_04)
    
    # Get any related orders identified through Probat
    # Pakkelinjer is used to find either grinding or roasting orders used directly in packaging
    # Mølleordrer is used to find roasting orders used for grinding orders
    query_probat_orders = f""" WITH [CTE_ORDERS_PACK] AS (
                           SELECT [ORDER_NAME] AS [Ordrenummer],[S_ORDER_NAME] AS [Relateret ordre]
                           ,'Probat formalet pakkelinje' AS [Kilde]
                           FROM [dbo].[PRO_EXP_ORDER_SEND_PG]
                           GROUP BY	[ORDER_NAME],[S_ORDER_NAME]
                           UNION ALL
                           SELECT [ORDER_NAME],[S_ORDER_NAME],'Probat helbønne pakkelinje'
                           FROM [dbo].[PRO_EXP_ORDER_SEND_PB]
                           GROUP BY	[ORDER_NAME],[S_ORDER_NAME] )
    					   ,[CTE_ORDERS] AS (
                           SELECT [Ordrenummer],[Relateret ordre],[Kilde]
                           FROM [CTE_ORDERS_PACK]
                           WHERE [Relateret ordre] IN (SELECT [Relateret ordre] 
                           FROM [CTE_ORDERS_PACK] WHERE [Ordrenummer] = '{req_reference_no}'))
    					   SELECT * 
    					   FROM [CTE_ORDERS]
                           WHERE [Relateret ordre] <> 'Retour Ground'
    					   UNION ALL
    					   SELECT [ORDER_NAME] AS [Ordrenummer]
    					   ,[S_ORDER_NAME] AS [Relateret ordre]
    					   ,'Probat mølle' AS [Kilde]
    					   FROM [dbo].[PRO_EXP_ORDER_LOAD_G]
    					   WHERE [ORDER_NAME] IN (SELECT [Relateret ordre] FROM [CTE_ORDERS])
    					   GROUP BY [S_ORDER_NAME],[ORDER_NAME] """
    df_probat_orders = pd.read_sql(query_probat_orders, con_probat)
    
    # Get lists of orders and related orders (if any) from Probat, first create dataframe with top level orders:
    df_temp_top_level = df_probat_orders.loc[df_probat_orders['Kilde'] != 'Probat mølle']
    probat_orders_top = df_temp_top_level['Ordrenummer'].unique().tolist()
    probat_orders_related = df_probat_orders['Relateret ordre'].unique().tolist()
    
    # Get related orders from Navision
    query_nav_order_related = f"""WITH [CTE_ORDER] AS (SELECT [Prod_ Order No_]
                       ,[Reserved Prod_ Order No_]
                       FROM [dbo].[BKI foods a_s$Reserved Prod_ Order No_]
                       WHERE [Prod_ Order No_] = '{req_reference_no}' 
                       AND [Invalid] = 0)
                       SELECT [Prod_ Order No_] AS [Ordrenummer] 
                       ,[Reserved Prod_ Order No_] AS [Relateret ordre]
                       ,'Navision reservationer' AS [Kilde]
                       FROM [dbo].[BKI foods a_s$Reserved Prod_ Order No_]
                       WHERE [Reserved Prod_ Order No_] IN 
                       (SELECT [Reserved Prod_ Order No_] FROM [CTE_ORDER] )
                       AND [Invalid] = 0 """
    df_nav_order_related = pd.read_sql(query_nav_order_related, con_nav)
    
    # Get list of orders and append to lists if they do not already exist
    # Merge Probat and NAV orders before merging
    nav_orders_top = df_nav_order_related['Ordrenummer'].unique().tolist()
    nav_orders_related = df_nav_order_related['Relateret ordre'].unique().tolist()
    
    # Create strings dependent on request relationsship type, defined when report is requested by user
    if req_ordrelationstype == 0: # All
        temp_orders_top = probat_orders_top + nav_orders_top
        temp_orders_related = probat_orders_related + nav_orders_related
    elif req_ordrelationstype == 1: # Just Probat
        temp_orders_top = probat_orders_top
        temp_orders_related = probat_orders_related
    elif req_ordrelationstype == 2: # Just Navision
        temp_orders_top = nav_orders_top
        temp_orders_related = nav_orders_related
    
    # If order doesn't exist in list, append:
    for order in temp_orders_top:
        if order not in  orders_top_level and order != '':
            orders_top_level.append(order)
    
    for order in temp_orders_related:
        if order not in orders_related:
            orders_related.append(order)
    
    req_orders_total = string_to_sql(orders_top_level) # String used for querying Navision, only finished goods
    
    # Recursive query to find all relevant produced orders related to the requested order
    # First is identified all lotnumbers related to the orders identified through NAV reservations (only production orders)
    # Next is a recursive part which identifies any document numbers which have consumed these lotnumbers (ILE_C)
    # Which is then queried again to find all lotnumbers produced on the orders from which these lotnumbers originally came.
    
    #First we find all relevant lot nos and store in string to be used in queries below
    query_nav_lotnos_total = f""" WITH [LOT_ORG] AS ( SELECT [Lot No_]
                                  FROM [dbo].[BKI foods a_s$Item Ledger Entry] (NOLOCK)
                                  WHERE [Order No_] IN ({req_orders_total})
                                  AND [Entry Type] = 6
                                  UNION ALL
                                  SELECT ILE_O.[Lot No_]
                                  FROM [LOT_ORG]
                                  INNER JOIN [dbo].[BKI foods a_s$Item Ledger Entry] (NOLOCK) AS ILE_C
                                      ON [LOT_ORG].[Lot No_] = ILE_C.[Lot No_]
                                      AND [ILE_C].[Entry Type] IN (5,8)
                                  INNER JOIN [dbo].[BKI foods a_s$Item Ledger Entry] (NOLOCK) AS ILE_O
                                	  ON ILE_C.[Document No_] = ILE_O.[Document No_]
                                      AND ILE_O.[Entry Type] IN (6,9)
                                  INNER JOIN [dbo].[BKI foods a_s$Item] (NOLOCK) AS I
    								  ON ILE_O.[Item No_] = I.[No_]
    								  WHERE I.[Item Category Code] = 'FÆR KAFFE')
                                  SELECT [Lot No_] AS [Lot]
                                  FROM [LOT_ORG] GROUP BY [Lot No_] """
    df_nav_lotnos_total = pd.read_sql(query_nav_lotnos_total, con_nav)
    nav_lotnots_total_sql_string = string_to_sql(df_nav_lotnos_total['Lot'].unique().tolist())
    
    query_nav_færdigvaretilgang = f""" WITH [LOT_SINGLE] AS ( SELECT [Lot No_], [Document No_] AS [Ordrenummer]
                                  FROM [dbo].[BKI foods a_s$Item Ledger Entry] (NOLOCK) 
    							  WHERE [Entry Type] IN (6,9)
    							  GROUP BY [Lot No_], [Document No_])
                                  SELECT ILE.[Item No_] AS [Varenummer],I.[Description] AS [Varenavn], LOT_SINGLE.[Ordrenummer]
                            	  ,SUM(CASE WHEN ILE.[Entry Type] IN (0,6,9)
                            		THEN ILE.[Quantity] * I.[Net Weight]
                            		ELSE 0 END) AS [Produceret]
                            	,SUM(CASE WHEN ILE.[Entry Type] = 1
                            		THEN ILE.[Quantity] * I.[Net Weight] * -1
                            		ELSE 0 END) AS [Salg]
                            	,SUM(CASE WHEN ILE.[Entry Type] NOT IN (0,1,6,9)
                            		THEN ILE.[Quantity] * I.[Net Weight] * -1
                            		ELSE 0 END) AS [Regulering & ompak]
                            	,SUM(ILE.[Remaining Quantity] * I.[Net Weight]) AS [Restlager]
                                FROM [dbo].[BKI foods a_s$Item Ledger Entry] (NOLOCK) AS ILE
                                INNER JOIN [dbo].[BKI foods a_s$Item] (NOLOCK) AS I
                                	ON ILE.[Item No_] = I.[No_]
                                INNER JOIN [LOT_SINGLE]
                                	ON ILE.[Lot No_] = [LOT_SINGLE].[Lot No_]
    							WHERE ILE.[Lot No_] IN ({nav_lotnots_total_sql_string})
                                GROUP BY ILE.[Item No_],I.[Description], LOT_SINGLE.[Ordrenummer] """
    df_nav_færdigvaretilgang = pd.read_sql(query_nav_færdigvaretilgang, con_nav)
    
    # Recursive query to get all customer who purchased identified lotnumbers.
    # See explanation of query above
    query_nav_debitorer = f"""   WITH [LOT_SINGLE] AS ( SELECT [Lot No_], [Document No_] AS [Produktionsordrenummer]
                          FROM [dbo].[BKI foods a_s$Item Ledger Entry] (NOLOCK)
    					  WHERE [Entry Type] IN (6,9) 
    					  GROUP BY [Lot No_],[Document No_] )
    
                          SELECT C.[No_] AS [Debitornummer],C.[Name] AS [Debitornavn], LOT_SINGLE.[Produktionsordrenummer]
                        	  ,ILE.[Posting Date] AS [Dato]
                        	  ,ILE.[Item No_] AS [Varenummer]
                        	  ,SUM(ILE.[Quantity] * -1) AS [Enheder]
                        	  ,SUM(ILE.[Quantity] * I.[Net Weight] * -1) AS [Kilo]
                          FROM [dbo].[BKI foods a_s$Item Ledger Entry] (NOLOCK) AS ILE
                          INNER JOIN [dbo].[BKI foods a_s$Item] (NOLOCK) AS I
                        	  ON ILE.[Item No_] = I.[No_]
                          INNER JOIN [LOT_SINGLE]
                          	  ON ILE.[Lot No_] = [LOT_SINGLE].[Lot No_]
                          INNER JOIN [dbo].[BKI foods a_s$Customer] (NOLOCK) AS C
                        	  ON ILE.[Source No_] = C.[No_]
                          WHERE ILE.[Entry Type] = 1
    						AND ILE.[Lot No_] IN ({nav_lotnots_total_sql_string})
                          GROUP BY  C.[No_] ,C.[Name],ILE.[Posting Date],ILE.[Item No_], LOT_SINGLE.[Produktionsordrenummer]  """
    df_nav_debitorer = pd.read_sql(query_nav_debitorer, con_nav)
    
    # Query to show relation between requested order and any orders which have used it as components
    query_nav_orders = f""" WITH [LOT_ORG] AS ( SELECT [Lot No_]
                                  FROM [dbo].[BKI foods a_s$Item Ledger Entry] (NOLOCK)
                                  WHERE [Order No_] IN ({req_orders_total})
                                  AND [Entry Type] = 6
                                  UNION ALL
                                  SELECT ILE_O.[Lot No_]
                                  FROM [LOT_ORG]
                                  INNER JOIN [dbo].[BKI foods a_s$Item Ledger Entry] (NOLOCK) AS ILE_C
                                      ON [LOT_ORG].[Lot No_] = ILE_C.[Lot No_]
                                      AND [ILE_C].[Entry Type] IN (5,8)
                                  INNER JOIN [dbo].[BKI foods a_s$Item Ledger Entry] (NOLOCK) AS ILE_O
                                	  ON ILE_C.[Document No_] = ILE_O.[Document No_]
                                      AND ILE_O.[Entry Type] IN (6,9) )
                                  ,[DOC_CONS] AS ( SELECT [Lot No_], [Document No_]
                                  FROM [dbo].[BKI foods a_s$Item Ledger Entry] (NOLOCK)
                                  WHERE [Entry Type] IN (5,8)
                                  GROUP BY [Lot No_], [Document No_] )
                                  ,[DOC_OUT] AS ( SELECT [Lot No_], [Document No_]
                                  FROM [dbo].[BKI foods a_s$Item Ledger Entry] (NOLOCK)
                                  WHERE [Entry Type] IN (6,9)
                                  GROUP BY [Lot No_], [Document No_] )
                                  SELECT DO.[Document No_] AS [Relateret ordre]
                                  ,DC.[Document No_] AS [Ordrenummer]
                                  ,'Navision forbrug' AS [Kilde]
                                  FROM [LOT_ORG] AS L
                                  INNER JOIN [DOC_OUT] AS DO
                                      ON L.[Lot No_] = DO.[Lot No_]
                                  LEFT JOIN [DOC_CONS] AS DC
                                      ON L.[Lot No_] = DC.[Lot No_]
                                  WHERE DC.[Document No_] IS NOT NULL
                                  GROUP BY DO.[Document No_] ,DC.[Document No_] """
    df_nav_orders = pd.read_sql(query_nav_orders, con_nav)
    
    # Lotnumber information for the originally requested order
    query_nav_lotno = f""" SELECT ILE.[Lot No_] AS [Lotnummer]
                	  ,LI.[Certificate Number] AS [Pallenummer]
                  	  ,[Quantity] * I.[Net Weight] AS [Kilo]
                	  ,CAST(ROUND(ILE.[Quantity] / IUM.[Qty_ per Unit of Measure],0) AS INT) AS [Antal poser]
                	  ,DATEADD(hour, 1, ILE.[Produktionsdato_-tid]) AS [Produktionstidspunkt]
                      FROM [dbo].[BKI foods a_s$Item Ledger Entry] (NOLOCK) ILE
                      INNER JOIN [dbo].[BKI foods a_s$Item] (NOLOCK) AS I
                          ON ILE.[Item No_] = I.[No_]
                      LEFT JOIN [dbo].[BKI foods a_s$Lot No_ Information] (NOLOCK) AS LI
                    	  ON ILE.[Lot No_] = LI.[Lot No_]
                          AND ILE.[Item No_] = LI.[Item No_]
                      LEFT JOIN [dbo].[BKI foods a_s$Item Unit of Measure] (NOLOCK) AS IUM
                    	  ON ILE.[Item No_] = IUM.[Item No_]
                          AND IUM.[Code] = 'PS'
                      WHERE ILE.[Order Type] = 1
                    	  AND ILE.[Entry Type] = 6
                          AND ILE.[Order No_] = '{req_reference_no}' """
    df_nav_lotno = pd.read_sql(query_nav_lotno, con_nav)
    
    # Primary packaging components used for the originally requested order
    query_nav_components = f""" SELECT POC.[Item No_] AS [Varenummer]
                    	   ,I.[Description] AS [Varenavn]
                           ,POAC.[Purchase Order No_] AS [Købsordre]
                           ,POAC.[Roll No_] AS [Rullenummer]
                           ,CAST(POAC.[Roll Lenght] AS INT) AS [Rullelængde]
                           ,POAC.[Batch_Lot No_] AS [Lotnummer]
                           ,POAC.[Packaging Date] AS [Pakkedato]
                           FROM [dbo].[BKI foods a_s$Prod_ Order Add_ Comp_] (NOLOCK) AS POAC
                           INNER JOIN [dbo].[BKI foods a_s$Prod_ Order Component] (NOLOCK) AS POC
                               ON POAC.[Prod_ Order No_] = POC.[Prod_ Order No_]
                               AND POAC.[Prod_ Order Line No_] = POC.[Prod_ Order Line No_]
                               AND POAC.[Prod_ Order Component Line No_] = POC.[Line No_]
                           INNER JOIN [dbo].[BKI foods a_s$Item] (NOLOCK) AS I
                               ON POC.[Item No_] = I.[No_]
                           WHERE POAC.[Prod_ Order No_] = '{req_reference_no}' """
    df_nav_components = pd.read_sql(query_nav_components, con_nav)
    
    # Components used for the originally requested order
    query_nav_consumption = f""" SELECT	ILE.[Item No_] AS [Varenummer]
                        	,I.[Description] AS [Varenavn]
                            ,I.[Base Unit of Measure] AS [Basisenhed]
                            ,SUM(ILE.[Quantity]) * -1 AS [Antal]
                            FROM [dbo].[BKI foods a_s$Item Ledger Entry] (NOLOCK) AS ILE
                            INNER JOIN [dbo].[BKI foods a_s$Item] (NOLOCK) AS I
                            	ON ILE.[Item No_] = I.[No_]
                            WHERE ILE.[Order No_] = '{req_reference_no}'
                            	AND ILE.[Entry Type] = 5
                            GROUP BY ILE.[Item No_] ,I.[Description],I.[Base Unit of Measure] """
    df_nav_consumption = pd.read_sql(query_nav_consumption, con_nav)
    
    q_related_orders = string_to_sql(orders_related)
    
    # Related grinding orders - information for batches out of grinder to include rework
    query_probat_ulg = f""" SELECT DATEADD(D, DATEDIFF(D, 0, [RECORDING_DATE] ), 0) AS [Dato]
                       ,[PRODUCTION_ORDER_ID] AS [Probat id] ,[SOURCE_NAME] AS [Mølle]
                       ,[ORDER_NAME] AS [Ordrenummer] ,[D_CUSTOMER_CODE] AS [Receptnummer]
                       ,[DEST_NAME] AS [Silo],SUM([WEIGHT]) / 1000.0 AS [Kilo]
                       FROM [dbo].[PRO_EXP_ORDER_UNLOAD_G]
                       WHERE [ORDER_NAME] IN ({q_related_orders})
                       GROUP BY [PRODUCTION_ORDER_ID],[ORDER_NAME],[DEST_NAME],[SOURCE_NAME]
                       ,[D_CUSTOMER_CODE], DATEADD(D, DATEDIFF(D, 0, [RECORDING_DATE] ), 0) """
    df_probat_ulg = pd.DataFrame(columns=['Dato','Probat id','Mølle','Ordrenummer',
                                          'Receptnummer','Silo','Kilo'])
    if len(q_related_orders) != 0:
        df_probat_ulg = pd.read_sql(query_probat_ulg, con_probat)
    
    # Find related roasting orders from any related grinding orders
    query_probat_lg = f""" SELECT [S_ORDER_NAME]
                           FROM [dbo].[PRO_EXP_ORDER_LOAD_G]
                           WHERE [ORDER_NAME] IN ({q_related_orders})
                           GROUP BY	[S_ORDER_NAME] """
    df_probat_lg = pd.DataFrame(columns=['S_ORDER_NAME'])
    if len(q_related_orders) != 0:
        df_probat_lg = pd.read_sql(query_probat_lg, con_probat)
    
    if len(df_probat_ulg) != 0: # Add to list only if dataframe is not empty
        for order in df_probat_lg['S_ORDER_NAME'].unique().tolist():
            if order not in orders_related:
                orders_related.append(order)
    
    q_related_orders = string_to_sql(orders_related)
    
    # Find information for identified roasting orders, batches out of roaster
    query_probat_ulr = f""" SELECT [S_CUSTOMER_CODE] AS [Receptnummer]
                            ,DATEADD(D, DATEDIFF(D, 0, [RECORDING_DATE] ), 0) AS [Dato]
                            ,[SOURCE_NAME] AS [Rister] ,[PRODUCTION_ORDER_ID] AS [Probat id]
                        	,[ORDER_NAME] AS [Ordrenummer] ,SUM([WEIGHT]) / 1000.0 AS [Kilo]
    						,[DEST_NAME] AS [Silo]
                            FROM [dbo].[PRO_EXP_ORDER_UNLOAD_R]
                            WHERE [ORDER_NAME] IN ({q_related_orders})
                            GROUP BY [S_CUSTOMER_CODE],[SOURCE_NAME],[PRODUCTION_ORDER_ID]
                            ,[ORDER_NAME],DATEADD(D, DATEDIFF(D, 0, [RECORDING_DATE] ), 0)
    						,[DEST_NAME] """
    df_probat_ulr = pd.DataFrame(columns=['Receptnummer','Dato','Rister','Probat id',
                                          'Ordrenummer','Kilo','Silo'])
    if len(q_related_orders) != 0:
        df_probat_ulr = pd.read_sql(query_probat_ulr, con_probat)
    
    # Find green coffee related to orders
    query_probat_lr = f""" SELECT [S_TYPE_CELL] AS [Sortnummer] ,[Source] AS [Silo]
                    ,[S_CONTRACT_NO] AS [Kontraktnummer]
                    ,[S_DELIVERY_NAME] AS [Modtagelse],[ORDER_NAME] AS [Ordrenummer]
                	,SUM([WEIGHT]) / 1000.0 AS [Kilo]
                    FROM [dbo].[PRO_EXP_ORDER_LOAD_R]
                    WHERE [ORDER_NAME] IN ({q_related_orders})
                    GROUP BY [S_TYPE_CELL],[Source],[S_CONTRACT_NO]
                    	,[S_DELIVERY_NAME],[ORDER_NAME] """
    df_probat_lr = pd.DataFrame(columns=['Sortnummer','Silo','Kontraktnummer',
                                         'Modtagelse','Ordrenummer','Kilo'])
    if len(q_related_orders) != 0:
        df_probat_lr = pd.read_sql(query_probat_lr, con_probat)
    
    # =============================================================================
    # Section 1: Generelt
    # =============================================================================
    section_id = 1
    section_name = get_section_name(section_id)
    column_order = ['Varenummer', 'Varenavn', 'Basisenhed','Stregkode', 'Receptnummer',
                    'Pakkelinje', 'Produktionsdato', 'Pakketidspunkt', 'Ordrenummer',
                    'Prod.ordre status', 'Smagning status', 'Opstartssilo',
                    'Igangsat af', 'Taravægt', 'Nitrogen', 'Henstandsprøver',
                    'Referenceprøver', 'Kontrolprøver', 'Bemærkning opstart',
                    'Lotnumre produceret', 'Slat forbrug','Slat afgang',
                    'Rework forbrug', 'Rework afgang']
    columns_1_dec = ['Slat forbrug', 'Slat afgang', 'Rework forbrug', 'Rework afgang',
                     'Taravægt']
    columns_0_dec = ['Henstandsprøver','Referenceprøver','Kontrolprøver']
    columns_0_pct = ['Nitrogen']
    
    if get_section_status_code(df_nav_generelt) == 99:
        try:
            df_nav_generelt['Pakkelinje'] = df_results_generelt['Pakkelinje'].iloc[0]
            df_nav_generelt['Pakketidspunkt'] = df_results_generelt['Pakketidspunkt'].iloc[0]
            df_nav_generelt['Ordrenummer'] = req_reference_no
            df_nav_generelt['Smagning status'] = df_results_generelt['Smagning status'].iloc[0]
            df_nav_generelt['Opstartssilo'] = df_results_generelt['Opstartssilo'].iloc[0]
            df_nav_generelt['Igangsat af'] = df_results_generelt['Igangsat af'].iloc[0]
            df_nav_generelt['Taravægt'] = df_results_generelt['Taravægt'].iloc[0]
            df_nav_generelt['Nitrogen'] = df_results_generelt['Nitrogen'].iloc[0]
            df_nav_generelt['Lotnumre produceret'] = len(df_nav_lotno)
            df_nav_generelt['Henstandsprøver'] = df_results_generelt['Henstandsprøver'].iloc[0]
            df_nav_generelt['Referenceprøver'] = df_results_generelt['Referenceprøver'].iloc[0]
            df_nav_generelt['Kontrolprøver'] = df_results_generelt['Kontrolprøver'].iloc[0]
            df_nav_generelt['Bemærkning opstart'] = df_results_generelt['Bemærkning opstart'].iloc[0]
            # Apply column formating
            for col in columns_1_dec:
                df_nav_generelt[col] = df_nav_generelt[col].apply(lambda x: number_format(x, 'dec_1'))
            for col in columns_0_dec:
                df_nav_generelt[col] = df_nav_generelt[col].apply(lambda x: number_format(x, 'dec_0'))
            for col in columns_0_pct:
                df_nav_generelt[col] = df_nav_generelt[col].apply(lambda x: number_format(x, 'pct_0'))
            df_nav_generelt['Produktionsdato'] = df_nav_generelt['Produktionsdato'].dt.strftime('%d-%m-%Y')
            # Transpose dataframe
            df_nav_generelt = df_nav_generelt[column_order].transpose()
            df_nav_generelt = df_nav_generelt.reset_index()
            df_nav_generelt.columns = ['Sektion','Værdi']
            # Write results to Word and Excel
            insert_dataframe_into_excel (df_nav_generelt, section_name, True)
            add_section_to_word(df_nav_generelt, section_name, True, [0])
            # Write status into log
            section_log_insert(section_id, 0)
        except Exception as e: # Insert error into log
            section_log_insert(section_id, 2, e)
    else: # Write into log if no data is found or section is out of scope
        section_log_insert(section_id, get_section_status_code(df_nav_generelt))
    
    # =============================================================================
    # Section 2: Relaterede ordrer NAV --> Probat
    # =============================================================================
    section_id = 2
    section_name = get_section_name(section_id)
    column_order = ['Ordrenummer','Varenummer','Navn','Relateret ordre',
                    'Relateret vare','Relateret navn','Kilde']
    
    if req_ordrelationstype == 0:
        df_temp_orders = pd.concat([df_nav_orders,df_probat_orders,df_nav_order_related])
    elif req_ordrelationstype == 1:
        df_temp_orders = pd.concat([df_nav_orders,df_probat_orders])
    elif req_ordrelationstype == 2:
        df_temp_orders = pd.concat([df_nav_orders,df_nav_order_related
                                    ,df_probat_orders.loc[df_probat_orders['Kilde'] == 'Probat mølle']]) # Only Probat orders which are not related to finished goods
    
    if get_section_status_code(df_temp_orders) == 99:
        try:
            df_temp_orders['Varenummer'] = df_temp_orders['Ordrenummer'].apply(lambda x: get_nav_order_info(x))
            df_temp_orders['Navn'] = df_temp_orders['Varenummer'].apply(lambda x: get_nav_item_info(x, 'Beskrivelse'))
            df_temp_orders['Relateret vare'] = df_temp_orders['Relateret ordre'].apply(lambda x: get_nav_order_info(x))
            df_temp_orders['Relateret navn'] = df_temp_orders['Relateret vare'].apply(lambda x: get_nav_item_info(x, 'Beskrivelse'))
            df_temp_orders = df_temp_orders[column_order]
            df_temp_orders.sort_values(by=['Ordrenummer','Relateret ordre'], inplace=True)
            # Write results to Word and Excel
            insert_dataframe_into_excel (df_temp_orders, section_name, False)
            add_section_to_word(df_temp_orders, section_name, True, [0])
            # Write status into log
            section_log_insert(section_id, 0)
    # =================================================================
    # Section 19: Relation visualization
    # =================================================================
            #Try to create .png with relations illustrated and add to .docx as well
            try:
                df_temp_order_relation = df_temp_orders[['Ordrenummer','Varenummer','Relateret ordre','Relateret vare']]
                df_temp_order_relation['Ordretype'] = df_temp_order_relation['Varenummer'].apply(lambda x: get_nav_item_info(x, 'Varetype'))
                df_temp_order_relation['Relateret ordretype'] = df_temp_order_relation['Relateret vare'].apply(lambda x: get_nav_item_info(x, 'Varetype'))
                df_temp_order_relation['Primær'] = df_temp_order_relation['Ordretype'] + '\n' + df_temp_order_relation['Ordrenummer']
                df_temp_order_relation['Sekundær'] = df_temp_order_relation['Relateret ordretype'] + '\n' + df_temp_order_relation['Relateret ordre']
                df_temp_order_relation = df_temp_order_relation[['Primær','Sekundær']]
                # Add green coffees
                df_temp_gc_orders = pd.DataFrame(columns=['Primær','Sekundær'])
                df_temp_gc_orders['Primær'] = 'Ristet kaffe' + '\n' + df_probat_lr['Ordrenummer']
                df_temp_gc_orders['Sekundær'] = 'Råkaffe' + '\n' + df_probat_lr['Kontraktnummer'] + '/' + df_probat_lr['Modtagelse']
                df_order_relations = pd.concat([df_temp_order_relation,df_temp_gc_orders])
                # Create relation visualization
                array_for_drawing = list(df_order_relations.itertuples(index=False, name=None))
                graph = nx.DiGraph()
                graph.add_edges_from(array_for_drawing)
                relations_plot = nx.drawing.nx_pydot.to_pydot(graph)
                relations_plot.write_png(path_png_relations)
                # Add image to word document
                doc.add_picture(path_png_relations, width=Inches(11.0), height=Inches(6.50))
                # Write to log
                section_log_insert(19, 0)
            except Exception as e: # Insert error into log. Same section_id as others..
                section_log_insert(19, 2, e)         
        except Exception as e: # Insert error into log
            section_log_insert(section_id, 2, e)
    else: # Write into log if no data is found or section is out of scope
        section_log_insert(section_id, get_section_status_code(df_temp_orders))
    
    # =============================================================================
    # Section 3: Færdigvaretilgang
    # =============================================================================
    section_id = 3
    section_name = get_section_name(section_id)
    column_order = ['Varenummer','Varenavn','Ordrenummer','Produceret','Salg','Restlager','Regulering & ompak']
    columns_1_dec = ['Produceret','Salg','Restlager','Regulering & ompak']
    columns_strip = ['Ordrenummer']
    
    if get_section_status_code(df_nav_færdigvaretilgang) == 99:
        try:
            # Concat order numbers to one string
            df_nav_færdigvaretilgang = df_nav_færdigvaretilgang.groupby(['Varenummer','Varenavn']).agg(
               {'Ordrenummer': lambda x: ','.join(sorted(pd.Series.unique(x))),
                'Produceret': 'sum',
                'Salg': 'sum',
                'Restlager': 'sum',
                'Regulering & ompak': 'sum'
               }).reset_index()
            # Remove trailing and leading commas
            for col in columns_strip:
                df_nav_færdigvaretilgang[col] = df_nav_færdigvaretilgang[col].apply(lambda x: strip_comma_from_string(x))
            # Create total for dataframe
            dict_færdigvare_total = {'Produceret': [df_nav_færdigvaretilgang['Produceret'].sum()],
                                     'Salg': [df_nav_færdigvaretilgang['Salg'].sum()],
                                     'Restlager': [df_nav_færdigvaretilgang['Restlager'].sum()],
                                     'Regulering & ompak': [df_nav_færdigvaretilgang['Regulering & ompak'].sum()]}
            df_temp_total = pd.concat([df_nav_færdigvaretilgang,
                                      pd.DataFrame.from_dict(data=dict_færdigvare_total, orient='columns')])
            df_temp_total = df_temp_total[column_order]
            df_temp_total.sort_values(by=['Varenummer'], inplace=True)
            # Data formating
            for col in columns_1_dec:
                df_temp_total[col] = df_temp_total[col].apply(lambda x: number_format(x, 'dec_1'))
            # Write results to Word and Excel
            insert_dataframe_into_excel (df_temp_total, section_name, False)
            add_section_to_word(df_temp_total, section_name, True, [-1,0])
            # Write status into log
            section_log_insert(section_id, 0)
        except Exception as e: # Insert error into log
            section_log_insert(section_id, 2, e)
    else: # Write into log if no data is found or section is out of scope
        section_log_insert(section_id, get_section_status_code(df_temp_total))
    
    # =============================================================================
    # Section 4: Mølleordrer
    # =============================================================================
    section_id = 4
    section_name = get_section_name(section_id)
    column_order = ['Receptnummer', 'Receptnavn', 'Dato', 'Mølle',
                    'Probat id', 'Ordrenummer', 'Silo', 'Kilo']
    columns_1_dec = ['Kilo']
    columns_strip = ['Dato','Silo','Mølle']
    
    if get_section_status_code(df_probat_ulg) == 99:
        try:
            # Create total for dataframe
            dict_mølle_total = {'Kilo': [df_probat_ulg['Kilo'].sum()],'Probat id':None}
            # Look up column values and string format datecolumn for export
            df_probat_ulg['Receptnavn'] = df_probat_ulg['Receptnummer'].apply(get_nav_item_info, field='Beskrivelse')
            df_probat_ulg['Dato'] = df_probat_ulg['Dato'].dt.strftime('%d-%m-%Y')
            # Join multiple dates or silos to one commaseparated string
            df_probat_ulg = df_probat_ulg.groupby(['Receptnummer', 'Receptnavn', 
                                                   'Probat id', 'Ordrenummer']).agg(
                                                       {'Silo': lambda x: ','.join(sorted(pd.Series.unique(x))),
                                                        'Dato': lambda x: ','.join(sorted(pd.Series.unique(x))),
                                                        'Mølle': lambda x: ','.join(sorted(pd.Series.unique(x))),
                                                        'Kilo': 'sum'
                                                       }).reset_index()
            # Remove trailing and leading commas
            for col in columns_strip:
                df_probat_ulg[col] = df_probat_ulg[col].apply(lambda x: strip_comma_from_string(x))
            # Create temp dataframe with total
            df_temp_total = pd.concat([df_probat_ulg, pd.DataFrame.from_dict(data=dict_mølle_total, orient='columns')])
            df_temp_total = df_temp_total[column_order]
            df_temp_total.sort_values(by=['Receptnummer','Ordrenummer'], inplace=True)
            # Data formating
            for col in columns_1_dec:
                df_temp_total[col] = df_temp_total[col].apply(lambda x: number_format(x, 'dec_1'))
            # Write results to Word and Excel
            insert_dataframe_into_excel (df_temp_total, section_name, False)
            add_section_to_word(df_temp_total, section_name, False, [-1,0])
            # Write status into log
            section_log_insert(section_id, 0)
        except Exception as e: # Insert error into log
            section_log_insert(section_id, 2, e)
    else: # Write into log if no data is found or section is out of scope
        section_log_insert(section_id, get_section_status_code(df_probat_ulg))
    
    # =============================================================================
    # Section 5: Risteordrer
    # =============================================================================
    section_id = 5
    section_name = get_section_name(section_id)
    column_order = ['Receptnummer', 'Receptnavn', 'Dato', 'Rister',
                    'Probat id', 'Ordrenummer', 'Silo', 'Kilo']
    columns_1_dec = ['Kilo']
    columns_strip = ['Dato','Silo']
    
    if get_section_status_code(df_probat_ulr) == 99:
        try:
            # Create total for dataframe
            dict_rister_total = {'Kilo':[df_probat_ulr['Kilo'].sum()],'Probat id':None}
            # Look up column values and string format datecolumn for export
            df_probat_ulr['Receptnavn'] = df_probat_ulr['Receptnummer'].apply(get_nav_item_info, field='Beskrivelse')
            df_probat_ulr['Dato'] = df_probat_ulr['Dato'].dt.strftime('%d-%m-%Y')
            # Join multiple dates or silos to one commaseparated string
            df_probat_ulr = df_probat_ulr.groupby(['Receptnummer', 'Receptnavn', 
                                                   'Rister','Probat id', 'Ordrenummer']).agg(
                                                       {'Silo': lambda x: ','.join(sorted(pd.Series.unique(x))),
                                                        'Dato': lambda x: ','.join(sorted(pd.Series.unique(x))),
                                                        'Kilo': 'sum'
                                                       }).reset_index()
            # Remove trailing and leading commas
            for col in columns_strip:
                df_probat_ulr[col] = df_probat_ulr[col].apply(lambda x: strip_comma_from_string(x))
            # Create temp dataframe with total
            df_temp_total = pd.concat([df_probat_ulr, pd.DataFrame.from_dict(data=dict_rister_total, orient='columns')])
            df_temp_total = df_temp_total[column_order]
            df_temp_total.sort_values(by=['Receptnummer','Ordrenummer'], inplace=True)
            # Data formating
            for col in columns_1_dec:
                df_temp_total[col] = df_temp_total[col].apply(lambda x: number_format(x, 'dec_1'))
            # Write results to Word and Excel
            insert_dataframe_into_excel (df_temp_total, section_name, False)
            add_section_to_word(df_temp_total, section_name, True, [-1,0])
            # Write status into log
            section_log_insert(section_id, 0)
        except Exception as e: # Insert error into log
            section_log_insert(section_id, 2, e)
    else: # Write into log if no data is found or section is out of scope
        section_log_insert(section_id, get_section_status_code(df_temp_total))
    
    # =============================================================================
    # Section 6: Råkaffeforbrug
    # =============================================================================
    section_id = 6
    section_name = get_section_name(section_id)
    column_order = ['Sortnummer','Sortnavn','Kontraktnummer','Modtagelse', 'Silo',
                    'Ordrenummer','Kilo']
    columns_1_dec = ['Kilo']
    
    if get_section_status_code(df_probat_lr) == 99:
        try:
            # Create total for dataframe
            dict_rister_ind_total = {'Kilo':[df_probat_lr['Kilo'].sum()],'Silo':None}
             # Look up column values
            df_probat_lr['Sortnavn'] = df_probat_lr['Sortnummer'].apply(get_nav_item_info, field='Beskrivelse')
            # Create temp dataframe with total
            df_temp_total = pd.concat([df_probat_lr, pd.DataFrame.from_dict(data=dict_rister_ind_total, orient='columns')])
            df_temp_total = df_temp_total[column_order]
            df_temp_total.sort_values(by=['Ordrenummer','Sortnummer','Kilo'], inplace=True)
            # Data formating
            for col in columns_1_dec:
                df_temp_total[col] = df_temp_total[col].apply(lambda x: number_format(x, 'dec_1'))
            # Write results to Word and Excel
            insert_dataframe_into_excel (df_temp_total, section_name, False)
            add_section_to_word(df_temp_total, section_name, True, [-1,0])
            # Write status into log
            section_log_insert(section_id, 0)
        except Exception as e: # Insert error into log
            section_log_insert(section_id, 2, e)
    else: # Write into log if no data is found or section is out of scope
        section_log_insert(section_id, get_section_status_code(df_temp_total))
    
    # =============================================================================
    # Section 7: Debitorer
    # =============================================================================
    section_id = 7
    section_name = get_section_name(section_id)
    column_order = ['Debitornummer','Debitornavn','Dato','Varenummer','Varenavn','Produktionsordrenummer',
                        'Enheder','Kilo']
    columns_1_dec = ['Enheder','Kilo']
    columns_strip = ['Produktionsordrenummer']
    
    if get_section_status_code(df_nav_debitorer) == 99:
        try:
            # Concat Order nos to one string
            df_nav_debitorer = df_nav_debitorer.groupby(['Debitornummer','Debitornavn','Dato','Varenummer']).agg(
                {'Produktionsordrenummer': lambda x: ','.join(sorted(pd.Series.unique(x))),
                 'Enheder': 'sum',
                 'Kilo': 'sum'
                }).reset_index()
            # Remove trailing and leading commas
            for col in columns_strip:
                df_nav_debitorer[col] = df_nav_debitorer[col].apply(lambda x: strip_comma_from_string(x))
            # Create total for dataframe
            dict_debitor_total = {'Enheder': [df_nav_debitorer['Enheder'].sum()],
                                  'Kilo':[df_nav_debitorer['Kilo'].sum()]}
            # Add varenavn
            df_nav_debitorer['Varenavn'] = df_nav_debitorer['Varenummer'].apply(get_nav_item_info, field='Beskrivelse')
             # Look up column values and string format datecolumn for export
            df_nav_debitorer['Dato'] = df_nav_debitorer['Dato'].dt.strftime('%d-%m-%Y')
            # Create temp dataframe with total
            df_temp_total = pd.concat([df_nav_debitorer, pd.DataFrame.from_dict(data=dict_debitor_total, orient='columns')])
            df_temp_total = df_temp_total[column_order]
            df_temp_total.sort_values(by=['Varenummer','Debitornummer','Dato'], inplace=True)
            # Data formating
            for col in columns_1_dec:
                df_temp_total[col] = df_temp_total[col].apply(lambda x: number_format(x, 'dec_1'))
            # Write results to Word and Excel
            insert_dataframe_into_excel (df_temp_total, section_name, False)
            add_section_to_word(df_temp_total, section_name, True, [-1,0])
            # Write status into log
            section_log_insert(section_id, 0)
        except Exception as e: # Insert error into log
            section_log_insert(section_id, 2, e)
    else: # Write into log if no data is found or section is out of scope
        section_log_insert(section_id, get_section_status_code(df_temp_total))
    
    # =============================================================================
    # Section 8: Massebalance
    # =============================================================================
    section_id = 8
    section_name = get_section_name(section_id)
    columns_1_dec = ['[1] Råkaffe','[2] Ristet kaffe','[3] Difference','[5] Færdigvaretilgang',
                     '[6] Difference','[8] Salg','[9] Regulering & ompak','[10] Restlager','[11] Difference']
    columns_2_pct = ['[4] Difference pct','[7] Difference pct','[12] Difference pct']
    
    dict_massebalance = {'[1] Råkaffe': df_probat_lr['Kilo'].sum(),
                         '[2] Ristet kaffe': df_probat_ulr['Kilo'].sum(),
                         '[3] Difference': None,
                         '[4] Difference pct': None,
                         '[5] Færdigvaretilgang': df_nav_færdigvaretilgang['Produceret'].sum(),
                         '[6] Difference': None,
                         '[7] Difference pct': None,
                         '[8] Salg': df_nav_færdigvaretilgang['Salg'].sum(),
                         '[9] Regulering & ompak': df_nav_færdigvaretilgang['Regulering & ompak'].sum(),
                         '[10] Restlager': df_nav_færdigvaretilgang['Restlager'].sum(),
                         '[11] Difference': None,
                         '[12] Difference pct': None}
    dict_massebalance['[3] Difference'] = dict_massebalance['[1] Råkaffe'] - dict_massebalance['[2] Ristet kaffe']
    dict_massebalance['[4] Difference pct'] = zero_division(dict_massebalance['[3] Difference'], dict_massebalance['[1] Råkaffe'], 'None')
    dict_massebalance['[6] Difference'] = dict_massebalance['[2] Ristet kaffe'] - dict_massebalance['[5] Færdigvaretilgang']
    dict_massebalance['[7] Difference pct'] = zero_division(dict_massebalance['[6] Difference'], dict_massebalance['[2] Ristet kaffe'] ,'None')
    dict_massebalance['[11] Difference'] = ( dict_massebalance['[5] Færdigvaretilgang']
        - dict_massebalance['[8] Salg'] - dict_massebalance['[9] Regulering & ompak']
        - dict_massebalance['[10] Restlager'] )
    dict_massebalance['[12] Difference pct'] = zero_division(dict_massebalance['[11] Difference'], dict_massebalance['[5] Færdigvaretilgang'], 'None')
    #Number formating
    for col in columns_1_dec:
        dict_massebalance[col] = number_format(dict_massebalance[col] ,'dec_1')
    for col in columns_2_pct:
        dict_massebalance[col] = number_format(dict_massebalance[col] ,'pct_2')
    
    df_massebalance = pd.DataFrame.from_dict(data=dict_massebalance, orient='index').reset_index()
    df_massebalance.columns = ['Sektion','Værdi']
    df_massebalance['Note'] = [None, None, '[1] - [2]', '[3] / [1]', None, '[2] - [5]',
                               '[6] / [2]', None, None, None, '[5] - [8] - [9] - [10]',
                               '[11] / [5]']
    df_massebalance['Bemærkning'] = None
    
    if get_section_status_code(df_massebalance) == 99:
        try:
            # Write results to Word and Excel
            insert_dataframe_into_excel (df_massebalance, section_name, True)
            add_section_to_word(df_massebalance, section_name, True, [0,3,4,6,7,11,12])
            # Write status into log
            section_log_insert(section_id, 0)
        except Exception as e: # Insert error into log
            section_log_insert(section_id, 2, e)
    else: # Write into log if no data is found or section is out of scope
        section_log_insert(section_id, get_section_status_code(df_massebalance))
    
    # =============================================================================
    # Section 10: Vægtkontrol
    # =============================================================================
    section_id = 10
    section_name = get_section_name(section_id)
    column_order = ['Registreringstidspunkt','Serienummer','Vægt','Status','Registreret af']
    columns_2_dec = ['Vægt']
    
    if get_section_status_code(df_ds_vægtkontrol) == 99:
        try:
            df_ds_vægtkontrol = df_ds_vægtkontrol[column_order]
            df_ds_vægtkontrol['Registreringstidspunkt'] = df_ds_vægtkontrol['Registreringstidspunkt'].dt.strftime('%d-%m-%Y %H:%M')
            #Column formating
            for col in columns_2_dec:
                df_ds_vægtkontrol[col] = df_ds_vægtkontrol[col].apply(lambda x: number_format(x, 'dec_2'))
            # Write results to Word and Excel
            insert_dataframe_into_excel (df_ds_vægtkontrol, section_name, False)
            add_section_to_word(df_ds_vægtkontrol, section_name, False, [0])
            # Write status into log
            section_log_insert(section_id, 0)
        except Exception as e: # Insert error into log
            section_log_insert(section_id, 2, e)
    else: # Write into log if no data is found or section is out of scope
        section_log_insert(section_id, get_section_status_code(df_ds_vægtkontrol))
    
    # =============================================================================
    # Section 11: Ordrestatistik fra e-vejning (poser)
    # =============================================================================
    section_id = 11
    section_name = get_section_name(section_id)
    column_order = ['Total vægt kg', 'Antal poser', 'Middelvægt g', 'Standardafvigelse g',
                    'Gns. godvægt per enhed g', 'Godvægt total g', 'Nominel vægt g', 'Taravægt g']
    columns_2_dec = ['Total vægt kg', 'Antal poser', 'Middelvægt g', 'Standardafvigelse g',
                    'Gns. godvægt per enhed g', 'Godvægt total g', 'Nominel vægt g', 'Taravægt g']
    
    if get_section_status_code(df_com_statistics) == 99:
        try:
            df_com_statistics = df_com_statistics[column_order]
            #Column formating
            for col in columns_2_dec:
                df_com_statistics[col] = df_com_statistics[col].apply(lambda x: number_format(x, 'dec_2'))
            # Transpose dataframe
            df_com_statistics = df_com_statistics.transpose()
            df_com_statistics = df_com_statistics.reset_index()
            df_com_statistics.columns = ['Sektion','Værdi']
            # Write results to Word and Excel
            insert_dataframe_into_excel (df_com_statistics, section_name, False)
            add_section_to_word(df_com_statistics, section_name, False, [0])
            # Write status into log
            section_log_insert(section_id, 0)
        except Exception as e: # Insert error into log
            section_log_insert(section_id, 2, e)
    else: # Write into log if no data is found or section is out of scope
        section_log_insert(section_id, get_section_status_code(df_com_statistics))
    
    # =============================================================================
    # Section 12: Karakterer
    # =============================================================================
    section_id = 12
    section_name = get_section_name(section_id)
    columns_1_dec = ['Syre','Krop','Aroma','Eftersmag','Robusta']
    
    if get_section_status_code(df_karakterer) == 99:
        try:
            # Column formating
            df_karakterer['Dato'] = df_karakterer['Dato'].dt.strftime('%d-%m-%Y')
            for col in columns_1_dec:
                df_karakterer[col] = df_karakterer[col].apply(lambda x: number_format(x, 'dec_1'))
            # Write results to Word and Excel
            insert_dataframe_into_excel (df_karakterer, section_name, False)
            add_section_to_word(df_karakterer, section_name, True, [0])
            # Write status into log
            section_log_insert(section_id, 0)
        except Exception as e: # Insert error into log
            section_log_insert(section_id, 2, e)
    else: # Write into log if no data is found or section is out of scope
        section_log_insert(section_id, get_section_status_code(df_karakterer))
    
    # =============================================================================
    # Section 13: Komponentforbrug
    # =============================================================================
    section_id = 13
    section_name = get_section_name(section_id)
    column_order = ['Varenummer','Varenavn','Basisenhed','Antal']
    columns_1_dec = ['Antal']
    
    if get_section_status_code(df_nav_consumption) == 99:
        try:
            df_nav_consumption = df_nav_consumption[column_order]
            # Data formating
            for col in columns_1_dec:
                df_nav_consumption[col] = df_nav_consumption[col].apply(lambda x: number_format(x, 'dec_1'))
            # Write results to Word and Excel
            insert_dataframe_into_excel (df_nav_consumption, section_name, False)
            add_section_to_word(df_nav_consumption, section_name, True, [0])
            # Write status into log
            section_log_insert(section_id, 0)
        except Exception as e: # Insert error into log
            section_log_insert(section_id, 2, e)
    else: # Write into log if no data is found or section is out of scope
        section_log_insert(section_id, get_section_status_code(df_nav_consumption))
    
    # =============================================================================
    # Section 14: Anvendt primæremballage
    # =============================================================================
    section_id = 14
    section_name = get_section_name(section_id)
    column_order = ['Varenummer','Varenavn','Lotnummer','Rullenummer','Rullelængde',
                    'Pakkedato','Købsordre']
    
    if get_section_status_code(df_nav_components) == 99:
        try:
            df_nav_components = pd.concat([df_nav_components, df_ds_ventil])
            df_nav_components['Varenavn'] = df_nav_components['Varenummer'].apply(get_nav_item_info, field='Beskrivelse')
            df_nav_components = df_nav_components[column_order]
            # Write results to Word and Excel
            insert_dataframe_into_excel (df_nav_components, section_name, False)
            add_section_to_word(df_nav_components, section_name, True, [0])
            # Write status into log
            section_log_insert(section_id, 0)
        except Exception as e: # Insert error into log
            section_log_insert(section_id, 2, e)
    else: # Write into log if no data is found or section is out of scope
        section_log_insert(section_id, get_section_status_code(df_nav_components))
    
    # =============================================================================
    # Section 15: Lotnumre
    # =============================================================================
    section_id = 15
    section_name = get_section_name(section_id)
    column_order = ['Lotnummer', 'Pallenummer', 'Produktionstidspunkt', 'Kontrolleret af',
                    'Kontrol bemærkning', 'Kontroltidspunkt', 'Kilo', 'Antal poser',
                    'Antal leakers', 'Leakers pct', 'Resultat af kontrol']
    columns_0_dec = ['Antal poser','Antal leakers']
    columns_1_dec = ['Kilo']
    columns_2_pct = ['Leakers pct']
    
    if get_section_status_code(df_nav_lotno) == 99:
        try:
            df_nav_lotno = pd.merge(df_nav_lotno, df_ds_vacslip, left_on = 'Lotnummer',
                                    right_on = 'Lotnummer', how='left', suffixes=('', '_y'))
            df_nav_lotno['Antal leakers'].fillna(value=0, inplace=True)
            df_nav_lotno['Resultat af kontrol'].fillna(value='Ej kontrolleret', inplace=True)
            df_nav_lotno['Leakers pct'] = df_nav_lotno.apply(lambda x: zero_division(x['Antal leakers'], x['Antal poser'], 'Zero'), axis=1)
            df_nav_lotno['Pallenummer'] = df_nav_lotno['Pallenummer_y'].fillna(df_nav_lotno['Pallenummer'])
            df_nav_lotno['Produktionstidspunkt'] = df_nav_lotno['Produktionstidspunkt'].dt.strftime('%d-%m-%Y %H:%M')
            df_nav_lotno = df_nav_lotno[column_order]
            # Data formating
            for col in columns_1_dec:
                df_nav_lotno[col] = df_nav_lotno[col].apply(lambda x: number_format(x, 'dec_1'))
            # Data formating
            for col in columns_0_dec:
                df_nav_lotno[col] = df_nav_lotno[col].apply(lambda x: number_format(x, 'dec_0'))
            # Data formating
            for col in columns_2_pct:
                df_nav_lotno[col] = df_nav_lotno[col].apply(lambda x: number_format(x, 'pct_2'))
            # Write results to Word and Excel
            insert_dataframe_into_excel (df_nav_lotno, section_name, False)
            add_section_to_word(df_nav_lotno, section_name, True, [0])
            # Write status into log
            section_log_insert(section_id, 0)
        except Exception as e: # Insert error into log
            section_log_insert(section_id, 2, e)
    else: # Write into log if no data is found or section is out of scope
        section_log_insert(section_id, get_section_status_code(df_nav_lotno))
    
    # =============================================================================
    # Section 16: Reference- og henstandsprøver
    # =============================================================================
    section_id = 16
    section_name = get_section_name(section_id)
    column_order = ['Id', 'Registreringstidspunkt', 'Operatør', 'Silo', 'Prøvetype',
                    'Bemærkning', 'Smagning status', 'Antal prøver']
    df_temp = df_prøver[df_prøver['Prøvetype int'] != 0]
    
    if get_section_status_code(df_temp) == 99:
        try:
            df_temp['Registreringstidspunkt'] = df_temp['Registreringstidspunkt'].dt.strftime('%d-%m-%Y %H:%M')
            df_temp = df_temp[column_order]
            # Write results to Word and Excel
            insert_dataframe_into_excel (df_temp, section_name, False)
            add_section_to_word(df_temp, section_name, False, [0])
            # Write status into log
            section_log_insert(section_id, 0)
        except Exception as e: # Insert error into log
            section_log_insert(section_id, 2, e)
    else: # Write into log if no data is found or section is out of scope
        section_log_insert(section_id, get_section_status_code(df_temp))
    
    # =============================================================================
    # Section 17: Udtagne kontrolprøver
    # =============================================================================
    section_id = 17
    section_name = get_section_name(section_id)
    column_order = ['Id','Registreringstidspunkt', 'Operatør', 'Bemærkning',
                    'Mærkning', 'Rygsvejsning', 'Tæthed', 'Ventil', 'Peelbar',
                    'Tintie', 'Vægt', 'Ilt']
    columns_2_dec = ['Vægt']
    columns_0_pct = ['Ilt']
    df_temp = df_prøver[df_prøver['Prøvetype int'] == 0]
    
    if get_section_status_code(df_temp) == 99:
        try:
            df_temp['Registreringstidspunkt'] = df_temp['Registreringstidspunkt'].dt.strftime('%d-%m-%Y %H:%M')
            df_temp = df_temp[column_order]
            # Data formating
            for col in columns_2_dec:
                df_temp[col] = df_temp[col].apply(lambda x: number_format(x, 'dec_2'))
            # Data formating
            for col in columns_0_pct:
                df_temp[col] = df_temp[col].apply(lambda x: number_format(x, 'pct_0'))
            # Write results to Word and Excel
            insert_dataframe_into_excel (df_temp, section_name, False)
            add_section_to_word(df_temp, section_name, True, [0])
            # Write status into log
            section_log_insert(section_id, 0)
        except Exception as e: # Insert error into log
            section_log_insert(section_id, 2, e)
    else: # Write into log if no data is found or section is out of scope
        section_log_insert(section_id, get_section_status_code(df_temp))
    
    # =============================================================================
    # Section 18: Sektionslog
    # =============================================================================
    section_id = 18
    df_section_log = pd.read_sql(query_ds_section_log, con_04)
    section_name = get_section_name(section_id)
    
    if get_section_status_code(df_section_log) == 99:
        try:
            df_section_log['Registreringstidspunkt'] = df_section_log['Registreringstidspunkt'].dt.strftime('%H:%M%:%S')
            df_section_log.sort_values(by=['Sektionskode'], inplace=True)
            # Write results to Word and Excel
            insert_dataframe_into_excel (df_section_log, section_name, False)
            add_section_to_word(df_section_log, section_name, False, [0])
            # Write status into log
            section_log_insert(section_id, 0)
        except Exception as e: # Insert error into log
            section_log_insert(section_id, 2, e)
    else: # Write into log if no data is found or section is out of scope
        section_log_insert(section_id, get_section_status_code(df_section_log))
    
    #Save files
    excel_writer.save()
    log_insert(script_name, f'Excel file {file_name} created')
    
    doc.save(path_file_doc)
    log_insert(script_name, f'Word document {file_name} created')



# =============================================================================
# Execute correct script type
# =============================================================================

if req_type == 0:
    rapport_færdigkaffe()
elif req_type == 1:
    Sporbarhed_råkaffe.initiate_report(req_id)
elif req_type == 2:
    pass


# Exit script
raise SystemExit(0)
