#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import urllib
from datetime import datetime
import pandas as pd
from sqlalchemy import create_engine
import pyodbc
import docx
from docx.shared import Inches
import networkx as nx
import Sporbarhed_shared_functions as ssf



def initiate_report(initiate_id):

    # =============================================================================
    # Variables for query connections
    # =============================================================================
    server_04 = 'sqlsrv04'
    db_04 = 'BKI_Datastore'
    con_04 = pyodbc.connect(f'DRIVER=SQL Server;SERVER={server_04};DATABASE={db_04};autocommit=True')
    cursor_04 = con_04.cursor()
    
    server_nav = r'SQLSRV03\NAVISION'
    db_nav = 'NAV100-DRIFT'
    con_nav = pyodbc.connect(f'DRIVER=SQL Server;SERVER={server_nav};DATABASE={db_nav};Trusted_Connection=yes')
    
    server_probat = '192.168.125.161'
    db_probat = 'BKI_IMP_EXP'
    con_probat = pyodbc.connect(f'DRIVER=SQL Server;SERVER={server_probat};DATABASE={db_probat};uid=bki_read;pwd=Probat2016')
    
    # =============================================================================
    # Read data from request
    # =============================================================================
    query_ds_request =  f""" SELECT TOP 1 [Id] ,[Forespørgselstype],[Rapport_modtager]
                        ,[Referencenummer] ,[Note_forespørgsel] ,[Modtagelse]  ,[Ordrerelationstype]
                        FROM [trc].[Sporbarhed_forespørgsel]
                        WHERE [Id] = {initiate_id} """
    df_request = pd.read_sql(query_ds_request, con_04)
    
    # Exit script if no request data is found
    if len(df_request) == 0:
        raise SystemExit(0)
    
    # =============================================================================
    # Set request variables
    # =============================================================================
    req_type = df_request.loc[0, 'Forespørgselstype']
    req_reference_no = df_request.loc[0, 'Referencenummer'].rstrip(' ')
    req_recipients = df_request.loc[0, 'Rapport_modtager']
    req_note = df_request.loc[0, 'Note_forespørgsel']
    req_id = df_request.loc[0, 'Id']
    req_modtagelse = df_request.loc[0, 'Modtagelse']
    req_ordrelationstype = df_request.loc[0, 'Ordrerelationstype']
    
    script_name = 'Sporbarhed_samlet.py'
    timestamp = datetime.now().strftime('%d-%m-%Y %H:%M:%S')
    orders_top_level = [req_reference_no]
    orders_related = []
    df_sections = ssf.get_ds_reporttype(req_id)
    
    # =============================================================================
    # Update request that it is initiated and write into log
    # =============================================================================
    cursor_04.execute(f"""UPDATE [trc].[Sporbarhed_forespørgsel]
                      SET [Forespørgsel_igangsat] = getdate()
                      WHERE [Id] = {req_id} """)
    cursor_04.commit()
    ssf.log_insert(script_name, f'Request id: {req_id} initiated')
    
    # =============================================================================
    # Variables for files generated
    # =============================================================================
    filepath = r'\\filsrv01\BKI\11. Økonomi\04 - Controlling\NMO\4. Kvalitet\Sporbarhedstest\Tests via PowerApps'
    file_name = f'Rapport_{req_reference_no}_{req_id}'
    
    doc = docx.Document()
    doc.add_heading(f'Rapport for produktionsordre {req_reference_no}',0)
    doc.sections[0].header.paragraphs[0].text = f'{script_name}'
    doc.sections[0].footer.paragraphs[0].text = f'{timestamp}'
    doc.sections[0].page_width = docx.shared.Mm(297)
    doc.sections[0].page_height = docx.shared.Mm(210)
    doc.sections[0].top_margin = docx.shared.Mm(15)
    doc.sections[0].bottom_margin = docx.shared.Mm(15)
    doc.sections[0].left_margin = docx.shared.Mm(10)
    doc.sections[0].right_margin = docx.shared.Mm(10)
    doc.sections[0].orientation = docx.enum.section.WD_ORIENT.LANDSCAPE
    
    doc_name = f'{file_name}.docx'
    path_file_doc = filepath + r'\\' + doc_name
    
    wb_name = f'{file_name}.xlsx'
    path_file_wb = filepath + r'\\' + wb_name
    excel_writer = pd.ExcelWriter(path_file_wb, engine='xlsxwriter')
    
    png_relations_name = f'{file_name}.png'
    path_png_relations = filepath + r'\\' + png_relations_name
    
    # =============================================================================
    # Read setup for section for reporttypes. NAV querys with NOLOCK to prevent deadlocks
    # =============================================================================
    query_ds_reporttypes =  f"""SELECT SRS.[Sektion], SS.[Beskrivelse] AS [Sektion navn]
                           FROM [trc].[Sporbarhed_rapport_sektion] AS SRS
    					   INNER JOIN [trc].[Sporbarhed_sektion] AS SS
    					   ON SRS.[Sektion] = SS.[Id]
                           WHERE [Forespørgselstype] = {req_type} """
    df_sections = pd.read_sql(query_ds_reporttypes, con_04)
    
    
    # Query section log for each section logged per script-run.
    # Query is only executed at the end of each class
    query_ds_section_log = f""" SELECT	SL.[Sektion] AS [Sektionskode]
                           ,S.[Beskrivelse] AS [Sektion],SS.[Beskrivelse] AS [Status]
                           ,SL.[Fejlkode_script] AS [Fejlkode script], SL.[Registreringstidspunkt]
                           FROM [trc].[Sporbarhed_sektion_log] AS SL
                           INNER JOIN [trc].[Sporbarhed_sektion] AS S
                             	ON SL.[Sektion] = S.[Id]
                           INNER JOIN [trc].[Sporbarhed_statuskode] AS SS
                                ON SL.[Statuskode] = SS.[Id]
                           WHERE SL.[Forespørgsels_id] = {req_id} """