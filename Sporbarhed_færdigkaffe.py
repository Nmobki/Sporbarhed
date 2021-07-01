#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import urllib
from datetime import datetime
import pandas as pd
from sqlalchemy import create_engine
import pyodbc
import docx
import openpyxl


# =============================================================================
# Define functions
# =============================================================================

# Get visibility for section from query
def get_section_visibility(dataframe, section):
    return dataframe['Sektion_synlig'].iloc[section]

# Get section name for section from query
def get_section_name(section):
    x = df_sections['Sektion navn'].iloc[section-1]
    if len(x) == 0 or len(x) > 31:
        return 'Sektion ' + str(section)
    else:
        return x

# Find statuscode for section log
def get_section_status_code(dataframe, visibility):
    if visibility == 0:
        return 3 # Not active for selected reporting type
    if len(dataframe) == 0:
        return 1 # Empty dataframe
    else:
        return 99 # Continue

# Write into section log
def section_log_insert(start_time, section, statuscode):
    df = pd.DataFrame(data={'Forespørgsels_id':req_id,'Sektion':section, 'Statuskode':statuscode, 'Start_tid':start_time}, index=[0])
    df.to_sql('Sporbarhed_sektion_log', con=engine_04, schema='trc', if_exists='append', index=False)

# Write dataframe into Excel sheet
def insert_dataframe_into_excel (dataframe, sheetname, include_index):
    dataframe.to_excel(excel_writer, sheet_name=sheetname, index=include_index)

# Convert list into string for SQL IN operator
def string_to_sql(list_with_values):
    if len(list_with_values) == 0:
        return ''
    else:
        return "'{}'".format("','".join(list_with_values))

# Write into dbo.log                **** ÆNDRE SCHEMA TIL dbo VED DRIFT
def log_insert(event, note):
    dict_log = {'Note': note
                ,'Event': event}
    pd.DataFrame(data=dict_log, index=[0]).to_sql('Log', con=engine_04, schema='dev', if_exists='append', index=False)

# Get info from item table
def get_nav_item_info(item_no, field):
    df_temp = df_nav_items[df_nav_items['Nummer'] == item_no]
    return df_temp[field].iloc[0]

# Convert placeholder values from dataframe to empty string for Word document
def convert_placeholders_word(string):
    if string in ['None','nan','NaT']:
        return ''
    else:
        return string

# Change Word page format between landscape and horizontal
def page_orientation_word(orientation):
    x = 7772400
    y = 10972800
    if orientation == 'Portrait':
        return [x,y]
    elif orientation == 'Landscape':
        return [y,x]
    else:
        return None

# Add dataframe to word document
def add_section_to_word(dataframe, section, pagebreak, orientation):
    # Add section header
    doc.add_heading(section, 1)
    # Add a table with an extra row for headers
    table = doc.add_table(dataframe.shape[0]+1, dataframe.shape[1])
    table.style = 'Table Grid'
    # Add headers to top row
    for i in range(dataframe.shape[-1]):
        table.cell(0,i).text = dataframe.columns[i]
        table.cell(0,i).paragraphs[0].runs[0].font.bold = True # Bold header
    # Add data from dataframe to the table
    for x in range(dataframe.shape[0]):
        for y in range(dataframe.shape[-1]):
            table.cell(x+1,y).text =  convert_placeholders_word(str(dataframe.values[x,y]))
    # Add page break
    if pagebreak == True:
        doc.add_page_break()  
        # doc.sections[-1].orientation = docx.enum.section.WD_ORIENT.LANDSCAPE # Tænker måske denne her er ligegyldig
        doc.sections[-1].page_width = page_orientation_word(orientation)[0] #10058400
        doc.sections[-1].page_height = page_orientation_word(orientation)[1] #7772400

# =============================================================================
# Variables for query connections
# =============================================================================
server_04 = 'sqlsrv04'
db_04 = 'BKI_Datastore'
con_04 = pyodbc.connect(f'DRIVER=SQL Server;SERVER={server_04};DATABASE={db_04};autocommit=True')
params_04 = urllib.parse.quote_plus(f'DRIVER=SQL Server Native Client 11.0;SERVER={server_04};DATABASE={db_04};Trusted_Connection=yes')
engine_04 = create_engine(f'mssql+pyodbc:///?odbc_connect={params_04}')
cursor_04 = con_04.cursor()

server_nav = r'SQLSRV03\NAVISION'
db_nav = 'NAV100-DRIFT'
con_nav = pyodbc.connect(f'DRIVER=SQL Server;SERVER={server_nav};DATABASE={db_nav};Trusted_Connection=yes')
params_nav = urllib.parse.quote_plus(f'DRIVER=SQL Server Native Client 11.0;SERVER={server_nav};DATABASE={db_nav};Trusted_Connection=yes')
engine_nav = create_engine(f'mssql+pyodbc:///?odbc_connect={params_nav}')

server_comscale = r'comscale-bki\sqlexpress'
db_comscale = 'ComScaleDB'
con_comscale = pyodbc.connect(f'DRIVER=SQL Server;SERVER={server_comscale};DATABASE={db_comscale}')
params_comscale = urllib.parse.quote_plus(f'DRIVER=SQL Server Native Client 11.0;SERVER={server_comscale};DATABASE={db_comscale};Trusted_Connection=yes')
engine_comscale = create_engine(f'mssql+pyodbc:///?odbc_connect={params_comscale}')

server_probat = '192.168.125.161'
db_probat = 'BKI_IMP_EXP'
con_probat = pyodbc.connect(f'DRIVER=SQL Server;SERVER={server_probat};DATABASE={db_probat};uid=bki_read;pwd=Probat2016')
params_probat = urllib.parse.quote_plus(f'DRIVER=SQL Server Native Client 11.0;SERVER={server_probat};DATABASE={db_probat};Trusted_Connection=yes')
engine_probat = create_engine(f'mssql+pyodbc:///?odbc_connect={params_probat}')

# =============================================================================
# Read data from request
# =============================================================================
query_ds_request =  """ SELECT TOP 1 [Id] ,[Forespørgselstype] ,[Rapporttype]
                    ,[Rapport_modtager] 
                    ,[Referencenummer] AS [Ordrenummer] ,[Note_forespørgsel] 
                    FROM [trc].[Sporbarhed_forespørgsel]
                    WHERE [Forespørgsel_igangsat] IS NULL
                    AND [Referencetype] = 0 AND [Forespørgselstype] = 0 """
df_request = pd.read_sql(query_ds_request, con_04)

# Exit script if no request data is found
if len(df_request) == 0:
    raise SystemExit(0)

# =============================================================================
# Set request variables
# =============================================================================
req_type = df_request.loc[0, 'Forespørgselstype']
req_report_type = df_request.loc[0, 'Rapporttype']
req_order_no = df_request.loc[0, 'Ordrenummer']
req_recipients = df_request.loc[0, 'Rapport_modtager']
req_note = df_request.loc[0, 'Note_forespørgsel']
req_id = df_request.loc[0, 'Id']

script_name = 'Sporbarhed_færdigkaffe.py'
timestamp = datetime.now().strftime('%d-%m-%Y %H:%M:%S')

# =============================================================================
# Update request that it is initiated and write into log
# =============================================================================
cursor_04.execute(f"""UPDATE [trc].[Sporbarhed_forespørgsel]
                  SET [Forespørgsel_igangsat] = getdate()
                  WHERE [Id] = {req_id}""")
cursor_04.commit()
log_insert(script_name, f'Request id: {req_id} initiated')


# =============================================================================
# Variables for files generated
# =============================================================================
filepath = r'\\filsrv01\BKI\11. Økonomi\04 - Controlling\NMO\4. Kvalitet\Sporbarhedstest\Tests' # Ændre ifbm. drift
file_name = f'Sporbarhedstest_{req_order_no}_{req_id}'

doc = docx.Document()
doc.add_heading(f'Rapport for produktionsordre {req_order_no}',0)
doc.add_paragraph('')
doc.sections[0].header.paragraphs[0].text = f'\t{script_name}\t'
doc.sections[0].footer.paragraphs[0].text = f'\t{timestamp}\t'
doc_name = f'{file_name}.docx'
path_file_doc = filepath + r'\\' + doc_name



wb = openpyxl.Workbook()
wb_name = f'{file_name}.xlsx'
path_file_wb = filepath + r'\\' + wb_name
excel_writer = pd.ExcelWriter(path_file_wb, engine='xlsxwriter')

# =============================================================================
# Read setup for section for reporttypes
# =============================================================================
query_ds_reporttypes =  f"""SELECT SRS.[Sektion] ,SRS.[Sektion_synlig] ,SS.[Beskrivelse] AS [Sektion navn]
                       FROM [trc].[Sporbarhed_rapport_sektion] AS SRS
					   INNER JOIN [trc].[Sporbarhed_sektion] AS SS
					   ON SRS.[Sektion] = SS.[Id]
                       WHERE [Rapporttype] = {req_type} 
                       AND [Forespørgselstype] = {req_report_type}"""
df_sections = pd.read_sql(query_ds_reporttypes, con_04)

# =============================================================================
# Queries for different parts of report
# =============================================================================
query_ds_generelt = f""" WITH [KP] AS ( SELECT [Ordrenummer]
                	,SUM( CASE WHEN [Prøvetype] = 0 THEN [Antal_prøver] ELSE 0 END) AS [Kontrolprøve]
                	,SUM( CASE WHEN [Prøvetype] = 1 THEN [Antal_prøver] ELSE 0 END) AS [Referenceprøve]
                	,SUM( CASE WHEN [Prøvetype] = 2 THEN [Antal_prøver] ELSE 0 END) AS [Henstandsprøve]
                    FROM [cof].[Kontrolskema_prøver]
                    GROUP BY [Ordrenummer] )
                    ,[SK] AS ( SELECT [Referencenummer] ,MAX([Status]) AS [Status]
                    FROM [cof].[Smageskema] WHERE [Referencetype] = 2
                    GROUP BY [Referencenummer] )
                    SELECT SF.[Referencenummer] AS [Ordrenummer] ,SF.[Pakketidspunkt] ,KH.[Igangsat_af] AS [Igangsat af]
                    ,KH.[Silo_opstart] AS [Opstartssilo] ,KH.[Taravægt] ,KH.[Nitrogen]
                    ,KH.[Bemærkning] AS [Bemærkning opstart] ,ISNULL(KP.[Kontrolprøve] ,0) AS [Kontrolprøver]
                    ,ISNULL(KP.[Referenceprøve] ,0) AS [Referenceprøver]
                    ,ISNULL(KP.[Henstandsprøve] ,0) AS [Henstandsprøver]
                    ,CASE WHEN SK.[Status] = 1 THEN 'Godkendt' WHEN SK.[Status] = 0 THEN 'Afvist'
                    ELSE 'Ej smagt' END AS [Smagning status], KH.[Pakkelinje]
                    FROM [trc].[Sporbarhed_forespørgsel] AS SF
                    LEFT JOIN [cof].[Kontrolskema_hoved] AS KH 
                        ON SF.[Referencenummer] = KH.[Ordrenummer]
                    LEFT JOIN [KP] 
                        ON SF.[Referencenummer] = KP.[Ordrenummer]
                    LEFT JOIN [SK] 
                        ON SF.[Referencenummer] = SK.[Referencenummer]
                    WHERE SF.[Id] = {req_id} """
df_results_generelt = pd.read_sql(query_ds_generelt, con_04)

query_ds_samples = f""" SELECT KP.[Id],KP.[Ordrenummer],KP.[Registreringstidspunkt]
            	   ,KP.[Registreret_af] AS [Operatør],KP.[Bemærkning]
                   ,KP.[Prøvetype] AS [Prøvetype int],P.[Beskrivelse] AS [Prøvetype]
                   ,CASE WHEN KP.[Kontrol_mærkning] = 1 THEN 'Ok' 
                   WHEN KP.[Kontrol_mærkning] = 0 THEN 'Ej ok' END AS [Mærkning]
                   ,CASE WHEN KP.[Kontrol_rygsvejning] = 1	THEN 'Ok'
                   WHEN KP.[Kontrol_rygsvejning] = 0 THEN 'Ej ok' END AS [Rygsvejsning]
                   ,CASE WHEN KP.[Kontrol_ventil] = 1 THEN 'Ok'
                   WHEN KP.[Kontrol_ventil] = 0 THEN 'Ej ok' END AS [Ventil]
                   ,CASE WHEN KP.[Kontrol_peelbar] = 1	THEN 'Ok'
                   WHEN KP.[Kontrol_peelbar] = 0 THEN 'Ej ok' END AS [Peelbar]
                   ,CASE WHEN KP.[Kontrol_tintie] = 1 THEN 'Ok'
                   WHEN KP.[Kontrol_tintie] = 0 THEN 'Ej ok' END AS [Tintie]
				   ,CASE WHEN KP.[Kontrol_tæthed] = 1 THEN 'Ok'
                   WHEN KP.[Kontrol_tæthed] = 0 THEN 'Ej ok' END AS [Tæthed]
                   ,KP.[Vægt_aflæst] AS [Vægt],KP.[Kontrol_ilt] AS [Ilt],KP.[Silo]
                   ,CASE WHEN SK.[Status] = 1 THEN 'Godkendt' WHEN SK.[Status] = 0
                   THEN 'Afvist' ELSE 'Ej smagt' END AS [Smagning status]
				   ,KP.[Antal_prøver] AS [Antal prøver]
                   FROM [cof].[Kontrolskema_prøver] AS KP
                   INNER JOIN [cof].[Prøvetype] AS P 
                       ON KP.[Prøvetype] = P.[Id]
                   LEFT JOIN [cof].[Smageskema] AS SK
                       ON KP.[Id] = SK.[Id_org]
                       AND SK.[Id_org_kildenummer] = 6
                   WHERE KP.[Ordrenummer] = '{req_order_no}' """
df_prøver = pd.read_sql(query_ds_samples, con_04)

query_ds_karakterer = f""" SELECT [Id] ,[Dato] ,[Bruger] ,[Smag_Syre] AS [Syre]
                      ,[Smag_Krop] AS [Krop] ,[Smag_Aroma] AS [Aroma] 
                      ,[Smag_Eftersmag] AS [Eftersmag],[Smag_Robusta] AS [Robusta] ,[Bemærkning]
                      FROM [cof].[Smageskema]
                      WHERE [Referencetype] = 2	
                          AND [Referencenummer] = '{req_order_no}'
                          AND COALESCE([Smag_Syre],[Smag_Krop],[Smag_Aroma],
                            [Smag_Eftersmag],[Smag_Robusta]) IS NOT NULL"""
df_karakterer = pd.read_sql(query_ds_karakterer, con_04)

query_ds_vacslip = """ SELECT [Registreringstidspunkt] AS [Kontroltidspunkt]
                   ,[Initialer] AS [Kontrolleret af],[Lotnummer]
                   ,[Pallenummer],[Antal_poser] AS [Antal leakers]
                   ,[Bemærkning] AS [Kontrol bemærkning]
				   ,CASE WHEN [Overført_email_log] = 1 THEN
				   'Over grænseværdi' ELSE 'Ok' END AS [Resultat af kontrol]
                   FROM [cof].[Vac_slip] """
df_ds_vacslip = pd.read_sql(query_ds_vacslip, con_04)

query_ds_ventil = f""" SELECT [Varenummer] ,[Batchnr_stregkode] AS [Lotnummer]
                  FROM [cof].[Ventil_registrering]
                  WHERE [Ordrenummer] = '{req_order_no}' """
df_ds_ventil = pd.read_sql(query_ds_ventil, con_04)

query_ds_section_log = f""" SELECT	SL.[Sektion] AS [Sektionskode]
                       ,S.[Beskrivelse] AS [Sektion],SS.[Beskrivelse] AS [Status]
                       ,SL.[Start_tid],SL.[Registreringstidspunkt] AS [Slut tid]
                	   ,DATEDIFF(ms, SL.[Start_tid] ,SL.[Registreringstidspunkt]) / 1000.0 AS [Sekunder]
                       FROM [trc].[Sporbarhed_sektion_log] AS SL
                       INNER JOIN [trc].[Sporbarhed_sektion] AS S
                         	ON SL.[Sektion] = S.[Id]
                       INNER JOIN [trc].[Sporbarhed_statuskode] AS SS
                            ON SL.[Statuskode] = SS.[Id]
                       WHERE SL.[Forespørgsels_id] = {req_id} """

query_com_statistics = f""" WITH CTE AS ( SELECT SD.[Nominal] ,SD.[Tare]
                       ,SUM( SD.[MeanValueTrade] * SD.[CounterGoodTrade] ) AS [Total vægt]
                       ,SUM( SD.[StandardDeviationTrade] * SD.[CounterGoodTrade] ) AS [Std afv]
                       ,SUM( SD.[CounterGoodTrade] ) AS [Antal enheder]
                       FROM [ComScaleDB].[dbo].[StatisticData] AS SD
                       INNER JOIN [dbo].[Statistic] AS S ON SD.[Statistic_ID] = S.[ID]
                       WHERE S.[Order] = '{req_order_no}' AND lower(S.[ArticleNumber]) NOT LIKE '%k'
                       GROUP BY S.[Order],SD.[Nominal],SD.[Tare] )
                       SELECT CTE.[Total vægt],CTE.[Antal enheder]
                       ,CASE WHEN CTE.[Antal enheder] = 0 
                       THEN NULL ELSE CTE.[Total vægt] / CTE.[Antal enheder] END AS [Middelvægt]
                       ,CASE WHEN CTE.[Antal enheder] = 0 
                       THEN NULL ELSE CTE.[Std afv] / CTE.[Antal enheder] END AS [Standardafvigelse]
                       ,CASE WHEN CTE.[Antal enheder] = 0 
                       THEN NULL ELSE CTE.[Total vægt] / CTE.[Antal enheder] END - CTE.[Nominal] AS [Gns. godvægt per enhed]
                       ,CTE.[Total vægt] - CTE.[Nominal] * CTE.[Antal enheder] AS [Godvægt total]
                       ,CTE.[Nominal] AS [Nominel vægt],CTE.[Tare] AS [Taravægt]
                       FROM CTE """
df_com_statistics = pd.read_sql(query_com_statistics, con_comscale)

query_nav_items = """ SELECT [No_] AS [Nummer],[Description] AS [Beskrivelse]
                  FROM [dbo].[BKI foods a_s$Item] """
df_nav_items = pd.read_sql(query_nav_items, con_nav)

query_nav_generelt = f""" WITH [RECEPT] AS (
                     SELECT	POC.[Prod_ Order No_],I.[No_]
                     FROM [dbo].[BKI foods a_s$Prod_ Order Component] AS POC
                     INNER JOIN [dbo].[BKI foods a_s$Item] AS I
                    	ON POC.[Item No_] = I.[No_]
                     WHERE [POC].[Prod_ Order Line No_] = 10000
                    	AND I.[Sequence Code] = 4)
                     ,[ILE] AS ( SELECT [Order No_],MIN([Posting Date]) AS [Posting Date]
                     ,SUM(CASE WHEN [Entry Type] = 5 AND [Location Code] = 'REWORK' 
                          THEN [Quantity] ELSE 0 END) AS [Rework forbrug]
                     ,SUM(CASE WHEN [Entry Type] = 6 AND [Location Code] = 'REWORK' 
                          THEN [Quantity] ELSE 0 END) AS [Rework afgang]
                     ,SUM(CASE WHEN [Entry Type] = 5 AND [Location Code] = 'SLAT' 
                          THEN [Quantity] ELSE 0 END) AS [Slat forbrug]
                     ,SUM(CASE WHEN [Entry Type] = 6 AND [Location Code] = 'SLAT' 
                          THEN [Quantity] ELSE 0 END) AS [Slat afgang]
                     FROM [dbo].[BKI foods a_s$Item Ledger Entry]
                     WHERE [Order Type] = 1 GROUP BY [Order No_] )
                     SELECT PO.[Source No_] AS [Varenummer]
                     ,I.[Description] AS [Varenavn]
                     ,I.[Base Unit of Measure] AS [Basisenhed]
                     ,CASE WHEN PO.[Status] = 0 THEN 'Simuleret'
                     WHEN PO.[Status] = 1 THEN 'Planlagt'
                     WHEN PO.[Status] = 2 THEN 'Fastlagt'
                     WHEN PO.[Status] = 3 THEN 'Frigivet'
                     WHEN PO.[Status] = 4 THEN 'Færdig'
                     END AS [Prod.ordre status]
                     ,ICR.[Cross-Reference No_] AS [Stregkode]
                     ,RECEPT.[No_] AS [Receptnummer],ILE.[Rework afgang]
                     ,ILE.[Posting Date] AS [Produktionsdato]
                     ,ILE.[Rework forbrug],ILE.[Slat afgang],ILE.[Slat forbrug]
                     FROM [dbo].[BKI foods a_s$Production Order] AS PO
                     INNER JOIN [dbo].[BKI foods a_s$Item] AS I
                    	ON PO.[Source No_] = I.[No_]
                     LEFT JOIN [dbo].[BKI foods a_s$Item Cross Reference] AS ICR
                    	ON I.[No_] = ICR.[Item No_] AND ICR.[Unit of Measure] = 'PS'
                    	AND ICR.[Cross-Reference Type] = 3
                     LEFT JOIN [RECEPT] ON PO.[No_] = RECEPT.[Prod_ Order No_]
                     LEFT JOIN [ILE] ON PO.[No_] = ILE.[Order No_]
                     WHERE I.[Item Category Code] = 'FÆR KAFFE' AND PO.[No_] = '{req_order_no}' """
df_nav_generelt = pd.read_sql(query_nav_generelt, con_nav)


req_orders_total = string_to_sql([req_order_no,'036720']) # **** SKAL ÆNDRES NÅR NAV UDVIKLING ER PÅ PLADS

# Recursive query to find all relevant produced orders related to the requested order
# First is identified all lotnumbers related to the orders identified through NAV reservations (only production orders)
# Next is a recursive part which identifies any document numbers which have consumed these lotnumbers (ILE_C)
# Which is then queried again to find all lotnumbers produced on the orders from which these lotnumbers originally came.
query_nav_færdigvaretilgang = f""" WITH [LOT_ORG] AS ( SELECT [Lot No_]
                              FROM [dbo].[BKI foods a_s$Item Ledger Entry]
                              WHERE [Order No_] IN({req_orders_total})
                              AND [Entry Type] = 6
                              UNION ALL
                              SELECT ILE_O.[Lot No_]
                              FROM [LOT_ORG]
                              INNER JOIN [dbo].[BKI foods a_s$Item Ledger Entry] AS ILE_C
                                  ON [LOT_ORG].[Lot No_] = ILE_C.[Lot No_]
                                  AND [ILE_C].[Entry Type] IN (5,8)
                              INNER JOIN [dbo].[BKI foods a_s$Item Ledger Entry]  AS ILE_O
                            	  ON ILE_C.[Document No_] = ILE_O.[Document No_]
                                  AND ILE_O.[Entry Type] IN (6,9) )
                              ,[LOT_SINGLE] AS ( SELECT [Lot No_]
                              FROM [LOT_ORG] GROUP BY [Lot No_] )
                              SELECT ILE.[Item No_] AS [Varenummer],I.[Description] AS [Varenavn]
                        	  ,SUM(CASE WHEN ILE.[Entry Type] IN (0,6,9)
                        		THEN ILE.[Quantity] * I.[Net Weight]
                        		ELSE 0 END) AS [Produceret]
                        	,SUM(CASE WHEN ILE.[Entry Type] = 1
                        		THEN ILE.[Quantity] * I.[Net Weight] * -1
                        		ELSE 0 END) AS [Salg]
                        	,SUM(CASE WHEN ILE.[Entry Type] NOT IN (0,1,6,9)
                        		THEN ILE.[Quantity] * I.[Net Weight] * -1
                        		ELSE 0 END) AS [Regulering & ompak]
                        	,SUM(ILE.[Remaining Quantity] * I.[Net Weight]) AS [Restlager]
                            FROM [dbo].[BKI foods a_s$Item Ledger Entry] AS ILE
                            INNER JOIN [dbo].[BKI foods a_s$Item] AS I
                            	ON ILE.[Item No_] = I.[No_]
                            INNER JOIN [LOT_SINGLE]
                            	ON ILE.[Lot No_] = [LOT_SINGLE].[Lot No_]
                            GROUP BY ILE.[Item No_],I.[Description] """
df_nav_færdigvaretilgang = pd.read_sql(query_nav_færdigvaretilgang, con_nav)

# Recursive query to get all customer who purchased identified lotnumbers.
# See explanation of query above
query_nav_debitorer = f""" WITH [LOT_ORG] AS ( SELECT [Lot No_]
                      FROM [dbo].[BKI foods a_s$Item Ledger Entry]
                      WHERE [Order No_] IN({req_orders_total}) AND [Entry Type] = 6
                      UNION ALL
                      SELECT ILE_O.[Lot No_]
                      FROM [LOT_ORG]
                      INNER JOIN [dbo].[BKI foods a_s$Item Ledger Entry] AS ILE_C
                          ON [LOT_ORG].[Lot No_] = ILE_C.[Lot No_]
                    	  AND [ILE_C].[Entry Type] IN (5,8)
                      INNER JOIN [dbo].[BKI foods a_s$Item Ledger Entry]  AS ILE_O
                    	  ON ILE_C.[Document No_] = ILE_O.[Document No_]
                    	  AND ILE_O.[Entry Type] IN (6,9) )
                      ,[LOT_SINGLE] AS ( SELECT [Lot No_]
                      FROM [LOT_ORG] GROUP BY [Lot No_] )
                      SELECT C.[No_] AS [Debitornummer],C.[Name] AS [Debitornavn]
                    	  ,ILE.[Posting Date] AS [Dato]
                    	  ,ILE.[Item No_] AS [Varenummer]
                    	  ,SUM(ILE.[Quantity] * -1) AS [Enheder]
                    	  ,SUM(ILE.[Quantity] * I.[Net Weight] * -1) AS [Kilo]
                      FROM [dbo].[BKI foods a_s$Item Ledger Entry] AS ILE
                      INNER JOIN [dbo].[BKI foods a_s$Item] AS I
                    	  ON ILE.[Item No_] = I.[No_]
                      INNER JOIN [LOT_SINGLE]
                      	  ON ILE.[Lot No_] = [LOT_SINGLE].[Lot No_]
                      INNER JOIN [dbo].[BKI foods a_s$Customer] AS C
                    	  ON ILE.[Source No_] = C.[No_]
                      WHERE ILE.[Entry Type] = 1
                      GROUP BY  C.[No_] ,C.[Name],ILE.[Posting Date],ILE.[Item No_] """
df_nav_debitorer = pd.read_sql(query_nav_debitorer, con_nav)

query_nav_lotno = f""" SELECT ILE.[Lot No_] AS [Lotnummer]
            	  ,LI.[Certificate Number] AS [Pallenummer]
              	  ,[Quantity] * I.[Net Weight] AS [Kilo]
            	  ,CAST(ROUND(ILE.[Quantity] / IUM.[Qty_ per Unit of Measure],0) AS INT) AS [Antal poser]
            	  ,DATEADD(hour, 1, ILE.[Produktionsdato_-tid]) AS [Produktionstidspunkt]
                  FROM [dbo].[BKI foods a_s$Item Ledger Entry] ILE
                  INNER JOIN [dbo].[BKI foods a_s$Item] AS I
                      ON ILE.[Item No_] = I.[No_]
                  LEFT JOIN [dbo].[BKI foods a_s$Lot No_ Information] AS LI
                	  ON ILE.[Lot No_] = LI.[Lot No_]
                      AND ILE.[Item No_] = LI.[Item No_]
                  LEFT JOIN [dbo].[BKI foods a_s$Item Unit of Measure] AS IUM
                	  ON ILE.[Item No_] = IUM.[Item No_]
                      AND IUM.[Code] = 'PS'
                  WHERE ILE.[Order Type] = 1
                	  AND ILE.[Entry Type] = 6
                      AND ILE.[Order No_] = '{req_order_no}' """
df_nav_lotno = pd.read_sql(query_nav_lotno, con_nav)

query_nav_components = f""" SELECT POC.[Item No_] AS [Varenummer]
                	   ,I.[Description] AS [Varenavn]
                       ,POAC.[Purchase Order No_] AS [Købsordre]
                       ,POAC.[Roll No_] AS [Rullenummer]
                       ,CAST(POAC.[Roll Lenght] AS INT) AS [Rullelængde]
                       ,POAC.[Batch_Lot No_] AS [Lotnummer]
                       ,POAC.[Packaging Date] AS [Pakkedato]
                       FROM [dbo].[BKI foods a_s$Prod_ Order Add_ Comp_] AS POAC
                       INNER JOIN [dbo].[BKI foods a_s$Prod_ Order Component] AS POC
                           ON POAC.[Prod_ Order No_] = POC.[Prod_ Order No_]
                           AND POAC.[Prod_ Order Line No_] = POC.[Prod_ Order Line No_]
                           AND POAC.[Prod_ Order Component Line No_] = POC.[Line No_]
                       INNER JOIN [dbo].[BKI foods a_s$Item] AS I
                           ON POC.[Item No_] = I.[No_]
                       WHERE POAC.[Prod_ Order No_] = '{req_order_no}' """
df_nav_components = pd.read_sql(query_nav_components, con_nav)

query_nav_consumption = f""" SELECT	ILE.[Item No_] AS [Varenummer]
                    	,I.[Description] AS [Varenavn]
                        ,I.[Base Unit of Measure] AS [Basisenhed]
                        ,SUM(ILE.[Quantity]) * -1 AS [Antal]
                        FROM [dbo].[BKI foods a_s$Item Ledger Entry] AS ILE
                        INNER JOIN [dbo].[BKI foods a_s$Item] AS I
                        	ON ILE.[Item No_] = I.[No_]
                        WHERE ILE.[Order No_] = '{req_order_no}'
                        	AND ILE.[Entry Type] = 5
                        GROUP BY ILE.[Item No_] ,I.[Description],I.[Base Unit of Measure] """
df_nav_consumption = pd.read_sql(query_nav_consumption, con_nav)


# OBS!!! Denne liste skal dannes ud fra NAV forespørgsel når Jira er på plads!!!!
related_orders = string_to_sql(['041367','041344','041234'])

query_probat_ulg = f""" SELECT DATEADD(D, DATEDIFF(D, 0, [RECORDING_DATE] ), 0) AS [Dato]
                   ,[PRODUCTION_ORDER_ID] AS [Probat id] ,[SOURCE_NAME] AS [Mølle]
                   ,[ORDER_NAME] AS [Ordrenummer] ,[D_CUSTOMER_CODE] AS [Receptnummer]
                   ,[DEST_NAME] AS [Silo],SUM([WEIGHT]) / 1000.0 AS [Kilo]
                   FROM [dbo].[PRO_EXP_ORDER_UNLOAD_G]
                   WHERE [ORDER_NAME] IN ({related_orders})
                   GROUP BY [PRODUCTION_ORDER_ID],[ORDER_NAME],[DEST_NAME],[SOURCE_NAME]
                   ,[D_CUSTOMER_CODE], DATEADD(D, DATEDIFF(D, 0, [RECORDING_DATE] ), 0) """
df_probat_ulg = pd.read_sql(query_probat_ulg, con_probat)


query_probat_lg = f""" SELECT [S_ORDER_NAME]
                       FROM [dbo].[PRO_EXP_ORDER_LOAD_G]
                       WHERE [ORDER_NAME] IN ({related_orders})
                       GROUP BY	[S_ORDER_NAME] """
if len(df_probat_ulg) != 0: # Add to list only if dataframe is not empty
    df_probat_lg = pd.read_sql(query_probat_lg, con_probat)
    related_orders = related_orders + ',' + string_to_sql(df_probat_lg['S_ORDER_NAME'].unique().tolist())


query_probat_ulr = f""" SELECT [S_CUSTOMER_CODE] AS [Receptnummer]
                        ,DATEADD(D, DATEDIFF(D, 0, [RECORDING_DATE] ), 0) AS [Dato]
                        ,[SOURCE_NAME] AS [Rister] ,[PRODUCTION_ORDER_ID] AS [Probat id]
                    	,[ORDER_NAME] AS [Ordrenummer] ,SUM([WEIGHT]) / 1000.0 AS [Kilo]
						,[DEST_NAME] AS [Silo]
                        FROM [dbo].[PRO_EXP_ORDER_UNLOAD_R]
                        WHERE [ORDER_NAME] IN ({related_orders})
                        GROUP BY [S_CUSTOMER_CODE],[SOURCE_NAME],[PRODUCTION_ORDER_ID]
                        ,[ORDER_NAME],DATEADD(D, DATEDIFF(D, 0, [RECORDING_DATE] ), 0)
						,[DEST_NAME] """
df_probat_ulr = pd.read_sql(query_probat_ulr, con_probat)


query_probat_lr = f""" SELECT [S_TYPE_CELL] AS [Sortnummer] ,[Source] AS [Silo]
                ,[S_CONTRACT_NO] AS [Kontraktnummer]
                ,[S_DELIVERY_NAME] AS [Modtagelse],[ORDER_NAME] AS [Ordrenummer]
            	,SUM([WEIGHT]) / 1000.0 AS [Kilo]
                FROM [dbo].[PRO_EXP_ORDER_LOAD_R]
                WHERE [ORDER_NAME] IN ({related_orders})
                GROUP BY [S_TYPE_CELL],[Source],[S_CONTRACT_NO]
                	,[S_DELIVERY_NAME],[ORDER_NAME] """
df_probat_lr = pd.read_sql(query_probat_lr, con_probat)



# =============================================================================
# Section 1: Generelt
# =============================================================================
section_id = 1
section_name = get_section_name(section_id)
timestamp = datetime.now()
column_order = ['Varenummer', 'Varenavn', 'Basisenhed', 'Receptnummer', 'Pakkelinje',
                'Produktionsdato', 'Pakketidspunkt', 'Stregkode', 'Ordrenummer',
                'Smagning status', 'Opstartssilo', 'Igangsat af', 'Taravægt',
                'Nitrogen', 'Henstandsprøver', 'Referenceprøver', 'Kontrolprøver',
                'Bemærkning opstart', 'Lotnumre produceret', 'Slat forbrug',
                'Slat afgang', 'Rework forbrug', 'Rework afgang' ,'Prod.ordre status']
if get_section_status_code(df_results_generelt, get_section_visibility(df_sections, section_id)) == 99:
    try:
        df_results_generelt['Varenummer'] = df_nav_generelt['Varenummer'].iloc[0]
        df_results_generelt['Varenavn'] = df_nav_generelt['Varenavn'].iloc[0]
        df_results_generelt['Basisenhed'] = df_nav_generelt['Basisenhed'].iloc[0]
        df_results_generelt['Receptnummer'] = df_nav_generelt['Receptnummer'].iloc[0]
        df_results_generelt['Produktionsdato'] = df_nav_generelt['Produktionsdato'].iloc[0]
        df_results_generelt['Produktionsdato'].dt.strftime('%d-%m-%Y')
        df_results_generelt['Stregkode'] = df_nav_generelt['Stregkode'].iloc[0]
        df_results_generelt['Lotnumre produceret'] = len(df_nav_lotno)
        df_results_generelt['Slat forbrug'] = df_nav_generelt['Slat forbrug'].iloc[0]
        df_results_generelt['Slat afgang'] = df_nav_generelt['Slat afgang'].iloc[0]
        df_results_generelt['Rework forbrug'] = df_nav_generelt['Rework forbrug'].iloc[0]
        df_results_generelt['Rework afgang'] = df_nav_generelt['Rework afgang'].iloc[0]
        df_results_generelt['Prod.ordre status'] = df_nav_generelt['Prod.ordre status'].iloc[0]
        df_results_generelt = df_results_generelt[column_order].transpose()        
        # Write results to Word and Excel
        insert_dataframe_into_excel (df_results_generelt, section_name, True)
        add_section_to_word(df_results_generelt, section_name, True, 'Portrait')
        # Write status into log
        section_log_insert(timestamp, section_id, 0)
    except: # Insert error into log
        section_log_insert(timestamp, section_id, 2)
else: # Write into log if no data is found or section is out of scope
    section_log_insert(timestamp, section_id, get_section_status_code(df_results_generelt, get_section_visibility(df_sections, section_id)))


# =============================================================================
# Section 3: Færdigvaretilgang
# =============================================================================
section_id = 3
section_name = get_section_name(section_id)
timestamp = datetime.now()
column_order = ['Varenummer','Varenavn','Produceret','Salg','Restlager','Regulering & ompak']

if get_section_status_code(df_nav_færdigvaretilgang, get_section_visibility(df_sections, section_id)) == 99:
    try:
        df_nav_færdigvaretilgang = df_nav_færdigvaretilgang[column_order]
        # Write results to Word and Excel
        insert_dataframe_into_excel (df_nav_færdigvaretilgang, section_name, False)
        add_section_to_word(df_nav_færdigvaretilgang, section_name, True, 'Landscape')
        # Write status into log
        section_log_insert(timestamp, section_id, 0)
    except: # Insert error into log
        section_log_insert(timestamp, section_id, 2)
else: # Write into log if no data is found or section is out of scope
    section_log_insert(timestamp, section_id, get_section_status_code(df_nav_færdigvaretilgang, get_section_visibility(df_sections, section_id)))


# =============================================================================
# Section 4: Mølleordrer
# =============================================================================
section_id = 4
section_name = get_section_name(section_id)
timestamp = datetime.now()
column_order = ['Receptnummer', 'Receptnavn', 'Dato', 'Mølle',
                'Probat id', 'Ordrenummer', 'Kilo']

if get_section_status_code(df_probat_ulg, get_section_visibility(df_sections, section_id)) == 99:
    try:
        df_probat_ulg['Receptnavn'] = df_probat_ulg['Receptnummer'].apply(get_nav_item_info, field='Beskrivelse')
        df_probat_ulg['Dato'] = df_probat_ulg['Dato'].dt.strftime('%d-%m-%Y')
        df_probat_ulg = df_probat_ulg[column_order]
        # Write results to Word and Excel
        insert_dataframe_into_excel (df_probat_ulg, section_name, False)
        add_section_to_word(df_probat_ulg, section_name, True, 'Landscape')
        # Write status into log
        section_log_insert(timestamp, section_id, 0)
    except: # Insert error into log
        section_log_insert(timestamp, section_id, 2)
else: # Write into log if no data is found or section is out of scope
    section_log_insert(timestamp, section_id, get_section_status_code(df_probat_ulg, get_section_visibility(df_sections, section_id)))


# =============================================================================
# Section 5: Risteordrer
# =============================================================================
section_id = 5
section_name = get_section_name(section_id)
timestamp = datetime.now()
column_order = ['Receptnummer', 'Receptnavn', 'Dato', 'Rister',
                'Probat id', 'Ordrenummer', 'Kilo']

if get_section_status_code(df_probat_ulr, get_section_visibility(df_sections, section_id)) == 99:
    try:
        df_probat_ulr['Receptnavn'] = df_probat_ulr['Receptnummer'].apply(get_nav_item_info, field='Beskrivelse')
        df_probat_ulr['Dato'] = df_probat_ulr['Dato'].dt.strftime('%d-%m-%Y')
        df_probat_ulr = df_probat_ulr[column_order]
        # Write results to Word and Excel
        insert_dataframe_into_excel (df_probat_ulr, section_name, False)
        add_section_to_word(df_probat_ulr, section_name, True, 'Landscape')
        # Write status into log
        section_log_insert(timestamp, section_id, 0)
    except: # Insert error into log
        section_log_insert(timestamp, section_id, 2)
else: # Write into log if no data is found or section is out of scope
    section_log_insert(timestamp, section_id, get_section_status_code(df_probat_ulr, get_section_visibility(df_sections, section_id)))


# =============================================================================
# Section 6: Råkaffeforbrug
# =============================================================================
section_id = 6
section_name = get_section_name(section_id)
timestamp = datetime.now()
column_order = ['Sortnummer','Sortnavn','Silo','Kontraktnummer','Modtagelse',
                'Ordrenummer','Kilo']
if get_section_status_code(df_probat_lr, get_section_visibility(df_sections, section_id)) == 99:
    try:
        df_probat_lr['Sortnavn'] = df_probat_lr['Sortnummer'].apply(get_nav_item_info, field='Beskrivelse')
        df_probat_lr = df_probat_lr[column_order]
        # Write results to Word and Excel
        insert_dataframe_into_excel (df_probat_lr, section_name, False)
        add_section_to_word(df_probat_lr, section_name, True, 'Landscape')
        # Write status into log
        section_log_insert(timestamp, section_id, 0)
    except: # Insert error into log
        section_log_insert(timestamp, section_id, 2)
else: # Write into log if no data is found or section is out of scope
    section_log_insert(timestamp, section_id, get_section_status_code(df_probat_lr, get_section_visibility(df_sections, section_id)))


# =============================================================================
# Section 7: Debitorer
# =============================================================================
section_id = 7
section_name = get_section_name(section_id)
timestamp = datetime.now()
column_order = ['Debitornummer','Debitornavn','Dato','Varenummer','Enheder','Kilo']

if get_section_status_code(df_nav_debitorer, get_section_visibility(df_sections, section_id)) == 99:
    try:
        df_nav_debitorer['Dato'] = df_nav_debitorer['Dato'].dt.strftime('%d-%m-%Y')
        df_nav_debitorer = df_nav_debitorer[column_order]
        # Write results to Word and Excel
        insert_dataframe_into_excel (df_nav_debitorer, section_name, False)
        add_section_to_word(df_nav_debitorer, section_name, True, 'Landscape')
        # Write status into log
        section_log_insert(timestamp, section_id, 0)
    except: # Insert error into log
        section_log_insert(timestamp, section_id, 2)
else: # Write into log if no data is found or section is out of scope
    section_log_insert(timestamp, section_id, get_section_status_code(df_nav_debitorer, get_section_visibility(df_sections, section_id)))


# =============================================================================
# Section 8: Massebalance
# =============================================================================
section_id = 8
section_name = get_section_name(section_id)
timestamp = datetime.now()
dict_massebalance = {'[1] Råkaffe': df_probat_lr['Kilo'].sum(),
                     '[2] Ristet kaffe': df_probat_ulr['Kilo'].sum(),
                     '[3] Difference': None,
                     '[4] Difference pct': None,
                     '[5] Færdigvaretilgang': df_nav_færdigvaretilgang['Produceret'].sum(),
                     '[6] Difference': None,
                     '[7] Difference pct': None,
                     '[8] Salg': df_nav_færdigvaretilgang['Salg'].sum(),
                     '[9] Regulering & ompak': df_nav_færdigvaretilgang['Regulering & ompak'].sum(),
                     '[10] Restlager': df_nav_færdigvaretilgang['Restlager'].sum(),
                     '[11] Difference': None,
                     '[12] Difference pct': None}
dict_massebalance['[3] Difference'] = dict_massebalance['[1] Råkaffe'] - dict_massebalance['[2] Ristet kaffe']
dict_massebalance['[4] Difference pct'] = dict_massebalance['[3] Difference'] / dict_massebalance['[1] Råkaffe']
dict_massebalance['[6] Difference'] = dict_massebalance['[2] Ristet kaffe'] - dict_massebalance['[5] Færdigvaretilgang']
dict_massebalance['[7] Difference pct'] = dict_massebalance['[6] Difference'] / dict_massebalance['[2] Ristet kaffe']
dict_massebalance['[11] Difference'] = ( dict_massebalance['[5] Færdigvaretilgang']
    - dict_massebalance['[8] Salg'] - dict_massebalance['[9] Regulering & ompak']
    - dict_massebalance['[10] Restlager'] )
dict_massebalance['[12] Difference pct'] = dict_massebalance['[11] Difference'] / dict_massebalance['[5] Færdigvaretilgang']

df_massebalance = pd.DataFrame.from_dict(data=dict_massebalance, orient='index')

if get_section_status_code(df_massebalance, get_section_visibility(df_sections, section_id)) == 99:
    try:
        # Write results to Word and Excel
        insert_dataframe_into_excel (df_massebalance, section_name, True)
        add_section_to_word(df_massebalance, section_name, True, 'Landscape')
        # Write status into log
        section_log_insert(timestamp, section_id, 0)
    except: # Insert error into log
        section_log_insert(timestamp, section_id, 2)
else: # Write into log if no data is found or section is out of scope
    section_log_insert(timestamp, section_id, get_section_status_code(df_massebalance, get_section_visibility(df_sections, section_id)))


# =============================================================================
# Section 11: Ordrestatistik fra e-vejning
# =============================================================================
section_id = 11
section_name = get_section_name(section_id)
timestamp = datetime.now()
column_order = ['Total vægt', 'Antal enheder', 'Middelvægt', 'Standardafvigelse',
                'Gns. godvægt per enhed', 'Godvægt total', 'Nominel vægt', 'Taravægt']

if get_section_status_code(df_com_statistics, get_section_visibility(df_sections, section_id)) == 99:
    try:
        df_com_statistics = df_com_statistics[column_order]
        # Write results to Word and Excel
        insert_dataframe_into_excel (df_com_statistics, section_name, False)
        add_section_to_word(df_com_statistics, section_name, True, 'Landscape')
        # Write status into log
        section_log_insert(timestamp, section_id, 0)
    except: # Insert error into log
        section_log_insert(timestamp, section_id, 2)
else: # Write into log if no data is found or section is out of scope
    section_log_insert(timestamp, section_id, get_section_status_code(df_com_statistics, get_section_visibility(df_sections, section_id)))


# =============================================================================
# Section 12: Karakterer
# =============================================================================
section_id = 12
section_name = get_section_name(section_id)
timestamp = datetime.now()

if get_section_status_code(df_karakterer, get_section_visibility(df_sections, section_id)) == 99:
    try:
        df_karakterer['Dato'] = df_karakterer['Dato'].dt.strftime('%d-%m-%Y')
        # Write results to Word and Excel
        insert_dataframe_into_excel (df_karakterer, section_name, False)
        add_section_to_word(df_karakterer, section_name, True, 'Landscape')
        # Write status into log
        section_log_insert(timestamp, section_id, 0)
    except: # Insert error into log
        section_log_insert(timestamp, section_id, 2)
else: # Write into log if no data is found or section is out of scope
    section_log_insert(timestamp, section_id, get_section_status_code(df_karakterer, get_section_visibility(df_sections, section_id)))


# =============================================================================
# Section 13: Komponentforbrug
# =============================================================================
section_id = 13
section_name = get_section_name(section_id)
timestamp = datetime.now()
column_order = ['Varenummer','Varenavn','Basisenhed','Antal']

if get_section_status_code(df_nav_consumption, get_section_visibility(df_sections, section_id)) == 99:
    try:
        df_nav_consumption = df_nav_consumption[column_order]
        # Write results to Word and Excel
        insert_dataframe_into_excel (df_nav_consumption, section_name, False)
        add_section_to_word(df_nav_consumption, section_name, True, 'Landscape')
        # Write status into log
        section_log_insert(timestamp, section_id, 0)
    except: # Insert error into log
        section_log_insert(timestamp, section_id, 2)
else: # Write into log if no data is found or section is out of scope
    section_log_insert(timestamp, section_id, get_section_status_code(df_nav_consumption, get_section_visibility(df_sections, section_id)))


# =============================================================================
# Section 14: Anvendt primæremballage
# =============================================================================
section_id = 14
section_name = get_section_name(section_id)
timestamp = datetime.now()
column_order = ['Varenummer','Varenavn','Lotnummer','Rullenummer','Rullelængde',
                'Pakkedato','Købsordre']

if get_section_status_code(df_nav_components, get_section_visibility(df_sections, section_id)) == 99:
    try:
        df_nav_components = pd.concat([df_nav_components, df_ds_ventil])
        df_nav_components['Varenavn'] = df_nav_components['Varenummer'].apply(get_nav_item_info, field='Beskrivelse')
        df_nav_components = df_nav_components[column_order]
        # Write results to Word and Excel
        insert_dataframe_into_excel (df_nav_components, section_name, False)
        add_section_to_word(df_nav_components, section_name, True, 'Landscape')
        # Write status into log
        section_log_insert(timestamp, section_id, 0)
    except: # Insert error into log
        section_log_insert(timestamp, section_id, 2)
else: # Write into log if no data is found or section is out of scope
    section_log_insert(timestamp, section_id, get_section_status_code(df_nav_components, get_section_visibility(df_sections, section_id)))


# =============================================================================
# Section 15: Lotnumre
# =============================================================================
section_id = 15
section_name = get_section_name(section_id)
timestamp = datetime.now()
column_order = ['Lotnummer', 'Pallenummer', 'Produktionstidspunkt', 'Kontrolleret af',
                'Kontrol bemærkning', 'Kontroltidspunkt', 'Kilo', 'Antal poser',
                'Antal leakers', 'Leakers pct', 'Resultat af kontrol']

if get_section_status_code(df_karakterer, get_section_visibility(df_sections, section_id)) == 99:
    try:
        df_nav_lotno = pd.merge(df_nav_lotno, df_ds_vacslip, left_on = 'Lotnummer',
                                right_on = 'Lotnummer', how='left', suffixes=('', '_y'))
        df_nav_lotno['Resultat af kontrol'].fillna(value='Ej kontrolleret', inplace=True)
        df_nav_lotno['Leakers pct'] = df_nav_lotno['Antal leakers'] / df_nav_lotno['Antal poser']
        df_nav_lotno['Pallenummer'] = df_nav_lotno['Pallenummer_y'].fillna(df_nav_lotno['Pallenummer'])
        df_nav_lotno['Produktionstidspunkt'] = df_nav_lotno['Produktionstidspunkt'].dt.strftime('%d-%m-%Y %H:%M')
        df_nav_lotno = df_nav_lotno[column_order]
        # Write results to Word and Excel
        insert_dataframe_into_excel (df_nav_lotno, section_name, False)
        add_section_to_word(df_nav_lotno, section_name, True, 'Landscape')
        # Write status into log
        section_log_insert(timestamp, section_id, 0)
    except: # Insert error into log
        section_log_insert(timestamp, section_id, 2)
else: # Write into log if no data is found or section is out of scope
    section_log_insert(timestamp, section_id, get_section_status_code(df_nav_lotno, get_section_visibility(df_sections, section_id)))


# =============================================================================
# Section 16: Reference- og henstandsprøver
# =============================================================================
section_id = 16
section_name = get_section_name(section_id)
timestamp = datetime.now()
column_order = ['Id', 'Registreringstidspunkt', 'Operatør', 'Silo', 'Prøvetype',
                'Bemærkning', 'Smagning status', 'Antal prøver']
df_temp = df_prøver[df_prøver['Prøvetype int'] != 0]

if get_section_status_code(df_temp, get_section_visibility(df_sections, section_id)) == 99:
    try:
        df_temp['Registreringstidspunkt'] = df_temp['Registreringstidspunkt'].dt.strftime('%d-%m-%Y %H:%M')
        df_temp = df_temp[column_order]
        # Write results to Word and Excel
        insert_dataframe_into_excel (df_temp, section_name, False)
        add_section_to_word(df_temp, section_name, True, 'Landscape')
        # Write status into log
        section_log_insert(timestamp, section_id, 0)
    except: # Insert error into log
        section_log_insert(timestamp, section_id, 2)
else: # Write into log if no data is found or section is out of scope
    section_log_insert(timestamp, section_id, get_section_status_code(df_temp, get_section_visibility(df_sections, section_id)))


# =============================================================================
# Section 17: Udtagne kontrolprøver
# =============================================================================
section_id = 17
section_name = get_section_name(section_id)
timestamp = datetime.now()
column_order = ['Id','Registreringstidspunkt', 'Operatør', 'Bemærkning',
                'Mærkning', 'Rygsvejsning', 'Tæthed', 'Ventil', 'Peelbar',
                'Tintie', 'Vægt', 'Ilt']
df_temp = df_prøver[df_prøver['Prøvetype int'] == 0]

if get_section_status_code(df_temp, get_section_visibility(df_sections, section_id)) == 99:
    try:
        df_temp['Registreringstidspunkt'] = df_temp['Registreringstidspunkt'].dt.strftime('%d-%m-%Y %H:%M')
        df_temp = df_temp[column_order]
        # Write results to Word and Excel
        insert_dataframe_into_excel (df_temp, section_name, False)
        add_section_to_word(df_temp, section_name, True, 'Landscape')
        # Write status into log
        section_log_insert(timestamp, section_id, 0)
    except: # Insert error into log
        section_log_insert(timestamp, section_id, 2)
else: # Write into log if no data is found or section is out of scope
    section_log_insert(timestamp, section_id, get_section_status_code(df_temp, get_section_visibility(df_sections, section_id)))


# =============================================================================
# Section 18: Sektionslog
# =============================================================================
section_id = 18
df_section_log = pd.read_sql(query_ds_section_log, con_04)
section_name = get_section_name(section_id)
timestamp = datetime.now()

if get_section_status_code(df_section_log, get_section_visibility(df_sections, section_id)) == 99:
    try:
        # Write results to Word and Excel
        insert_dataframe_into_excel (df_section_log, section_name, False)
        add_section_to_word(df_section_log, section_name, True, 'Landscape')
        # Write status into log
        section_log_insert(timestamp, section_id, 0)
    except: # Insert error into log
        section_log_insert(timestamp, section_id, 2)
else: # Write into log if no data is found or section is out of scope
    section_log_insert(timestamp, section_id, get_section_status_code(df_section_log, get_section_visibility(df_sections, section_id)))


#Save files
excel_writer.save()
log_insert(script_name, f'Excel file {file_name} created')

doc.save(path_file_doc)
log_insert(script_name, f'Word document {file_name} created')
# *** TODO SAVE PDF FILE

# =============================================================================
# Write into email log
# =============================================================================
dict_email_log = {'Filsti': filepath
                  ,'Filnavn': file_name
                  ,'Modtager': req_recipients
                  ,'Emne': f'Anmodet rapport for ordre {req_order_no}'
                  ,'Forespørgsels_id': req_id
                  ,'Note':req_note}
# pd.DataFrame(data=dict_email_log, index=[0]).to_sql('Sporbarhed_email_log', con=engine_04, schema='trc', if_exists='append', index=False)
log_insert(script_name, f'Request id: {req_id} inserted into [trc].[Email_log]')

# =============================================================================
# Update request that dataprocessing has been completed
# =============================================================================
cursor_04.execute(f"""UPDATE [trc].[Sporbarhed_forespørgsel]
                  SET Data_færdigbehandlet = 1
                  WHERE [Id] = {req_id}""")
cursor_04.commit()
log_insert(script_name, f'Request id: {req_id} completed')
